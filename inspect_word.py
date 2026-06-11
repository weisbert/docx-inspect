#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
inspect_word.py — dump the structure of a Word (.docx) file to plain text.

Reports document defaults, page setup, style definitions (including East Asian
fonts), numbering/list definitions, headers & footers, the body outline
(paragraphs and tables in order), per-table layout (column widths, borders,
cell text), and image dimensions.

Handy for understanding how a document is formatted, debugging python-docx
output, or reproducing a layout in a template.

Usage:
    python inspect_word.py INPUT.docx [OUTPUT.txt]
    With no output path, writes "<input>_inspect.txt" next to the input.

Requires:
    pip install python-docx

Notes:
    - .docx only. For the legacy .doc format, re-save as .docx first.
"""

import sys
import os

try:
    import docx
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
except ImportError:
    print("Missing dependency python-docx. Install it with:  pip install python-docx")
    sys.exit(1)


# --------------------------------------------------------------------------
# Output buffer helpers
# --------------------------------------------------------------------------
_LINES = []


def emit(line=""):
    _LINES.append(str(line))


def hr(char="-", n=60):
    emit(char * n)


def section(title):
    emit()
    emit("=" * 60)
    emit("# " + title)
    emit("=" * 60)


# --------------------------------------------------------------------------
# Unit conversions
# --------------------------------------------------------------------------
def twips_to_cm(v):
    # Page units are twips = 1/1440 inch ; 1 inch = 2.54 cm
    try:
        return round(int(v) / 1440.0 * 2.54, 2)
    except (TypeError, ValueError):
        return None


def emu_to_cm(v):
    # Image units are EMU ; 1 cm = 360000 EMU
    try:
        return round(int(v) / 360000.0, 2)
    except (TypeError, ValueError):
        return None


def halfpt_to_pt(v):
    try:
        return int(v) / 2.0
    except (TypeError, ValueError):
        return None


# --------------------------------------------------------------------------
# Run properties (rPr): font, size, bold, color, etc.
# Note: for CJK documents we also read w:eastAsia, which python-docx's
# font.name does not expose (it only returns w:ascii).
# --------------------------------------------------------------------------
def describe_rpr(rPr):
    if rPr is None:
        return {}
    info = {}
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        for label, attr in (
            ("ascii", "w:ascii"),
            ("hAnsi", "w:hAnsi"),
            ("eastAsia", "w:eastAsia"),
            ("cs", "w:cs"),
        ):
            val = rFonts.get(qn(attr))
            if val:
                info[label] = val

    sz = rPr.find(qn("w:sz"))
    if sz is not None:
        pt = halfpt_to_pt(sz.get(qn("w:val")))
        if pt is not None:
            info["size_pt"] = pt

    # A toggle like <w:b/> means on; <w:b w:val="0"/> means explicitly off.
    for tag, label in (("w:b", "bold"), ("w:i", "italic")):
        el = rPr.find(qn(tag))
        if el is not None:
            val = el.get(qn("w:val"))
            info[label] = False if val in ("0", "false", "off") else True
    u = rPr.find(qn("w:u"))
    if u is not None:
        uval = u.get(qn("w:val")) or "single"
        if uval != "none":
            info["underline"] = uval

    color = rPr.find(qn("w:color"))
    if color is not None:
        val = color.get(qn("w:val"))
        if val and val != "auto":
            info["color"] = "#" + val

    highlight = rPr.find(qn("w:highlight"))
    if highlight is not None:
        info["highlight"] = highlight.get(qn("w:val"))

    return info


def fmt_dict(d):
    if not d:
        return "(none)"
    return ", ".join(f"{k}={v}" for k, v in d.items())


# --------------------------------------------------------------------------
# Paragraph properties (pPr): alignment, indent, spacing, numbering
# --------------------------------------------------------------------------
def describe_ppr(pPr):
    if pPr is None:
        return {}
    info = {}
    jc = pPr.find(qn("w:jc"))
    if jc is not None:
        info["align"] = jc.get(qn("w:val"))

    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        for label, attr in (
            ("indent_left", "w:left"),
            ("indent_right", "w:right"),
            ("first_line", "w:firstLine"),
            ("hanging", "w:hanging"),
        ):
            val = ind.get(qn(attr))
            if val:
                cm = twips_to_cm(val)
                info[label] = f"{cm}cm" if cm is not None else val

    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        before = spacing.get(qn("w:before"))
        after = spacing.get(qn("w:after"))
        line = spacing.get(qn("w:line"))
        line_rule = spacing.get(qn("w:lineRule"))
        if before:
            info["space_before"] = f"{round(int(before)/20,1)}pt"
        if after:
            info["space_after"] = f"{round(int(after)/20,1)}pt"
        if line:
            if line_rule in (None, "auto"):
                info["line"] = f"{round(int(line)/240,2)}x"
            else:
                info["line"] = f"{round(int(line)/20,1)}pt({line_rule})"

    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        ilvl = numPr.find(qn("w:ilvl"))
        numId = numPr.find(qn("w:numId"))
        info["numbering"] = (
            f"numId={numId.get(qn('w:val')) if numId is not None else '?'}"
            f", level={ilvl.get(qn('w:val')) if ilvl is not None else '0'}"
        )

    outline = pPr.find(qn("w:outlineLvl"))
    if outline is not None:
        info["outline_level"] = outline.get(qn("w:val"))

    # paragraph borders — e.g. the rule line under a heading is a bottom border,
    # not a font underline. (Use probe_style.py for the full border XML.)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is not None:
        edges = []
        for edge in ("top", "bottom", "left", "right"):
            b = pBdr.find(qn("w:" + edge))
            if b is not None:
                edges.append(f"{edge}={b.get(qn('w:val'))}/{b.get(qn('w:sz'))}/{b.get(qn('w:color'))}")
        if edges:
            info["border"] = "[" + " ".join(edges) + "]"

    shd = pPr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill != "auto":
            info["shading"] = "#" + fill

    return info


# --------------------------------------------------------------------------
# 1. Document defaults (font / size)
# --------------------------------------------------------------------------
def dump_defaults(document):
    section("1. Document Defaults")
    try:
        styles_el = document.styles.element
        docDefaults = styles_el.find(qn("w:docDefaults"))
        if docDefaults is None:
            emit("(no docDefaults found)")
            return
        rPrDefault = docDefaults.find(qn("w:rPrDefault"))
        if rPrDefault is not None:
            rPr = rPrDefault.find(qn("w:rPr"))
            emit("Default run properties: " + fmt_dict(describe_rpr(rPr)))
        pPrDefault = docDefaults.find(qn("w:pPrDefault"))
        if pPrDefault is not None:
            pPr = pPrDefault.find(qn("w:pPr"))
            emit("Default paragraph properties: " + fmt_dict(describe_ppr(pPr)))
    except Exception as e:
        emit(f"(error parsing defaults: {e})")


# --------------------------------------------------------------------------
# 2. Page / section setup
# --------------------------------------------------------------------------
def dump_sections(document):
    section("2. Page Setup (Sections)")
    for i, sec in enumerate(document.sections):
        emit(f"[Section {i+1}]")
        try:
            emit(f"  Paper: width {emu_to_cm(sec.page_width)}cm  height {emu_to_cm(sec.page_height)}cm")
            emit(f"  Orientation: {sec.orientation}")
            emit(
                f"  Margins: top {emu_to_cm(sec.top_margin)}  bottom {emu_to_cm(sec.bottom_margin)}"
                f"  left {emu_to_cm(sec.left_margin)}  right {emu_to_cm(sec.right_margin)} (cm)"
            )
            emit(
                f"  Header distance {emu_to_cm(sec.header_distance)}cm  footer distance {emu_to_cm(sec.footer_distance)}cm"
            )
        except Exception as e:
            emit(f"  (error parsing section: {e})")
        emit()


# --------------------------------------------------------------------------
# 3. Styles (styles.xml)
# --------------------------------------------------------------------------
STYLE_TYPE = {
    "paragraph": "paragraph",
    "character": "character",
    "table": "table",
    "numbering": "numbering",
}


def collect_used_styles(document):
    """Count paragraph/table styles actually used in the body (incl. table cells)."""
    para_styles = {}
    table_styles = {}
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, document)
            name = p.style.name if p.style is not None else None
            if name:
                para_styles[name] = para_styles.get(name, 0) + 1
        elif child.tag == qn("w:tbl"):
            t = Table(child, document)
            name = t.style.name if t.style is not None else None
            if name:
                table_styles[name] = table_styles.get(name, 0) + 1
    for t in document.tables:
        try:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        name = p.style.name if p.style is not None else None
                        if name:
                            para_styles[name] = para_styles.get(name, 0) + 1
        except Exception:
            pass
    return para_styles, table_styles


def dump_styles(document, used_para=None, used_table=None):
    used_para = used_para or {}
    used_table = used_table or {}
    used_names = set(used_para) | set(used_table)

    section("3. Styles")
    emit("Each style records its font / size / alignment / spacing.")
    emit("Markers: *=used in body  [custom]=user-defined style  *default*=default for its type")
    emit()
    emit("--- Styles actually used in the body ---")
    if used_para:
        for n, c in sorted(used_para.items(), key=lambda x: -x[1]):
            emit(f"    paragraph  {n}   (x{c})")
    if used_table:
        for n, c in sorted(used_table.items(), key=lambda x: -x[1]):
            emit(f"    table      {n}   (x{c})")
    if not used_names:
        emit("    (none detected; see full list below)")
    emit()
    emit("--- Full style definitions ---")

    styles_el = document.styles.element
    for style_el in styles_el.findall(qn("w:style")):
        try:
            stype = style_el.get(qn("w:type"))
            styleId = style_el.get(qn("w:styleId"))
            default = style_el.get(qn("w:default"))
            custom = style_el.get(qn("w:customStyle"))
            name_el = style_el.find(qn("w:name"))
            name = name_el.get(qn("w:val")) if name_el is not None else styleId
            based_el = style_el.find(qn("w:basedOn"))
            based = based_el.get(qn("w:val")) if based_el is not None else None

            mark = "* " if name in used_names else "  "
            tag = f"{mark}[{STYLE_TYPE.get(stype, stype)}] {name}"
            if styleId and styleId != name:
                tag += f"  (id={styleId})"
            if custom == "1":
                tag += "  [custom]"
            if default == "1":
                tag += "  *default*"
            if based:
                tag += f"  basedOn<{based}>"
            emit(tag)

            rPr = style_el.find(qn("w:rPr"))
            rinfo = describe_rpr(rPr)
            if rinfo:
                emit("    font: " + fmt_dict(rinfo))
            pPr = style_el.find(qn("w:pPr"))
            pinfo = describe_ppr(pPr)
            if pinfo:
                emit("    para: " + fmt_dict(pinfo))
        except Exception as e:
            emit(f"    (error parsing a style: {e})")
    emit()


# --------------------------------------------------------------------------
# 4. Numbering definitions (numbering.xml)
# --------------------------------------------------------------------------
def dump_numbering(document):
    section("4. Numbering / List Definitions")
    try:
        numbering_part = document.part.numbering_part
    except (NotImplementedError, KeyError, AttributeError):
        emit("(no numbering definitions)")
        return
    if numbering_part is None:
        emit("(no numbering definitions)")
        return
    root = numbering_part.element

    # abstractNum: the actual format definitions
    for absnum in root.findall(qn("w:abstractNum")):
        aid = absnum.get(qn("w:abstractNumId"))
        emit(f"[abstractNum id={aid}]")
        for lvl in absnum.findall(qn("w:lvl")):
            ilvl = lvl.get(qn("w:ilvl"))
            numFmt = lvl.find(qn("w:numFmt"))
            lvlText = lvl.find(qn("w:lvlText"))
            fmt = numFmt.get(qn("w:val")) if numFmt is not None else "?"
            text = lvlText.get(qn("w:val")) if lvlText is not None else "?"
            emit(f"    level {ilvl}: format={fmt}  text='{text}'")
    # num: numId -> abstractNumId mapping
    emit()
    for num in root.findall(qn("w:num")):
        nid = num.get(qn("w:numId"))
        absref = num.find(qn("w:abstractNumId"))
        absid = absref.get(qn("w:val")) if absref is not None else "?"
        emit(f"numId {nid} -> abstractNum {absid}")
    emit()


# --------------------------------------------------------------------------
# 5. Headers / footers
# --------------------------------------------------------------------------
def dump_headers_footers(document):
    section("5. Headers & Footers")
    for i, sec in enumerate(document.sections):
        emit(f"[Section {i+1}]")
        for label, hf in (
            ("header", sec.header),
            ("first-page header", sec.first_page_header),
            ("even-page header", sec.even_page_header),
            ("footer", sec.footer),
            ("first-page footer", sec.first_page_footer),
            ("even-page footer", sec.even_page_footer),
        ):
            try:
                if hf is None:
                    continue
                linked = getattr(hf, "is_linked_to_previous", None)
                texts = [p.text for p in hf.paragraphs if p.text.strip()]
                if not texts and linked:
                    continue  # linked to previous and empty -> skip
                emit(f"  {label} (linked_to_previous={linked}):")
                for t in texts:
                    emit(f"      | {t}")
                if not texts:
                    emit("      | (no text; may contain a page-number field or image)")
            except Exception as e:
                emit(f"  ({label} parse error: {e})")
        emit()


# --------------------------------------------------------------------------
# 6. Body outline (paragraphs and tables in order)
# --------------------------------------------------------------------------
def dump_body(document, max_text=120):
    section("6. Body Outline (paragraphs + tables in order)")
    emit("[P]=paragraph  [T]=table. Long text is truncated. Style name in <>.")
    emit()
    body = document.element.body
    table_index = 0
    para_index = 0
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            para_index += 1
            p = Paragraph(child, document)
            style = p.style.name if p.style is not None else "?"
            text = p.text.replace("\n", " ")
            if len(text) > max_text:
                text = text[:max_text] + " ..."
            line = f"[P{para_index}] <{style}>"
            ppr_info = describe_ppr(child.find(qn("w:pPr")))
            key = {k: v for k, v in ppr_info.items() if k in ("align", "numbering", "outline_level")}
            if key:
                line += " " + fmt_dict(key)
            if text.strip():
                line += f"  : {text}"
            else:
                line += "  : (empty paragraph)"
            emit(line)
            if child.findall(".//" + qn("w:drawing")):
                emit("        >> contains an image/drawing")
        elif tag == qn("w:tbl"):
            table_index += 1
            emit(f"[T{table_index}] === table, see section 7 ===")
    emit()
    emit(f"Totals: {para_index} paragraphs, {table_index} tables.")
    emit()


# --------------------------------------------------------------------------
# 7. Tables
# --------------------------------------------------------------------------
def describe_table_borders(tbl_el):
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        return "(no tblPr)"
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        return "(not set explicitly; inherits from style)"
    out = []
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.find(qn("w:" + edge))
        if b is not None:
            val = b.get(qn("w:val"))
            sz = b.get(qn("w:sz"))
            color = b.get(qn("w:color"))
            out.append(f"{edge}={val}/{sz}/{color}")
    return ", ".join(out) if out else "(none)"


def dump_tables(document, max_cell=40, max_rows_detail=8):
    section("7. Tables")
    if not document.tables:
        emit("(no tables)")
        emit()
        return
    for ti, table in enumerate(document.tables, start=1):
        tbl_el = table._tbl
        emit(f"[Table {ti}]")
        try:
            n_rows = len(table.rows)
            n_cols = len(table.columns)
            emit(f"  Size: {n_rows} rows x {n_cols} cols")

            style = table.style.name if table.style is not None else "?"
            emit(f"  Table style: {style}")

            grid = tbl_el.find(qn("w:tblGrid"))
            if grid is not None:
                widths = []
                for col in grid.findall(qn("w:gridCol")):
                    w = col.get(qn("w:w"))
                    cm = twips_to_cm(w)
                    widths.append(f"{cm}cm" if cm is not None else w)
                emit(f"  Column widths: {widths}")

            emit(f"  Borders: {describe_table_borders(tbl_el)}")

            tblPr = tbl_el.find(qn("w:tblPr"))
            if tblPr is not None:
                tblW = tblPr.find(qn("w:tblW"))
                if tblW is not None:
                    w = tblW.get(qn("w:w"))
                    wtype = tblW.get(qn("w:type"))
                    if wtype == "pct":
                        emit(f"  Table width: {int(w)/50.0}% (relative)")
                    elif wtype == "dxa":
                        emit(f"  Table width: {twips_to_cm(w)}cm")

            emit("  Cells:")
            for ri, row in enumerate(table.rows):
                if ri >= max_rows_detail:
                    emit(f"    ... ({n_rows - max_rows_detail} more rows omitted)")
                    break
                cells_text = []
                seen_tc = set()
                for cell in row.cells:
                    # merged cells return the same tc repeatedly; de-dupe
                    tc_id = id(cell._tc)
                    if tc_id in seen_tc:
                        cells_text.append("<<merged>>")
                        continue
                    seen_tc.add(tc_id)
                    txt = cell.text.replace("\n", " ").strip()
                    if len(txt) > max_cell:
                        txt = txt[:max_cell] + "..."
                    cells_text.append(txt if txt else ".")
                emit(f"    row{ri+1}: " + " | ".join(cells_text))

            # first-row shading (likely a header)
            try:
                first_row = table.rows[0]
                shades = []
                for cell in first_row.cells:
                    tcPr = cell._tc.find(qn("w:tcPr"))
                    if tcPr is not None:
                        shd = tcPr.find(qn("w:shd"))
                        if shd is not None:
                            fill = shd.get(qn("w:fill"))
                            if fill and fill != "auto":
                                shades.append("#" + fill)
                if shades:
                    emit(f"  First-row fill: {set(shades)}  (likely a header)")
            except Exception:
                pass

        except Exception as e:
            emit(f"  (error parsing table: {e})")
        emit()


# --------------------------------------------------------------------------
# 8. Images / drawings
# --------------------------------------------------------------------------
def dump_images(document):
    section("8. Images")
    try:
        shapes = document.inline_shapes
        emit(f"Inline images: {len(shapes)}")
        for i, sh in enumerate(shapes, start=1):
            try:
                w = emu_to_cm(sh.width)
                h = emu_to_cm(sh.height)
                emit(f"  image {i}: ~{w}cm x {h}cm  type={sh.type}")
            except Exception:
                emit(f"  image {i}: (could not read size)")
    except Exception as e:
        emit(f"(error reading inline images: {e})")

    try:
        img_parts = [
            p for p in document.part.package.iter_parts()
            if "image" in (p.content_type or "")
        ]
        emit(f"Total embedded image files (incl. floating/header): {len(img_parts)}")
    except Exception:
        pass
    emit()
    emit("Note: image positions and sizes are listed; binary image data is not included.")
    emit()


# --------------------------------------------------------------------------
# main
# --------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print(__doc__)
        print("\nError: please provide an input .docx path.")
        sys.exit(1)

    in_path = sys.argv[1]
    if not os.path.isfile(in_path):
        print(f"File not found: {in_path}")
        sys.exit(1)
    if in_path.lower().endswith(".doc") and not in_path.lower().endswith(".docx"):
        print("Legacy .doc detected. Re-save as .docx in Word first.")
        sys.exit(1)

    if len(sys.argv) >= 3:
        out_path = sys.argv[2]
    else:
        base = os.path.splitext(in_path)[0]
        out_path = base + "_inspect.txt"

    try:
        document = Document(in_path)
    except Exception as e:
        print(f"Failed to open document: {e}")
        sys.exit(1)

    emit("=" * 60)
    emit("DOCX STRUCTURE REPORT  (inspect_word.py)")
    emit("=" * 60)
    emit(f"Source: {os.path.abspath(in_path)}")

    dump_defaults(document)
    dump_sections(document)
    used_para, used_table = collect_used_styles(document)
    dump_styles(document, used_para, used_table)
    dump_numbering(document)
    dump_headers_footers(document)
    dump_body(document)
    dump_tables(document)
    dump_images(document)

    emit("=" * 60)
    emit("END OF REPORT")
    emit("=" * 60)

    text = "\n".join(_LINES)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"Done. Structure report written to:\n  {os.path.abspath(out_path)}")
    print(f"{len(_LINES)} lines. Open it in any text editor.")


if __name__ == "__main__":
    main()
