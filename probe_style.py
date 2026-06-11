#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
probe_style.py — deep-inspect a .docx: dump the FULL raw XML of styles or
paragraphs, so nothing is missed (paragraph borders / shading / tab stops /
run-level character formatting that a summary view skips).

Use it when a high-level inspector doesn't show enough — e.g. the underline
rule under a heading is a paragraph BOTTOM BORDER (w:pBdr), not a font
underline, and only shows up in the raw XML.

Usage:
    # 1) list every style (id | name | type)
    python probe_style.py file.docx

    # 2) dump one or more styles in full (resolves the basedOn chain too)
    python probe_style.py file.docx style "Heading 1" "heading 1" "Caption"

    # 3) find paragraphs containing some text and dump their full XML
    #    (reveals run-level / per-character formatting)
    python probe_style.py file.docx find "Summary"

    # 4) dump every RUN (per-character font / size / color) of the paragraphs
    #    between two anchor texts — shows which words are coloured red and the
    #    per-paragraph fonts. End anchor is optional (defaults to a few paras).
    python probe_style.py file.docx runs "Introduction" "Scope"

    # 5) dump every section's header & footer: full XML + any embedded logo
    #    (file name / size in cm / position) — to mirror a header/footer
    python probe_style.py file.docx headers

    # 6) dump raw formatting PARTS verbatim (styles.xml / numbering.xml /
    #    settings.xml) — to reuse the real formatting byte-for-byte instead of
    #    re-deriving (and silently missing) properties.
    python probe_style.py file.docx parts

    # 7) dump the full XML of tables (per-cell borders / shading / margins /
    #    alignment). Optional N limits to the first N tables (cover = 1-3).
    python probe_style.py file.docx tables 3

Output is printed AND written to "<docx>_probe.txt".

Requires:  pip install python-docx lxml
"""

import sys
import os

try:
    from docx import Document
    from docx.oxml.ns import qn, nsmap
    from docx.text.paragraph import Paragraph
    from lxml import etree
except ImportError:
    print("Missing dependency. Install with:  pip install python-docx lxml")
    sys.exit(1)

OUT = []


def emit(s=""):
    OUT.append(str(s))


def pretty(el):
    """Pretty-print an lxml element, keeping the w: prefix readable."""
    xml = etree.tostring(el, pretty_print=True, encoding="unicode")
    # collapse the long namespace declarations on the root tag for readability
    return xml.strip()


W = qn("w:x").split("}")[0][1:]  # the w namespace URI


def style_name(style_el):
    n = style_el.find(qn("w:name"))
    sid = style_el.get(qn("w:styleId"))
    return (n.get(qn("w:val")) if n is not None else sid), sid


def norm(s):
    return (s or "").replace(" ", "").lower()


def find_styles(document, wanted):
    """Match each wanted name against w:name OR w:styleId (case/space-insensitive)."""
    styles_el = document.styles.element
    result = []
    wn = [norm(w) for w in wanted]
    for st in styles_el.findall(qn("w:style")):
        name, sid = style_name(st)
        if norm(name) in wn or norm(sid) in wn:
            result.append(st)
    return result


def basedon_chain(document, style_el):
    """Yield the style and each ancestor it is basedOn, in order."""
    styles_el = document.styles.element
    by_id = {st.get(qn("w:styleId")): st for st in styles_el.findall(qn("w:style"))}
    seen = set()
    cur = style_el
    while cur is not None:
        sid = cur.get(qn("w:styleId"))
        if sid in seen:
            break
        seen.add(sid)
        yield cur
        b = cur.find(qn("w:basedOn"))
        cur = by_id.get(b.get(qn("w:val"))) if b is not None else None


def readable_summary(el):
    """A short human summary of the formatting bits people most often miss."""
    out = []
    pPr = el.find(qn("w:pPr"))
    rPr = el.find(qn("w:rPr"))
    # paragraph border (THIS is the heading underline)
    if pPr is not None:
        pbdr = pPr.find(qn("w:pBdr"))
        if pbdr is not None:
            edges = []
            for edge in ("top", "bottom", "left", "right"):
                b = pbdr.find(qn("w:" + edge))
                if b is not None:
                    edges.append(
                        f"{edge}:{b.get(qn('w:val'))}/sz{b.get(qn('w:sz'))}/"
                        f"space{b.get(qn('w:space'))}/{b.get(qn('w:color'))}"
                    )
            out.append("  paragraph border pBdr -> " + "; ".join(edges))
        shd = pPr.find(qn("w:shd"))
        if shd is not None:
            out.append(f"  paragraph shading shd -> fill={shd.get(qn('w:fill'))}")
        tabs = pPr.find(qn("w:tabs"))
        if tabs is not None:
            ts = [f"{t.get(qn('w:val'))}@{t.get(qn('w:pos'))}" for t in tabs.findall(qn("w:tab"))]
            out.append("  tab stops tabs -> " + ", ".join(ts))
    if rPr is not None:
        shd = rPr.find(qn("w:shd"))
        if shd is not None:
            out.append(f"  run shading -> fill={shd.get(qn('w:fill'))}")
    return out


def describe_run(r_el):
    """Per-run (per-character) formatting overrides from w:rPr."""
    info = {}
    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        return info
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        asc = rFonts.get(qn("w:ascii"))
        ea = rFonts.get(qn("w:eastAsia"))
        if asc:
            info["ascii"] = asc
        if ea:
            info["eastAsia"] = ea
    sz = rPr.find(qn("w:sz"))
    if sz is not None:
        try:
            info["sz_pt"] = int(sz.get(qn("w:val"))) / 2
        except (TypeError, ValueError):
            pass
    for tag, label in (("w:b", "bold"), ("w:i", "italic")):
        el = rPr.find(qn(tag))
        if el is not None:
            v = el.get(qn("w:val"))
            info[label] = False if v in ("0", "false", "off") else True
    u = rPr.find(qn("w:u"))
    if u is not None and (u.get(qn("w:val")) or "single") != "none":
        info["underline"] = u.get(qn("w:val")) or "single"
    col = rPr.find(qn("w:color"))
    if col is not None:
        v = col.get(qn("w:val"))
        if v and v != "auto":
            info["color"] = "#" + v
    hl = rPr.find(qn("w:highlight"))
    if hl is not None:
        info["highlight"] = hl.get(qn("w:val"))
    return info


def is_red(hexstr):
    try:
        h = hexstr.lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return r >= 128 and g < 110 and b < 110
    except Exception:
        return False


def dump_part_images(el, part):
    """List the images inside a header/footer element: file / size(cm) / rId."""
    lines = []
    for dr in el.findall(".//" + qn("w:drawing")):
        ext = dr.find(".//" + qn("wp:extent"))
        cx = ext.get("cx") if ext is not None else None
        cy = ext.get("cy") if ext is not None else None
        blip = dr.find(".//" + qn("a:blip"))
        rid = blip.get(qn("r:embed")) if blip is not None else None
        target = "?"
        if rid:
            try:
                target = part.rels[rid].target_ref
            except Exception:
                target = "(unresolved)"
        anchored = dr.find(".//" + qn("wp:anchor")) is not None
        w = round(int(cx) / 360000, 2) if cx else "?"
        h = round(int(cy) / 360000, 2) if cy else "?"
        kind = "floating" if anchored else "inline"
        lines.append(f"  [image] {target}  {w}cm x {h}cm  ({kind}, rId={rid})")
    return lines


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    path = sys.argv[1]
    if not os.path.isfile(path):
        print("File not found:", path)
        sys.exit(1)
    doc = Document(path)

    mode = sys.argv[2] if len(sys.argv) > 2 else "list"

    if mode == "list" or len(sys.argv) == 2:
        emit("==== ALL STYLES (id | name | type) ====")
        for st in doc.styles.element.findall(qn("w:style")):
            name, sid = style_name(st)
            emit(f"{st.get(qn('w:type')):10} | id={sid:28} | {name}")

    elif mode == "style":
        wanted = sys.argv[3:]
        if not wanted:
            print('Give style names, e.g.:  probe_style.py file.docx style "Heading 1"')
            sys.exit(1)
        # always show docDefaults first
        dd = doc.styles.element.find(qn("w:docDefaults"))
        if dd is not None:
            emit("==== docDefaults ====")
            emit(pretty(dd))
            emit()
        for st in find_styles(doc, wanted):
            name, sid = style_name(st)
            emit("=" * 70)
            emit(f"STYLE: {name}   (id={sid})")
            emit("=" * 70)
            for anc in basedon_chain(doc, st):
                aname, asid = style_name(anc)
                emit(f"\n----- {aname} (id={asid}) -----")
                summ = readable_summary(anc)
                if summ:
                    emit("[key formatting summary]")
                    for s in summ:
                        emit(s)
                emit("[full XML]")
                emit(pretty(anc))
            emit()

    elif mode == "find":
        needle = sys.argv[3] if len(sys.argv) > 3 else ""
        if not needle:
            print('Give a text fragment, e.g.:  probe_style.py file.docx find "Summary"')
            sys.exit(1)
        count = 0
        for p in doc.paragraphs:
            if needle in p.text:
                count += 1
                emit("=" * 70)
                emit(f"PARAGRAPH (style={p.style.name if p.style else '?'}): {p.text[:80]}")
                emit("=" * 70)
                summ = readable_summary(p._p)
                if summ:
                    emit("[key formatting summary]")
                    for s in summ:
                        emit(s)
                emit("[full XML]")
                emit(pretty(p._p))
                emit()
                if count >= 20:
                    emit("... (stopped at 20 matches)")
                    break
        if count == 0:
            emit(f"(no paragraph contains: {needle})")

    elif mode == "runs":
        if len(sys.argv) < 4:
            print('Usage: probe_style.py file.docx runs "start text" ["end text"]')
            sys.exit(1)
        start = sys.argv[3]
        end = sys.argv[4] if len(sys.argv) > 4 else None
        collecting = False
        any_hit = False

        def _is_toc(name):
            n = (name or "").replace(" ", "").lower()
            return n.startswith("toc") or "tableoffigures" in n or n.startswith("content")

        for p in doc.paragraphs:
            # skip TOC / table-of-figures entries so the anchor matches the body
            if _is_toc(p.style.name if p.style else ""):
                continue
            if collecting and end and end in p.text:
                break
            if not collecting and start in p.text:
                collecting = True
            if not collecting:
                continue
            any_hit = True
            emit("-" * 70)
            emit(f"[{p.style.name if p.style else '?'}]  «{p.text[:80]}»")
            for s in readable_summary(p._p):
                emit(s)
            for r in p._p.findall(".//" + qn("w:r")):
                txt = "".join((t.text or "") for t in r.findall(qn("w:t")))
                if txt == "":
                    continue
                info = describe_run(r)
                flag = "   *red text" if ("color" in info and is_red(info["color"])) else ""
                emit(f"    run «{txt}»  {info if info else '(inherits paragraph style)'}{flag}")
        if not any_hit:
            emit(f"(no paragraph contains: {start})")

    elif mode == "headers":
        items = [
            ("default header", "header"),
            ("first-page header", "first_page_header"),
            ("even-page header", "even_page_header"),
            ("default footer", "footer"),
            ("first-page footer", "first_page_footer"),
            ("even-page footer", "even_page_footer"),
        ]
        for i, sec in enumerate(doc.sections):
            emit("=" * 70)
            emit(f"SECTION {i + 1}")
            emit("=" * 70)
            for label, attr in items:
                hf = getattr(sec, attr, None)
                if hf is None:
                    continue
                linked = getattr(hf, "is_linked_to_previous", None)
                el = hf._element
                imgs = dump_part_images(el, hf.part)
                texts = [p.text for p in hf.paragraphs if p.text.strip()]
                if linked and not imgs and not texts:
                    emit(f"\n----- {label}: (linked to previous, no own content) -----")
                    continue
                emit(f"\n----- {label}  (linked_to_previous={linked}) -----")
                for t in texts:
                    emit(f"  text: | {t}")
                for s in imgs:
                    emit(s)
                emit("[full XML]")
                emit(pretty(el))
            emit()

    elif mode == "usedstyles":
        import zipfile
        import re
        # find the styleIds actually referenced (body / tables / headers / footers)
        refs = set()
        with zipfile.ZipFile(path) as z:
            for n in z.namelist():
                if n.startswith("word/") and n.endswith(".xml") and n != "word/styles.xml":
                    try:
                        data = z.read(n).decode("utf-8")
                    except Exception:
                        continue
                    refs.update(re.findall(r'w:(?:pStyle|rStyle|tblStyle)\s+w:val="([^"]+)"', data))
        styles_el = doc.styles.element
        by_id = {}
        for s in styles_el.findall(qn("w:style")):
            by_id[s.get(qn("w:styleId"))] = s
        # expand the basedOn / link / next inheritance chain
        want = set()
        stack = list(refs)
        while stack:
            sid = stack.pop()
            if sid in want or sid not in by_id:
                continue
            want.add(sid)
            for tag in ("w:basedOn", "w:link", "w:next"):
                e = by_id[sid].find(qn(tag))
                if e is not None:
                    v = e.get(qn("w:val"))
                    if v and v not in want:
                        stack.append(v)
        emit(f"# styles used: {len(want)} / {len(by_id)} total")
        emit("# these <w:style> elements can be embedded verbatim into a template.\n")
        dd = styles_el.find(qn("w:docDefaults"))
        if dd is not None:
            emit("===== docDefaults =====")
            emit(pretty(dd))
            emit()
        for sid, s in by_id.items():  # keep document order
            if sid in want:
                emit(pretty(s))
                emit()

    elif mode == "parts":
        import zipfile
        name_map = {
            "styles": "word/styles.xml",
            "numbering": "word/numbering.xml",
            "settings": "word/settings.xml",
            "theme": "word/theme/theme1.xml",
        }
        which = sys.argv[3:] or ["styles", "numbering", "settings"]
        with zipfile.ZipFile(path) as z:
            for key in which:
                part = name_map.get(key, key)
                try:
                    data = z.read(part).decode("utf-8")
                except KeyError:
                    emit(f"\n(no such part: {part})")
                    continue
                emit("=" * 70)
                emit(f"PART: {part}   ({len(data)} bytes)")
                emit("=" * 70)
                emit(data)  # verbatim — this is the whole point
                emit()

    elif mode == "tables":
        limit = None
        if len(sys.argv) > 3:
            try:
                limit = int(sys.argv[3])
            except ValueError:
                limit = None
        for i, t in enumerate(doc.tables, 1):
            if limit and i > limit:
                break
            emit("=" * 70)
            emit(f"TABLE {i}   ({len(t.rows)} rows x {len(t.columns)} cols)")
            emit("=" * 70)
            emit(pretty(t._tbl))
            emit()

    else:
        print("Unknown mode:", mode)
        print(__doc__)
        sys.exit(1)

    text = "\n".join(OUT)
    out_path = os.path.splitext(path)[0] + "_probe.txt"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    # Console encoding on Windows is often GBK and can't encode the output;
    # the .txt (UTF-8) is the real deliverable, so just point at it.
    print(f"[written to] {out_path}  ({len(OUT)} lines)")
    print("Open it in a UTF-8 editor (VS Code / Notepad) to read the full output.")


if __name__ == "__main__":
    main()
