#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Regression tests for web/live_preview.py -- the resident Word page renderer.

live_preview keeps ONE Word process of its own alive and turns a single outline
section into page images (engine section-only .docx -> Word PDF -> PNG). Two
properties are worth locking down, and neither of them is the speed:

  A. It must never touch a copy of Word it did not start. The people using this
     tool keep their documents open, very likely with unsaved changes, so a
     module that swept processes by image name would destroy real work. The
     guard here snapshots every WINWORD process id BEFORE the module runs and
     proves, after the module has shut its own instance down, that every one of
     them is still alive and that the module's own id was never one of them.

  B. It must degrade, not explode, on a machine that cannot do this at all --
     no pywin32, no PyMuPDF, not Windows. Importing the module still has to
     work, and the failure has to arrive as LivePreviewUnavailable carrying a
     sentence a person can read.

What is asserted:

  1. Importing live_preview raises nothing even when the capability probe fails.
     Checked in a subprocess whose import machinery refuses pywin32 and PyMuPDF,
     so the assertion runs on THIS machine too, where both are installed.
  2. A section render returns at least one page, a plausible page size, and PNG
     bytes (the magic number is checked after base64 decoding).
  3. The pid file names an id that was NOT a Word process before the test began;
     after shutdown() that id is gone and every pre-existing Word process is
     still running. This is the "do not close the user's own Word" guard.
  4. Two renders in a row reuse the same instance (same pid) and the second is
     faster, which is the entire point of keeping Word resident.
  5. shutdown() is idempotent.
  6. A section fragment shows the numbers the WHOLE report gives it. Word
     recalculates SEQ / STYLEREF while it paginates, which would restart the
     figure counter inside a fragment, so live_preview freezes those fields to
     the numbers the renderer computed over the complete outline. Asserted
     against a full render of the same fixture -- and this one needs no Word,
     so it runs everywhere.
  7. The instance's process id is read from a window that is STILL OPEN, before
     the throwaway document is closed. Reading it afterwards asks Windows about
     a handle it has already taken back: the call then reports failure only in
     its thread id and leaves the process id slot UNINITIALISED, so a stranger's
     id -- most plausibly one of the user's own Word windows -- would be recorded
     as ours and terminated later by the idle watchdog. Driven with a fake COM
     object, so it runs everywhere; a control proves the probe fails against the
     close-first order.
  8. A recorded creation time must match the live one EXACTLY. A tolerance there
     only widens the band in which an unrelated process is taken for ours, and
     every process that comparison answers True about is a candidate for
     termination.

Tests 2 to 5 need Word, pywin32 and PyMuPDF. Without them the file SKIPS them
and still exits 0, after running 1, 6, 7 and 8.

The fixture is the neutral, ASCII-only project already used by the golden render
test: four chapters, images and a cross-reference, with the target section
sitting mid-document so a restarted counter would be visible. No company data is
read or written; everything happens in a throwaway folder under the system temp
directory.

Run:
    python builder/tests/test_live_preview.py
"""

import base64
import json
import os
import shutil
import subprocess
import sys
import tempfile
import time

HERE = os.path.dirname(os.path.abspath(__file__))       # builder/tests
if HERE not in sys.path:
    sys.path.insert(0, HERE)
BUILDER = os.path.dirname(HERE)                          # builder/
if BUILDER not in sys.path:
    sys.path.insert(0, BUILDER)
import buildpath  # noqa: E402,F401  (registers core/docx_io/store/sync/web)

import live_preview  # noqa: E402  -- assertion 1: this import must never raise

# The fixture lives with the golden render test; reused rather than duplicated so
# the two tests cannot drift apart about what a renderable project looks like.
from test_render_golden import golden_config, section_only_project  # noqa: E402

#: A level-2 section in the middle of the fixture. It holds a figure, a table and
#: a paragraph referring to a figure in another chapter, so every numbering trap
#: named in the module docstring is exercised by rendering just this one node.
SECTION = "sec-delta-two"


def preview_config():
    """The fixture config as the server hands it over.

    ``engine._load_config`` resolves the logo to an absolute path and stores it
    under ``_logo_path``; a config built inline has to supply that key itself or
    the header builder raises. Empty means "no logo", which the fixture wants.
    """
    cfg = golden_config()
    cfg["_logo_path"] = ""
    return cfg


# ---------------------------------------------------------------------------
# Assertion harness
# ---------------------------------------------------------------------------
PASS = 0
FAIL = 0
SKIPPED = []


def check(cond, name, detail=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print("  ok   %s" % name)
    else:
        FAIL += 1
        print("  FAIL %s%s" % (name, ("  <- " + detail) if detail else ""))
    return bool(cond)


def skip(name, why):
    SKIPPED.append(name)
    print("  skip %s  (%s)" % (name, why))


# ---------------------------------------------------------------------------
# Process enumeration -- for the TEST only, deliberately independent of the
# module under test, so the guard cannot pass by agreeing with a broken module.
# ---------------------------------------------------------------------------
def _tasklist(filter_expr):
    """Raw `tasklist` rows matching one filter, as a list of CSV field lists."""
    if os.name != "nt":
        return []
    try:
        out = subprocess.run(
            ["tasklist", "/FI", filter_expr, "/NH", "/FO", "CSV"],
            capture_output=True, timeout=30).stdout
    except (OSError, subprocess.SubprocessError):
        return []
    rows = []
    for line in out.decode("ascii", "replace").splitlines():
        line = line.strip()
        if not line.startswith('"'):
            continue        # the "no tasks are running" notice, in any language
        rows.append([f.strip('"') for f in line.split('","')])
    return rows


def word_pids():
    """Every WINWORD process id on the machine right now."""
    found = set()
    for row in _tasklist("IMAGENAME eq WINWORD.EXE"):
        if len(row) >= 2 and row[1].isdigit():
            found.add(int(row[1]))
    return found


def pid_running(pid):
    """True when `pid` is a live process (any image)."""
    for row in _tasklist("PID eq %d" % int(pid)):
        if len(row) >= 2 and row[1].isdigit() and int(row[1]) == int(pid):
            return True
    return False


# ---------------------------------------------------------------------------
# Fixture on disk
# ---------------------------------------------------------------------------
def make_report(root):
    """Write the fixture project into a throwaway report folder."""
    os.makedirs(os.path.join(root, "images"), exist_ok=True)
    with open(os.path.join(root, "project.json"), "w", encoding="utf-8") as fh:
        json.dump(section_only_project(), fh, indent=1)
    return root


def pidfile_entries(project_dir):
    """The report's pid-file records, or None when the file does not exist."""
    path = os.path.join(project_dir, live_preview.PREVIEW_DIRNAME,
                        live_preview.PIDFILE_NAME)
    if not os.path.isfile(path):
        return None
    with open(path, "r", encoding="utf-8") as fh:
        data = json.load(fh)
    entries = data.get("entries") if isinstance(data, dict) else data
    return list(entries or [])


# ---------------------------------------------------------------------------
# 1. Import survives a failed capability probe
# ---------------------------------------------------------------------------
#: Run in a child interpreter whose import machinery refuses the optional
#: dependencies -- i.e. the machine this feature is NOT available on. argv[1] is
#: builder/, argv[2] a report folder.
_CAPFAIL = r'''
import sys

_BLOCKED = ("win32com", "win32api", "win32process", "pythoncom",
            "win32event", "fitz")


class _Refuse(object):
    def find_spec(self, name, path=None, target=None):
        if name.split(".")[0] in _BLOCKED:
            raise ImportError("blocked by the test: %s" % name)
        return None


sys.meta_path.insert(0, _Refuse())
sys.path.insert(0, sys.argv[1])
import buildpath
import live_preview                      # must not raise

assert live_preview.available() is False, "available() must report the failure"
st = live_preview.status()
assert st["available"] is False, st
assert isinstance(st["reason"], str) and len(st["reason"]) > 10, st
assert st["alive"] is False and st["pid"] is None, st

try:
    live_preview.render_section(sys.argv[2], None, "sec-delta-two")
except live_preview.LivePreviewUnavailable as exc:
    assert exc.reason and str(exc) == exc.reason, repr(exc)
else:
    raise AssertionError("render_section must refuse when it cannot render")

live_preview.shutdown()
live_preview.shutdown()                  # idempotent with nothing ever started
print("CAPFAIL-OK")
'''


def test_import_without_capability(report_dir):
    proc = subprocess.run([sys.executable, "-c", _CAPFAIL, BUILDER, report_dir],
                          capture_output=True, timeout=120)
    out = proc.stdout.decode("utf-8", "replace")
    err = proc.stderr.decode("utf-8", "replace")
    check(proc.returncode == 0 and "CAPFAIL-OK" in out,
          "import + refuse cleanly when pywin32 / PyMuPDF are missing",
          (err.strip().splitlines() or [""])[-1])


# ---------------------------------------------------------------------------
# 6. A fragment carries the whole report's numbers
# ---------------------------------------------------------------------------
def _captions(docx_path):
    """Caption paragraph texts, in order."""
    from docx import Document
    doc = Document(docx_path)
    out = []
    for para in doc.paragraphs:
        name = (para.style.name or "") if para.style is not None else ""
        if name.lower().startswith("caption"):
            out.append(para.text.strip())
    return out


def _para_starting(docx_path, prefix):
    from docx import Document
    for para in Document(docx_path).paragraphs:
        if para.text.strip().startswith(prefix):
            return para.text.strip()
    return None


def test_section_numbering(report_dir, tmp):
    """The fragment's captions and references must equal the full render's."""
    import engine

    project = section_only_project()
    cfg = preview_config()

    whole = os.path.join(tmp, "whole.docx")
    engine.render_report(project, cfg, report_dir, whole)
    part = os.path.join(tmp, "part.docx")
    engine.render_section_docx(project, cfg, report_dir, part, node_id=SECTION)

    frozen = live_preview._freeze_number_fields(part)
    check(frozen > 0, "the fragment's SEQ / STYLEREF / REF fields are frozen",
          "froze %r fields" % frozen)

    whole_caps = _captions(whole)
    part_caps = _captions(part)
    check(len(part_caps) >= 2,
          "the fragment still carries its figure and table captions",
          repr(part_caps))
    missing = [c for c in part_caps if c not in whole_caps]
    check(not missing,
          "every fragment caption reads exactly as it does in the full render",
          "not in the full render: %r (full: %r)" % (missing, whole_caps))
    # The trap this guards: a fragment renumbered from 1 would call a figure the
    # report knows as 4-3 "4-1". Assert it keeps its whole-document position.
    check(any(cap.endswith("Delta two figure") and "4-3" in cap
              for cap in part_caps),
          "the fragment's figure keeps its whole-document number (4-3)",
          repr(part_caps))

    whole_ref = _para_starting(whole, "See ")
    part_ref = _para_starting(part, "See ")
    check(part_ref is not None and part_ref == whole_ref,
          "a cross-reference inside the fragment reads as it does in the full "
          "render", "fragment=%r full=%r" % (part_ref, whole_ref))


# ---------------------------------------------------------------------------
# 7. The pid is read from a window that is still open
# ---------------------------------------------------------------------------
class _FakeWindow(object):
    """Stands in for Word's document window; its handle dies with the document."""

    def __init__(self):
        self.Hwnd = 0x4C12E0            # a plausible handle value
        self.alive = True


class _FakeDoc(object):
    def __init__(self, window):
        self.window = window
        self.closed = False

    def Close(self, save_changes):
        self.closed = True
        self.window.alive = False       # Windows may hand the handle out again


class _FakeDocuments(object):
    def __init__(self, app):
        self.app = app

    def Add(self):
        self.app.doc = _FakeDoc(self.app.window)
        return self.app.doc


class _FakeApp(object):
    """The two members live_preview touches while it works out its own pid."""

    def __init__(self):
        self.window = _FakeWindow()
        self.doc = None
        self.Documents = _FakeDocuments(self)

    @property
    def ActiveWindow(self):
        return self.window


def _pid_probe(app, record):
    """A stand-in for _window_pid that records the state it was called in."""
    def probe(hwnd):
        record["called"] = True
        record["hwnd"] = hwnd
        record["window_alive"] = app.window.alive
        record["doc_closed"] = app.doc.closed
        return 4242
    return probe


def _close_first(app):
    """The ordering this test exists to forbid -- kept only as a control.

    Close the throwaway document, then ask Windows which process the (already
    returned) handle belongs to.
    """
    blank = app.Documents.Add()
    try:
        hwnd = app.ActiveWindow.Hwnd
    finally:
        blank.Close(False)
    return live_preview._window_pid(hwnd)


def test_pid_read_before_close():
    """The pid must be taken while the window is open, and the probe must bite."""
    app = _FakeApp()
    record = {}
    original = live_preview._window_pid
    live_preview._window_pid = _pid_probe(app, record)
    try:
        pid = live_preview._identify_instance(app)
    finally:
        live_preview._window_pid = original

    check(record.get("called") is True and pid == 4242,
          "the instance's pid comes from the document window's handle",
          repr(record))
    check(record.get("window_alive") is True and record.get("doc_closed") is False,
          "the pid is read BEFORE the throwaway document is closed -- a handle "
          "already returned to Windows can name another process", repr(record))
    check(app.doc is not None and app.doc.closed is True,
          "the throwaway document is closed afterwards, not left open")

    # Control: the same probe, run against the close-first order, must object.
    # Without it the assertion above could pass by never being exercised.
    control_app = _FakeApp()
    control = {}
    live_preview._window_pid = _pid_probe(control_app, control)
    try:
        _close_first(control_app)
    finally:
        live_preview._window_pid = original
    check(control.get("doc_closed") is True and control.get("window_alive") is False,
          "the check has teeth: it catches the close-then-read order",
          repr(control))


def test_dead_handle_yields_no_pid():
    """A handle that names no window must produce None, never a number.

    The Win32 call signals failure with a thread id of 0 and does not write the
    process id at all, so pywin32 hands back whatever was in that slot -- an
    arbitrary value, not zero. Taking it would record a stranger as our Word.
    """
    if os.name != "nt":
        skip("dead window handle yields no pid", "needs Windows")
        return
    try:
        import win32process
    except ImportError as exc:
        skip("dead window handle yields no pid", "needs pywin32 (%s)" % exc)
        return
    dead = 0xDEADBEEF                   # no window has ever had this handle
    raw = win32process.GetWindowThreadProcessId(dead)
    check(live_preview._window_pid(dead) is None,
          "a dead window handle yields no pid at all",
          "the raw call answered thread=%r pid=%r" % (raw[0], raw[1]))


# ---------------------------------------------------------------------------
# 8. A recorded creation time must match exactly
# ---------------------------------------------------------------------------
def test_creation_time_match_is_exact():
    """(pid, creation time) is an identity, so the time may not be fuzzy."""
    if os.name != "nt":
        skip("creation time match is exact", "needs Windows")
        return
    me = os.getpid()
    created = live_preview._process_created(me)
    if created is None:
        check(False, "this process's creation time is readable")
        return
    check(live_preview._process_matches(me, created),
          "the exact recorded creation time identifies the process")
    check(live_preview._process_matches(me, json.loads(json.dumps(created))),
          "the creation time survives the pid file's JSON round trip intact")
    for drift in (0.5, -0.5, 0.001):
        check(not live_preview._process_matches(me, created + drift),
              "a creation time %+.3fs off is NOT accepted as the same process"
              % drift, "%r vs %r" % (created + drift, created))


# ---------------------------------------------------------------------------
# 2 to 5. The Word path
# ---------------------------------------------------------------------------
PNG_MAGIC = b"\x89PNG\r\n\x1a\n"


def check_page(page, label):
    ok = isinstance(page, dict) and "png_b64" in page
    if not check(ok, "%s is a page record" % label, repr(page)[:120]):
        return
    w, h = page.get("w"), page.get("h")
    check(isinstance(w, int) and isinstance(h, int)
          and 400 <= w <= 20000 and 400 <= h <= 20000 and h > w,
          "%s has a plausible portrait page size" % label, "%rx%r" % (w, h))
    try:
        raw = base64.b64decode(page["png_b64"], validate=True)
    except Exception as exc:
        check(False, "%s decodes as base64" % label, str(exc))
        return
    check(raw[:8] == PNG_MAGIC and len(raw) > 1000,
          "%s decodes to PNG bytes" % label,
          "%d bytes, header %r" % (len(raw), raw[:8]))


def test_word_path(report_dir, timings):
    """Renders twice, then shuts down, watching the process table throughout."""
    before = word_pids()
    print("  .. WINWORD processes before the test: %s"
          % (sorted(before) or "none"))

    cfg = preview_config()

    t0 = time.time()
    first = live_preview.render_section(report_dir, cfg, SECTION)
    wall_first = time.time() - t0

    check(first.get("scope") == "section",
          "the renderer produced a SECTION, not the whole document",
          "scope=%r" % first.get("scope"))

    pages = first.get("pages")
    check(isinstance(pages, list) and len(pages) >= 1,
          "the render returns at least one page",
          "pages=%r" % (len(pages) if isinstance(pages, list) else pages))
    if pages:
        check_page(pages[0], "page 1")

    # --- the pid file, while the instance is resident -----------------------
    pid = live_preview.status().get("pid")
    check(isinstance(pid, int) and pid > 0,
          "status() names the resident instance's pid", repr(pid))
    entries = pidfile_entries(report_dir)
    check(entries is not None, "the pid file exists while Word is resident")
    ours = [e for e in (entries or []) if int(e.get("pid", 0)) == (pid or -1)]
    check(len(ours) == 1, "the pid file records exactly our instance",
          repr(entries))
    check(bool(ours) and ours[0].get("host_pid") == os.getpid(),
          "the record names this process as its owner", repr(ours))

    # --- THE guard: our instance is not a Word the user had open -------------
    check(pid not in before,
          "the instance is a NEW process, not one that was already running",
          "pid=%r was in %r" % (pid, sorted(before)))
    check(pid_running(pid) if pid else False,
          "the recorded pid is a live process")
    # The recorded identity must be the real one, to the digit: sweep_stale,
    # _quit_word and shutdown all decide whether to terminate by comparing these.
    live_created = live_preview._process_created(pid) if pid else None
    check(bool(ours) and live_created is not None
          and ours[0].get("created") == live_created,
          "the pid file's creation time is EXACTLY the process's own",
          "recorded %r, live %r" % (ours[0].get("created") if ours else None,
                                    live_created))
    image = os.path.basename(live_preview._process_image(pid) or "") if pid else ""
    check(image.upper() == live_preview._WORD_IMAGE_NAME,
          "the recorded pid names a Word process, not some other program",
          "image=%r" % image)

    # --- residency: same instance, second render faster ---------------------
    t1 = time.time()
    second = live_preview.render_section(report_dir, cfg, SECTION)
    wall_second = time.time() - t1
    pid2 = live_preview.status().get("pid")
    check(pid2 == pid, "the second render reuses the same instance",
          "%r then %r" % (pid, pid2))
    check(second["ms"] < first["ms"],
          "the second render is faster than the first (Word stayed resident)",
          "%dms then %dms" % (first["ms"], second["ms"]))
    check(len(second.get("pages") or []) == len(pages or []),
          "both renders produce the same number of pages")

    timings.update({
        "cold_ms": first["ms"], "warm_ms": second["ms"],
        "cold_wall_ms": int(wall_first * 1000),
        "warm_wall_ms": int(wall_second * 1000),
        "cold_stages": (first["docx_ms"], first["word_ms"], first["png_ms"]),
        "warm_stages": (second["docx_ms"], second["word_ms"], second["png_ms"]),
        "pages": len(pages or []), "dpi": first.get("dpi"),
        "page_kb": [len(base64.b64decode(p["png_b64"])) // 1024
                    for p in (pages or [])],
    })

    # --- shutdown, twice ----------------------------------------------------
    live_preview.shutdown()
    after = live_preview.status()
    check(after["alive"] is False and after["pid"] is None,
          "shutdown() drops the resident instance", repr(after))
    check((not pid_running(pid)) if pid else False,
          "the instance's process is gone after shutdown()", "pid=%r" % pid)

    still = word_pids()
    lost = sorted(p for p in before if p not in still)
    check(not lost,
          "every Word process that was running before the test is STILL running",
          "these were closed: %r" % lost)
    for p in sorted(before):
        check(pid_running(p), "pre-existing Word pid %d survived" % p)

    left = [e for e in (pidfile_entries(report_dir) or [])
            if int(e.get("pid", 0)) == (pid or -1)]
    check(not left, "shutdown() removes our record from the pid file",
          repr(left))

    live_preview.shutdown()             # idempotent
    again = live_preview.status()
    check(again["alive"] is False and again["pid"] is None,
          "shutdown() is idempotent", repr(again))


# ---------------------------------------------------------------------------
def main():
    tmp = tempfile.mkdtemp(prefix="live_preview_test_")
    timings = {}
    try:
        report_dir = make_report(os.path.join(tmp, "MODULE_A"))

        print("\n-- import and capability --")
        test_import_without_capability(report_dir)

        print("\n-- section numbering (no Word needed) --")
        test_section_numbering(report_dir, tmp)

        print("\n-- instance identity (no Word needed) --")
        test_pid_read_before_close()
        test_dead_handle_yields_no_pid()
        test_creation_time_match_is_exact()

        print("\n-- resident Word --")
        if not live_preview.available():
            skip("resident Word render",
                 live_preview.status().get("reason") or "not available here")
        else:
            test_word_path(report_dir, timings)
    finally:
        try:
            live_preview.shutdown()
        except Exception:
            pass
        shutil.rmtree(tmp, ignore_errors=True)

    if timings:
        print("\n  timings: cold %(cold_ms)dms  warm %(warm_ms)dms  "
              "(%(pages)d page(s) at %(dpi)d dpi)" % timings)
        print("           cold stages docx/word/png = %r ms"
              % (timings["cold_stages"],))
        print("           warm stages docx/word/png = %r ms"
              % (timings["warm_stages"],))
        print("           page sizes = %r KB" % (timings["page_kb"],))

    print("\n=== LIVE PREVIEW SUMMARY ===")
    print("  PASS: %d   FAIL: %d   SKIP: %d" % (PASS, FAIL, len(SKIPPED)))
    print("============================")
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(main())
