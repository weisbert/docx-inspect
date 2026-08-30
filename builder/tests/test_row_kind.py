#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Regression tests for ROW-KIND driven shading of plain (free) tables.

A plain table used to shade its rows through ``row_fills`` -- a {row_index: hex6}
map -- so inserting a row shifted every colour below it. ``core/tables.py`` now
maps a per-row KIND (header / setting / result) to a fill through the template
config, and ``store/migrate_row_kind.py`` derives those kinds for existing data
by inverting the renderer's own map.

What this file locks down:

  1. A table whose rows all carry a kind renders the SAME fills as the equivalent
     index-keyed table (both with the config default map and with a template that
     names its own ``kind_fills``).
  2. A kind WINS over a ``row_fills`` entry for the same row -- including a kind
     that maps to "unshaded", which must beat a stored colour rather than fall
     through to it.
  3. ``row_fills`` on its own still renders exactly as it did before kinds
     existed (header rows take the header fill, every other row takes its stored
     entry, unknown rows stay unshaded).
  4. Inserting a row in the middle of a kind-driven table leaves every other
     row's fill unchanged. This is the bug the whole change exists to fix, so it
     is asserted directly, and the old index-keyed path is asserted to still show
     the shift -- otherwise the test would pass for the wrong reason.
  5. ``migrate()`` as a dry run touches nothing on disk; an apply changes only
     the tables it reported (and only by adding ``row_kinds``); rollback restores
     byte-identical files.
  6. ``migrate()`` is idempotent -- a second run reports nothing and writes
     nothing.
  7. The default reports root is the one the sibling layout migration resolves,
     never the repository checkout.
  8. The snapshot copies ONLY the reports about to be rewritten, and a copy it
     cannot make aborts the run instead of leaving a file with no pre-image.

Plus the migration's core promise, which is what makes 5 safe: a migrated table
rendered from its derived kinds produces the same fills as the same table
rendered from its original ``row_fills``.

The fixture data is self-contained, ASCII-only and uses neutral module names.
The migration is only ever pointed at a throwaway root under the system temp
directory, never at the repository.

Run:
    python builder/tests/test_row_kind.py
"""

import json
import os
import shutil
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))     # builder/tests
BUILDER = os.path.dirname(HERE)                       # builder/
if BUILDER not in sys.path:
    sys.path.insert(0, BUILDER)
import buildpath  # noqa: E402,F401  (registers core/docx_io/store/sync/web)

from docx import Document                # noqa: E402
from docx.oxml.ns import qn              # noqa: E402

import tables as T                       # noqa: E402
import migrate_row_kind as M             # noqa: E402


_fails = []


def check(cond, name):
    print(("  PASS  " if cond else "  FAIL  ") + name)
    if not cond:
        _fails.append(name)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
# Mirrors the shape of a real template's free_table section: only a header fill,
# so the renderer's default condition shade applies.
CFG_DEFAULT = {
    "header_fill": "D9D9D9",
    "border": {"val": "single", "sz": 4, "color": "000000"},
    "font_pt": 8.0,
}

# A template that names the whole kind -> fill map itself.
CFG_MAPPED = {
    "border": {"val": "single", "sz": 4, "color": "000000"},
    "font_pt": 8.0,
    "kind_fills": {"header": "FFFF00", "setting": "F2F2F2", "result": None},
}

# One header row, three condition rows, three result rows.
ROWS = [
    ["Item", "Condition", "Unit"],
    ["Supply", "1.80", "V"],
    ["Temperature", "25", "degC"],
    ["Corner", "TT", ""],
    ["Divider ratio", "8", ""],
    ["Current", "412", "uA"],
    ["Lock time", "3.1", "us"],
]
KINDS = ["header", "setting", "setting", "setting", "result", "result", "result"]
SETTING_ROWS = [1, 2, 3]


# ---------------------------------------------------------------------------
# docx inspection helpers
# ---------------------------------------------------------------------------
def cell_fill(cell):
    """The hex6 shading of one cell, or None when it carries no w:shd."""
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    val = shd.get(qn("w:fill"))
    return val.upper() if val else None


def cell_bold(cell):
    return any(bool(run.font.bold) for p in cell.paragraphs for run in p.runs)


def render_fills(rows, cfg, **kw):
    """Render a free table into a fresh document and return one fill per row.

    Asserts on the way out that a row is shaded uniformly -- the renderer shades
    whole rows, and a per-row answer is only meaningful if that stays true.
    """
    doc = Document()
    res = T.render_free_table(doc, rows, cfg, **kw)
    table = res["table"]
    out = []
    for r in range(len(rows)):
        seen = set(cell_fill(table.cell(r, c)) for c in range(len(table.columns)))
        assert len(seen) == 1, "row %d is not uniformly shaded: %r" % (r, seen)
        out.append(seen.pop())
    return out


def render_bolds(rows, cfg, **kw):
    doc = Document()
    table = T.render_free_table(doc, rows, cfg, **kw)["table"]
    return [cell_bold(table.cell(r, 0)) for r in range(len(rows))]


def text_to_fill(rows, fills):
    """Map a row's first-cell text -> its fill, so rows can be compared by
    identity rather than by position across an insert."""
    return dict((T._row_cells(row)[0], fill) for row, fill in zip(rows, fills))


# ---------------------------------------------------------------------------
# 1. kinds render the same fills as the equivalent index-keyed table
# ---------------------------------------------------------------------------
def test_kind_matches_row_fills():
    print("== kinds render the same fills as row_fills ==")

    by_kind = render_fills(ROWS, CFG_DEFAULT, header_rows=1, row_kinds=KINDS)
    by_index = render_fills(ROWS, CFG_DEFAULT, header_rows=1,
                            row_fills=dict((str(i), "EEECE1") for i in SETTING_ROWS))
    check(by_kind == by_index,
          "default config: kind-driven fills == index-driven fills")
    check(by_kind == ["D9D9D9", "EEECE1", "EEECE1", "EEECE1", None, None, None],
          "default config: header shaded, conditions beige, results unshaded")

    # A template that names its own map must be honoured, not the default shade.
    mapped_kind = render_fills(ROWS, CFG_MAPPED, header_rows=1, row_kinds=KINDS)
    mapped_index = render_fills(ROWS, CFG_MAPPED, header_rows=1,
                                header_fill="FFFF00",
                                row_fills=dict((str(i), "F2F2F2")
                                               for i in SETTING_ROWS))
    check(mapped_kind == mapped_index,
          "template kind_fills: kind-driven fills == index-driven fills")
    check(mapped_kind == ["FFFF00", "F2F2F2", "F2F2F2", "F2F2F2", None, None, None],
          "template kind_fills: the config map is what gets drawn")

    # The same kinds delivered as a {index: kind} map, not a parallel list.
    as_map = render_fills(ROWS, CFG_DEFAULT, header_rows=1,
                          row_kinds=dict((str(i), k) for i, k in enumerate(KINDS)))
    check(as_map == by_kind, "a {index: kind} map renders like a parallel list")

    # And carried inline on a dict row.
    dict_rows = [{"cells": cells, "kind": kind} for cells, kind in zip(ROWS, KINDS)]
    inline = render_fills(dict_rows, CFG_DEFAULT, header_rows=1)
    check(inline == by_kind, "a dict row's inline kind renders like a parallel list")

    check(T.free_kind_fills(CFG_DEFAULT) ==
          {"header": "D9D9D9", "setting": "EEECE1", "result": None},
          "free_kind_fills exposes the default map")
    check(T.free_kind_fills(CFG_MAPPED, header_fill="FFCC66")["header"] == "FFCC66",
          "a per-table header_fill still beats the config header entry")


# ---------------------------------------------------------------------------
# 2. kind wins over row_fills
# ---------------------------------------------------------------------------
def test_kind_beats_row_fills():
    print("== kind wins when both carriers are present ==")

    # Deliberately wrong index-keyed colours on every body row.
    wrong = dict((str(i), "FF0000") for i in range(1, len(ROWS)))
    both = render_fills(ROWS, CFG_DEFAULT, header_rows=1,
                        row_fills=wrong, row_kinds=KINDS)
    only_kind = render_fills(ROWS, CFG_DEFAULT, header_rows=1, row_kinds=KINDS)
    check(both == only_kind, "row_fills is ignored on every row that has a kind")
    check("FF0000" not in both, "no stored colour leaks through a kind")
    check(both[4:] == [None, None, None],
          "a kind mapped to unshaded beats a stored colour "
          "(it does not fall through to row_fills)")

    # A row with NO kind still falls back to its stored colour, which is what
    # keeps a partially migrated table rendering as before.
    partial = ["header", "setting", None, None, "result", "result", "result"]
    mixed = render_fills(ROWS, CFG_DEFAULT, header_rows=1,
                         row_fills={"2": "FFECD3", "3": "FFF6EA"},
                         row_kinds=partial)
    check(mixed == ["D9D9D9", "EEECE1", "FFECD3", "FFF6EA", None, None, None],
          "a row with no kind still falls back to its row_fills entry")


# ---------------------------------------------------------------------------
# 3. row_fills alone is unchanged (backward compatibility)
# ---------------------------------------------------------------------------
def test_row_fills_backward_compatible():
    print("== row_fills alone renders exactly as before ==")

    idx_fills = dict((str(i), "EEECE1") for i in SETTING_ROWS)
    legacy = render_fills(ROWS, CFG_DEFAULT, header_rows=1, row_fills=idx_fills)
    check(legacy == ["D9D9D9", "EEECE1", "EEECE1", "EEECE1", None, None, None],
          "header row takes the header fill, stored rows their colour, rest none")

    # Passing row_kinds=None explicitly must be identical to not passing it.
    check(render_fills(ROWS, CFG_DEFAULT, header_rows=1, row_kinds=None,
                       row_fills=idx_fills) == legacy,
          "row_kinds=None is the same call as omitting it")

    # A table-level header_fill plus alternating band colours that belong to no
    # kind at all: the exact shape the migration leaves alone.
    band = render_fills(ROWS[:5], CFG_DEFAULT, header_rows=1,
                        header_fill="FFCC66",
                        row_fills={"1": "FFECD3", "2": "FFF6EA",
                                   "3": "FFECD3", "4": "FFF6EA"})
    check(band == ["FFCC66", "FFECD3", "FFF6EA", "FFECD3", "FFF6EA"],
          "per-table header_fill and unmapped band colours still render")

    # Two header rows, integer keys, and a gap in the map.
    two_head = render_fills(ROWS, CFG_DEFAULT, header_rows=2,
                            row_fills={3: "EEECE1", 5: "EEECE1"})
    check(two_head == ["D9D9D9", "D9D9D9", None, "EEECE1", None, "EEECE1", None],
          "header_rows, integer keys and gaps behave as before")

    # Bold: the header rows are bold and nothing else is, with or without kinds.
    check(render_bolds(ROWS, CFG_DEFAULT, header_rows=1, row_fills=idx_fills) ==
          render_bolds(ROWS, CFG_DEFAULT, header_rows=1, row_kinds=KINDS),
          "kinds do not change which rows are bold")


# ---------------------------------------------------------------------------
# 4. THE BUG: inserting a row must not move anyone else's colour
# ---------------------------------------------------------------------------
def test_insert_row_keeps_fills():
    print("== inserting a row in the middle ==")

    new_row = ["Load capacitance", "0.6", "pF"]

    # Insert a further CONDITION row at index 3, i.e. in the middle of the
    # condition band and above every result row.
    at = 3
    grown = ROWS[:at] + [new_row] + ROWS[at:]
    grown_kinds = KINDS[:at] + ["setting"] + KINDS[at:]

    before = text_to_fill(ROWS, render_fills(ROWS, CFG_DEFAULT, header_rows=1,
                                             row_kinds=KINDS))
    after = text_to_fill(grown, render_fills(grown, CFG_DEFAULT, header_rows=1,
                                             row_kinds=grown_kinds))

    moved = dict((name, (before[name], after[name]))
                 for name in before if before[name] != after[name])
    check(not moved,
          "kind-driven: every pre-existing row keeps its fill after an insert "
          "(moved: %r)" % (moved,))
    check(after[new_row[0]] == "EEECE1",
          "kind-driven: the inserted condition row is shaded as a condition row")

    # Inserting a RESULT row directly above the result band is the harsher case:
    # under index addressing it pushes a beige row down into the results.
    at2 = 4
    grown2 = ROWS[:at2] + [new_row] + ROWS[at2:]
    grown2_kinds = KINDS[:at2] + ["result"] + KINDS[at2:]
    after2 = text_to_fill(grown2, render_fills(grown2, CFG_DEFAULT, header_rows=1,
                                               row_kinds=grown2_kinds))
    moved2 = dict((name, (before[name], after2[name]))
                  for name in before if before[name] != after2[name])
    check(not moved2,
          "kind-driven: a result row inserted at the band boundary moves nothing "
          "(moved: %r)" % (moved2,))
    check(after2[new_row[0]] is None,
          "kind-driven: the inserted result row is unshaded")

    # Control: the index-keyed path MUST still show the shift, otherwise the two
    # assertions above would be passing for the wrong reason.
    idx_fills = dict((str(i), "EEECE1") for i in SETTING_ROWS)
    legacy_before = text_to_fill(ROWS, render_fills(ROWS, CFG_DEFAULT,
                                                    header_rows=1,
                                                    row_fills=idx_fills))
    legacy_after = text_to_fill(grown, render_fills(grown, CFG_DEFAULT,
                                                    header_rows=1,
                                                    row_fills=idx_fills))
    check(any(legacy_before[n] != legacy_after[n] for n in legacy_before),
          "control: the index-keyed path DOES shift colours on the same insert")


# ---------------------------------------------------------------------------
# Migration fixture: a throwaway reports root
# ---------------------------------------------------------------------------
TEMPLATE_CONFIG = {"styles": {}, "free_table": dict(CFG_DEFAULT)}


def _table_block(bid, rows, row_fills, header_fill=None):
    block = {"id": bid, "type": "table", "rows": [list(r) for r in rows],
             "header_rows": 1, "col_w": [4.0, 3.0, 2.0]}
    if row_fills:
        block["row_fills"] = dict(row_fills)
    if header_fill:
        block["header_fill"] = header_fill
    return block


def _project(title, blocks):
    return {
        "meta": {"title": title, "stage": "CDR"},
        "template": "tpl_neutral",
        "outline": [
            {"id": "n1", "title": "Simulation setup", "blocks": blocks[:1],
             "children": [
                 {"id": "n2", "title": "Conditions", "blocks": blocks[1:],
                  "children": []},
             ]},
        ],
    }


def build_root(base):
    """Write a small reports root: two reports with index-keyed tables, one
    without, a nested three-level path and a flat one."""
    root = os.path.join(base, "reports")
    tpl = os.path.join(root, "templates", "tpl_neutral")
    os.makedirs(tpl)
    with open(os.path.join(tpl, "config.json"), "w", encoding="utf-8") as fh:
        json.dump(TEMPLATE_CONFIG, fh, indent=2)

    # Nested project/module/stage layout, two tables that shade by index.
    nested = os.path.join(root, "1108", "CLKDIV_5G", "CDR")
    os.makedirs(os.path.join(nested, "images"))
    # Bulk that lives beside a report and that the migration never touches: the
    # snapshot must not duplicate it (see test_snapshot_scope_and_failure).
    with open(os.path.join(nested, "images", "figure.png"), "wb") as fh:
        fh.write(b"PNGDATA" + b"x" * 4096)
    p1 = _project("CLKDIV_5G CDR", [
        _table_block("t-plain", ROWS,
                     dict((str(i), "EEECE1") for i in SETTING_ROWS)),
        _table_block("t-band", ROWS[:5],
                     {"1": "FFECD3", "2": "FFF6EA", "3": "FFECD3", "4": "FFF6EA"},
                     header_fill="FFCC66"),
    ])
    with open(os.path.join(nested, "project.json"), "w", encoding="utf-8") as fh:
        json.dump(p1, fh, indent=2)

    # Flat legacy layout, one table that shades by index.
    flat = os.path.join(root, "MODULE_A")
    os.makedirs(os.path.join(flat, "out"))
    with open(os.path.join(flat, "out", "notes.txt"), "wb") as fh:
        fh.write(b"unrelated working file")
    p2 = _project("MODULE_A CDR", [
        _table_block("t-only", ROWS, {"1": "EEECE1", "2": "EEECE1"}),
        {"id": "p1", "type": "para", "runs": [{"t": "No table here."}]},
    ])
    with open(os.path.join(flat, "project.json"), "w", encoding="utf-8") as fh:
        json.dump(p2, fh, indent=2)

    # A report whose tables carry no row_fills at all: must not be touched.
    untouched = os.path.join(root, "TXBUF_2G")
    os.makedirs(untouched)
    p3 = _project("TXBUF_2G CDR", [
        _table_block("t-nofills", ROWS, None),
        {"id": "p2", "type": "para", "runs": [{"t": "Nothing to migrate."}]},
    ])
    with open(os.path.join(untouched, "project.json"), "w", encoding="utf-8") as fh:
        json.dump(p3, fh, indent=2)
    return root


def snapshot_bytes(root):
    """{relative path: bytes} for every file under root, so a change anywhere
    on disk is visible -- not just in the files the migration means to touch."""
    out = {}
    for dirpath, dirnames, filenames in os.walk(root):
        for name in filenames:
            full = os.path.join(dirpath, name)
            rel = os.path.relpath(full, root).replace("\\", "/")
            with open(full, "rb") as fh:
                out[rel] = fh.read()
    return out


def load(root, rel):
    with open(os.path.join(root, rel), "r", encoding="utf-8") as fh:
        return json.load(fh)


def strip_kinds(node):
    """The project with every ``row_kinds`` removed, for an "only that key was
    added" comparison."""
    if isinstance(node, dict):
        return dict((k, strip_kinds(v)) for k, v in node.items()
                    if k != "row_kinds")
    if isinstance(node, list):
        return [strip_kinds(v) for v in node]
    return node


def blocks_of(project):
    out = {}

    def walk(node):
        for block in node.get("blocks") or []:
            if isinstance(block, dict) and block.get("id"):
                out[block["id"]] = block
        for child in node.get("children") or []:
            walk(child)
    for node in project.get("outline") or []:
        walk(node)
    return out


# ---------------------------------------------------------------------------
# 5 + 6. dry run / apply / rollback / idempotence
# ---------------------------------------------------------------------------
def test_migration(base):
    print("== migration: dry run, apply, rollback, idempotence ==")
    root = build_root(base)
    original = snapshot_bytes(root)

    # --- dry run ---------------------------------------------------------
    dry = M.migrate(root, dry_run=True)
    check(dry["scanned"] == 3,
          "dry run scans the 3 tables that shade by index (got %d)" % dry["scanned"])
    check(dry["changed"] == 3,
          "dry run plans 3 changes (got %d)" % dry["changed"])
    check(dry["snapshot"] is None, "dry run takes no snapshot")
    check(snapshot_bytes(root) == original, "dry run changes nothing on disk")
    check(not os.path.exists(os.path.join(root, "_migrations")),
          "dry run does not even create the snapshot directory")

    reported = set((e["report"], e["block"]) for e in dry["tables"])
    check(("1108/CLKDIV_5G/CDR", "t-plain") in reported and
          ("MODULE_A", "t-only") in reported,
          "dry run reaches both a nested and a flat report")
    check(not any(e["block"] == "t-nofills" for e in dry["tables"]),
          "a table without row_fills is not reported")
    band = [e for e in dry["tables"] if e["block"] == "t-band"]
    check(len(band) == 1 and band[0]["unmapped_rows"] == [1, 2, 3, 4],
          "band colours that match no kind are reported as unmapped")

    # --- apply -----------------------------------------------------------
    applied = M.migrate(root, dry_run=False)
    check(applied["changed"] == 3, "apply changes the 3 planned tables")
    snap = applied["snapshot"]
    check(bool(snap) and os.path.isdir(snap),
          "apply took a snapshot before writing")

    after = snapshot_bytes(root)
    touched = set(rel for rel in original if original[rel] != after.get(rel))
    check(touched == set(["1108/CLKDIV_5G/CDR/project.json",
                          "MODULE_A/project.json"]),
          "apply rewrote only the two reports it reported (touched: %r)"
          % sorted(touched))
    check(all(rel in after for rel in original), "apply deleted nothing")

    for rel in sorted(touched):
        check(strip_kinds(load(root, rel)) ==
              strip_kinds(json.loads(original[rel].decode("utf-8"))),
              "%s changed by nothing except added row_kinds" % rel)

    nested_blocks = blocks_of(load(root, "1108/CLKDIV_5G/CDR/project.json"))
    check(nested_blocks["t-plain"].get("row_kinds") == KINDS,
          "derived kinds match the stored colours (got %r)"
          % (nested_blocks["t-plain"].get("row_kinds"),))
    check(bool(nested_blocks["t-plain"].get("row_fills")),
          "row_fills is left in place so the older editor still draws the bands")
    check(nested_blocks["t-band"].get("row_kinds") ==
          ["header", None, None, None, None],
          "an unmapped band row keeps a null kind rather than being guessed at")
    flat_blocks = blocks_of(load(root, "MODULE_A/project.json"))
    check(flat_blocks["t-only"].get("row_kinds") ==
          ["header", "setting", "setting", "result", "result", "result", "result"],
          "an unshaded body row is derived as a result row")
    check("row_kinds" not in
          blocks_of(load(root, "TXBUF_2G/project.json"))["t-nofills"],
          "a table that never shaded by index gains no kinds")

    # The whole point of deriving kinds by inverting the render map: the
    # document must come out the same. Render each migrated table both ways.
    same = True
    for rel in ("1108/CLKDIV_5G/CDR/project.json", "MODULE_A/project.json"):
        for bid, block in sorted(blocks_of(load(root, rel)).items()):
            if block.get("type") != "table" or not block.get("row_fills"):
                continue
            kw = {"header_rows": block.get("header_rows", 1)}
            if block.get("header_fill"):
                kw["header_fill"] = block["header_fill"]
            old = render_fills(block["rows"], CFG_DEFAULT,
                               row_fills=block["row_fills"], **kw)
            new = render_fills(block["rows"], CFG_DEFAULT,
                               row_fills=block["row_fills"],
                               row_kinds=block["row_kinds"], **kw)
            if old != new:
                same = False
                print("      %s / %s: %r != %r" % (rel, bid, old, new))
    check(same,
          "every migrated table renders the same fills as before the migration")

    # --- idempotence -----------------------------------------------------
    again = M.migrate(root, dry_run=True)
    check(again["changed"] == 0,
          "a second dry run plans no change (got %d)" % again["changed"])
    check(again["scanned"] == 3, "the tables are still found, just not changed")
    check(all(e["already"] for e in again["tables"]),
          "every table is reported as already migrated")

    second = M.migrate(root, dry_run=False)
    check(second["changed"] == 0, "a second apply changes nothing")
    check(second["snapshot"] is None,
          "a second apply takes no snapshot, having nothing to write")
    check(snapshot_bytes(root) == after, "a second apply writes no bytes")

    # --- rollback --------------------------------------------------------
    res = M.rollback(snap)
    check(res["restored"] == 2,
          "rollback restores the 2 rewritten reports (got %d)" % res["restored"])
    restored = snapshot_bytes(root)
    diff = [rel for rel in original if original[rel] != restored.get(rel)]
    check(not diff, "rollback restores byte-identical files (differ: %r)" % diff)

    check(M.rollback(snap)["restored"] == 0, "rolling back twice is a no-op")

    # A mid-table divider row that happens to carry the HEADER colour must not
    # be derived as kind "header": the renderer bolds every "header" row, so the
    # migration would keep the fills but silently bold a body row, breaking the
    # promise that the rendered document is unchanged. It is left unmapped.
    divider_rows = [["Item", "Value"], ["a", "1"],
                    ["Sub-total", ""], ["b", "2"]]
    divider = {"type": "table", "rows": divider_rows, "header_rows": 1,
               "row_fills": {"2": "D9D9D9"}}
    kinds, unmapped = M.derive_kinds(divider, CFG_DEFAULT)
    check(kinds == ["header", "result", None, "result"] and unmapped == [2],
          "a body row carrying the header colour is left unmapped (got %r / %r)"
          % (kinds, unmapped))
    check(render_bolds(divider_rows, CFG_DEFAULT, header_rows=1,
                       row_fills=divider["row_fills"]) ==
          render_bolds(divider_rows, CFG_DEFAULT, header_rows=1,
                       row_fills=divider["row_fills"], row_kinds=kinds),
          "migrating a table with a divider row does not bold it")
    check(render_fills(divider_rows, CFG_DEFAULT, header_rows=1,
                       row_fills=divider["row_fills"]) ==
          render_fills(divider_rows, CFG_DEFAULT, header_rows=1,
                       row_fills=divider["row_fills"], row_kinds=kinds),
          "migrating a table with a divider row does not move its fills")

    # And the migration can be run again from the rolled-back state.
    redo = M.migrate(root, dry_run=True)
    check(redo["changed"] == 3, "the rolled-back root migrates again cleanly")


# ---------------------------------------------------------------------------
# 7. the default reports root
# ---------------------------------------------------------------------------
def test_default_root(base):
    print("== the default reports root ==")
    import migrate_layout as L                      # the sibling migration

    saved = os.environ.pop("BUILDER_REPORTS_ROOT", None)
    try:
        default = M.resolve_root(None)
        sibling = L.resolve_root(None)
        check(os.path.normcase(default) == os.path.normcase(sibling),
              "the default root is the one the sibling migration resolves "
              "(%r vs %r)" % (default, sibling))
        # The defect this pins down: the default used to be the repository
        # checkout, so an --apply with no --root walked -- and copied -- the
        # source tree and every report under it. Asserted wherever the two can
        # differ at all; on a bare checkout with no reports folder beside it,
        # both correctly fall back to the repo and there is nothing to compare.
        if os.path.isdir(os.path.join(M._REPO, M._REPORTS_DIRNAME)):
            check(os.path.normcase(default) != os.path.normcase(M._REPO),
                  "the default root is the reports root, not the repository")

        env_root = os.path.join(base, "env_root")
        os.makedirs(env_root)
        os.environ["BUILDER_REPORTS_ROOT"] = env_root
        check(M.resolve_root(None) == os.path.abspath(env_root),
              "BUILDER_REPORTS_ROOT still wins over the default")
        check(M.resolve_root(base) == os.path.abspath(base),
              "an explicit root still wins over the env var")
    finally:
        os.environ.pop("BUILDER_REPORTS_ROOT", None)
        if saved is not None:
            os.environ["BUILDER_REPORTS_ROOT"] = saved


# ---------------------------------------------------------------------------
# 8. the snapshot: only what is being rewritten, and never a partial one
# ---------------------------------------------------------------------------
def test_snapshot_scope_and_failure(base):
    print("== snapshot scope and an incomplete snapshot ==")
    root = build_root(base)
    before = snapshot_bytes(root)

    # A copy that fails must abort the whole run. Skipping it would leave a
    # rewritten report with no pre-image, i.e. nothing for rollback to restore.
    real_copy2 = shutil.copy2

    def failing_copy2(src, dst, *a, **kw):
        if os.path.basename(src) == "project.json":
            raise OSError(13, "simulated copy failure")
        return real_copy2(src, dst, *a, **kw)

    shutil.copy2 = failing_copy2
    raised = None
    try:
        M.migrate(root, dry_run=False)
    except M.MigrationError as ex:
        raised = str(ex)
    except Exception as ex:                          # any other type is a fail
        print("      unexpected exception: %r" % (ex,))
    finally:
        shutil.copy2 = real_copy2
    check(bool(raised),
          "a snapshot that cannot copy a needed file raises MigrationError")
    check(snapshot_bytes(root) == before,
          "the aborted run wrote nothing: no report is rewritten without a "
          "pre-image to roll back to")

    # The successful run copies the reports it rewrites -- and nothing else.
    applied = M.migrate(root, dry_run=False)
    snap = applied["snapshot"]
    copied = sorted(snapshot_bytes(snap))
    check(copied == ["1108/CLKDIV_5G/CDR/project.json",
                     "MODULE_A/project.json",
                     M.SNAPSHOT_MANIFEST],
          "the snapshot holds ONLY the reports being rewritten (got %r)"
          % (copied,))
    check(not any(rel.endswith((".png", ".txt")) for rel in copied),
          "figures and unrelated working files are not duplicated into it")

    res = M.rollback(snap)
    check(res["restored"] == 2,
          "rollback restores both rewritten reports (got %d)" % res["restored"])
    back = snapshot_bytes(root)
    diff = [rel for rel in before if before[rel] != back.get(rel)]
    check(not diff, "rollback restores byte-identical files (differ: %r)" % diff)


# ---------------------------------------------------------------------------
def main():
    test_kind_matches_row_fills()
    test_kind_beats_row_fills()
    test_row_fills_backward_compatible()
    test_insert_row_keeps_fills()
    base = tempfile.mkdtemp(prefix="row_kind_test_")
    try:
        test_migration(base)
    finally:
        shutil.rmtree(base, ignore_errors=True)

    for fn in (test_default_root, test_snapshot_scope_and_failure):
        base = tempfile.mkdtemp(prefix="row_kind_test_")
        try:
            fn(base)
        finally:
            shutil.rmtree(base, ignore_errors=True)

    print("")
    if _fails:
        print("FAILED (%d):" % len(_fails))
        for f in _fails:
            print("  - " + f)
        return 1
    print("all row-kind assertions passed")
    return 0


if __name__ == "__main__":
    sys.exit(main())
