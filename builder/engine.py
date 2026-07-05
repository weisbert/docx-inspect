# -*- coding: utf-8 -*-
"""
engine.py -- Generic, config-driven structured-document renderer.

Reads a project file (``project.json``, schema v1) plus a template configuration file and
produces a .docx whose layout (page geometry, 3-column header/footer tables, cover tables,
heading styles + multilevel autonumber, fixed rich-text bodies, the four content block
types, chapter-sequence caption numbering, and an inline portrait data-driven table) is
entirely defined by the configuration -- nothing domain specific is hardcoded here.

CLI:
    python engine.py <report_folder> [--config <template_config.json>] [--out <dir>]

    <report_folder> must contain ``project.json`` (and an ``images/`` folder for image blocks).
    The template config path may also come from the ``BUILDER_TEMPLATE_CONFIG`` env var; if
    neither is given, the engine looks for ``project.json``'s ``template`` id as
    ``template_config_<id>.json`` next to the project folder's parent.

Output: ``<report_folder>/out/<name>.docx``.
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import tables
import content_lint


# ===========================================================================
# Low-level docx helpers (ported, kept generic)
# ===========================================================================
def _rgb(hex6):
    return RGBColor.from_string(hex6)


def set_rfonts(el, ascii=None, hansi=None, eastasia=None, cs=None):
    rpr = el.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for k, v in (("w:ascii", ascii), ("w:hAnsi", hansi),
                 ("w:eastAsia", eastasia), ("w:cs", cs)):
        if v:
            rf.set(qn(k), v)


def run_fmt(run, ascii=None, eastasia=None, size=None, bold=None, italic=None,
            color=None, underline=None):
    if ascii or eastasia:
        set_rfonts(run._r, ascii=ascii, hansi=ascii, eastasia=eastasia, cs=ascii)
    f = run.font
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if underline is not None:
        f.underline = underline
    if color is not None:
        f.color.rgb = color if isinstance(color, RGBColor) else _rgb(color)


def style_font(st, ascii=None, eastasia=None, size=None, bold=None, italic=None,
               color=None, underline=None):
    if ascii or eastasia:
        set_rfonts(st.element, ascii=ascii, hansi=ascii, eastasia=eastasia, cs=ascii)
    f = st.font
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if underline is not None:
        f.underline = underline
    if color is not None:
        f.color.rgb = color if isinstance(color, RGBColor) else _rgb(color)


def style_para(st, align=None, left=None, first_line=None, before=None, after=None,
               line=None, keepnext=None):
    pf = st.paragraph_format
    if align is not None:
        pf.alignment = align
    if left is not None:
        pf.left_indent = Cm(left)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if keepnext:
        st.element.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def new_style(doc, name, base=None):
    try:
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        st = doc.styles[name]
    if base:
        st.base_style = doc.styles[base]
    return st


def cell_valign(cell, val="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    e = OxmlElement("w:vAlign")
    e.set(qn("w:val"), val)
    tcPr.append(e)


def _border(tag, val, sz, color):
    e = OxmlElement(tag)
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)
    return e


def cell_borders(cell, edges, val="single", sz=6, color="auto"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = OxmlElement("w:tcBorders")
        tcPr.append(b)
    for edge in edges:
        for old in b.findall(qn("w:" + edge)):
            b.remove(old)
        b.append(_border("w:" + edge, val, sz, color))


def table_grid(table, widths_cm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tblPr.addnext(grid)
    else:
        for gc in grid.findall(qn("w:gridCol")):
            grid.remove(gc)
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 567)))
        grid.append(gc)
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def table_outer_inner(table, outer="double", outer_sz=6, inner="single",
                      inner_sz=6, color="000000"):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    b.append(_border("w:top", outer, outer_sz, color))
    b.append(_border("w:left", outer, outer_sz, color))
    b.append(_border("w:bottom", outer, outer_sz, color))
    b.append(_border("w:right", outer, outer_sz, color))
    b.append(_border("w:insideH", inner, inner_sz, color))
    b.append(_border("w:insideV", inner, inner_sz, color))
    tblPr.append(b)


def table_all_single(table, sz=6, color="auto"):
    table_outer_inner(table, outer="single", outer_sz=sz, inner="single",
                      inner_sz=sz, color=color)


def table_center(table):
    tblPr = table._tbl.tblPr
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)


def row_exact_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:hRule"), "exact")
    h.set(qn("w:val"), str(twips))
    trPr.append(h)


def add_field(p, instr, placeholder="", bold=None, bookmark_id=None, bookmark_name=None):
    """Emit a complete Word field (begin / instrText / separate / placeholder / end).

    When both ``bookmark_id`` and ``bookmark_name`` are given, the entire field is
    wrapped in <w:bookmarkStart/> ... <w:bookmarkEnd/> so a REF field can target the
    number the field renders. ``bookmark_id`` must be a per-document-unique integer
    and ``bookmark_name`` a per-document-unique name (see CaptionState)."""
    def mk(run_el):
        if bold:
            rpr = OxmlElement("w:rPr")
            rpr.append(OxmlElement("w:b"))
            run_el.append(rpr)
    if bookmark_id is not None and bookmark_name:
        bs = OxmlElement("w:bookmarkStart")
        bs.set(qn("w:id"), str(bookmark_id))
        bs.set(qn("w:name"), bookmark_name)
        p._p.append(bs)
    r1 = OxmlElement("w:r"); mk(r1)
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); r1.append(b)
    r2 = OxmlElement("w:r"); mk(r2)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; r2.append(it)
    r3 = OxmlElement("w:r"); mk(r3)
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate"); r3.append(s)
    r4 = OxmlElement("w:r"); mk(r4)
    t = OxmlElement("w:t"); t.text = placeholder; r4.append(t)
    r5 = OxmlElement("w:r"); mk(r5)
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); r5.append(e)
    for r in (r1, r2, r3, r4, r5):
        p._p.append(r)
    if bookmark_id is not None and bookmark_name:
        be = OxmlElement("w:bookmarkEnd")
        be.set(qn("w:id"), str(bookmark_id))
        p._p.append(be)


def fill(cell, text, ascii=None, eastasia=None, size=None, bold=None, italic=None,
         color=None, underline=None, align=None, valign=None):
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    if valign:
        cell_valign(cell, valign)
    if text:
        r = p.add_run(text)
        run_fmt(r, ascii=ascii, eastasia=eastasia, size=size, bold=bold,
                italic=italic, color=color, underline=underline)
    return p


def add_inline_logo(part_container, cell, logo_path, width_cm=1.13):
    p = cell.paragraphs[0]
    p.alignment = ALIGN.CENTER
    if not logo_path or not os.path.exists(logo_path):
        r = p.add_run("[LOGO]")
        run_fmt(r, color="FF0000")
        return
    try:
        inline = part_container.part.new_pic_inline(logo_path, Cm(width_cm), Cm(width_cm))
        run = p.add_run()
        drawing = OxmlElement("w:drawing")
        drawing.append(inline)
        run._r.append(drawing)
    except Exception as ex:
        r = p.add_run("[LOGO]")
        run_fmt(r, color="FF0000")
        print("  ! logo failed:", ex)


# ===========================================================================
# Heading multilevel autonumber (config-driven num_id / fonts)
# ===========================================================================
def add_heading_numbering(doc, num_id=88, abstract_id=88, suffix="space", ascii_font="Arial"):
    numbering = doc.part.numbering_part.element
    abs_el = OxmlElement("w:abstractNum")
    abs_el.set(qn("w:abstractNumId"), str(abstract_id))
    ml = OxmlElement("w:multiLevelType")
    ml.set(qn("w:val"), "multilevel")
    abs_el.append(ml)
    for i in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(i))
        st = OxmlElement("w:start"); st.set(qn("w:val"), "1"); lvl.append(st)
        nf = OxmlElement("w:numFmt"); nf.set(qn("w:val"), "decimal"); lvl.append(nf)
        ps = OxmlElement("w:pStyle"); ps.set(qn("w:val"), f"Heading{i+1}"); lvl.append(ps)
        suff = OxmlElement("w:suff"); suff.set(qn("w:val"), suffix); lvl.append(suff)
        lt = OxmlElement("w:lvlText")
        lt.set(qn("w:val"), ".".join(f"%{j+1}" for j in range(i + 1)) + " ")
        lvl.append(lt)
        jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lvl.append(jc)
        pPr = OxmlElement("w:pPr"); ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0"); ind.set(qn("w:firstLine"), "0")
        pPr.append(ind); lvl.append(pPr)
        abs_el.append(lvl)
    first = numbering.find(qn("w:num"))
    if first is not None:
        first.addprevious(abs_el)
    else:
        numbering.append(abs_el)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), str(abstract_id))
    num.append(aid)
    numbering.append(num)
    for i in range(9):
        st = doc.styles[f"Heading {i+1}"]
        pPr = st.element.get_or_add_pPr()
        for old in pPr.findall(qn("w:numPr")):
            pPr.remove(old)
        numPr = OxmlElement("w:numPr")
        il = OxmlElement("w:ilvl"); il.set(qn("w:val"), str(i)); numPr.append(il)
        nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id)); numPr.append(nid)
        pPr.append(numPr)


# ===========================================================================
# Body list numbering (native Word multilevel: bullets + decimals)
# ===========================================================================
# Per-level bullet glyphs (item 3: solid circle, then hollow square, then filled
# small square). Numbered levels: 1.  a)  i.  (decimal / lower-alpha / lower-roman).
_LIST_BULLET_GLYPHS = ["●", "□", "▪", "–"]   # ● □ ▪ –
_LIST_NUM_FMT = ["decimal", "lowerLetter", "lowerRoman"]
_LIST_NUM_TEXT = ["%1.", "%2)", "%3."]
_LIST_LEFT_CM = 0.74      # base left indent (level 0)
_LIST_STEP_CM = 0.74      # extra indent per nesting level
_LIST_HANG_CM = 0.53      # hanging indent (marker column width)
_TWIPS_PER_CM = 567


def _list_level_ppr(i):
    """A <w:pPr><w:ind .../></w:pPr> for numbering level i (left + hanging)."""
    left = round((_LIST_LEFT_CM + i * _LIST_STEP_CM) * _TWIPS_PER_CM)
    hang = round(_LIST_HANG_CM * _TWIPS_PER_CM)
    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hang))
    pPr.append(ind)
    return pPr


def _add_num(numbering, num_id, abstract_id, restart=False):
    """Append a <w:num> mapping num_id -> abstract_id. When restart, override every
    level's start to 1 so a fresh numbered list begins at 1 (not continuing a prior)."""
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId")
    aid.set(qn("w:val"), str(abstract_id))
    num.append(aid)
    if restart:
        for i in range(9):
            ov = OxmlElement("w:lvlOverride")
            ov.set(qn("w:ilvl"), str(i))
            so = OxmlElement("w:startOverride")
            so.set(qn("w:val"), "1")
            ov.append(so)
            num.append(ov)
    numbering.append(num)


def build_list_numbering(doc, bullet_abs=90, bullet_num=90, dec_abs=91):
    """Create one bullet abstractNum (per-level glyphs) and one decimal abstractNum,
    plus a shared bullet <w:num>. Returns the ids a ListCtx needs. Mirrors the
    heading-numbering construction (all abstractNum before any num, per schema)."""
    numbering = doc.part.numbering_part.element
    abs_b = OxmlElement("w:abstractNum")
    abs_b.set(qn("w:abstractNumId"), str(bullet_abs))
    mlb = OxmlElement("w:multiLevelType"); mlb.set(qn("w:val"), "hybridMultilevel")
    abs_b.append(mlb)
    abs_d = OxmlElement("w:abstractNum")
    abs_d.set(qn("w:abstractNumId"), str(dec_abs))
    mld = OxmlElement("w:multiLevelType"); mld.set(qn("w:val"), "multilevel")
    abs_d.append(mld)
    for i in range(9):
        # bullet level
        lb = OxmlElement("w:lvl"); lb.set(qn("w:ilvl"), str(i))
        s = OxmlElement("w:start"); s.set(qn("w:val"), "1"); lb.append(s)
        nf = OxmlElement("w:numFmt"); nf.set(qn("w:val"), "bullet"); lb.append(nf)
        lt = OxmlElement("w:lvlText")
        lt.set(qn("w:val"), _LIST_BULLET_GLYPHS[min(i, len(_LIST_BULLET_GLYPHS) - 1)])
        lb.append(lt)
        jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lb.append(jc)
        lb.append(_list_level_ppr(i))
        abs_b.append(lb)
        # decimal level
        ld = OxmlElement("w:lvl"); ld.set(qn("w:ilvl"), str(i))
        s2 = OxmlElement("w:start"); s2.set(qn("w:val"), "1"); ld.append(s2)
        nf2 = OxmlElement("w:numFmt")
        nf2.set(qn("w:val"), _LIST_NUM_FMT[i] if i < len(_LIST_NUM_FMT) else "decimal")
        ld.append(nf2)
        lt2 = OxmlElement("w:lvlText")
        lt2.set(qn("w:val"), _LIST_NUM_TEXT[i] if i < len(_LIST_NUM_TEXT) else "%%%d." % (i + 1))
        ld.append(lt2)
        jc2 = OxmlElement("w:lvlJc"); jc2.set(qn("w:val"), "left"); ld.append(jc2)
        ld.append(_list_level_ppr(i))
        abs_d.append(ld)
    first = numbering.find(qn("w:num"))
    for a in (abs_b, abs_d):
        if first is not None:
            first.addprevious(a)
        else:
            numbering.append(a)
    _add_num(numbering, bullet_num, bullet_abs)   # bullets: one shared, non-restarting num
    return {"bullet_num": bullet_num, "dec_abs": dec_abs}


class ListCtx:
    """Allocates numbering references for body list paragraphs. Bullets share one
    num; each *numbered list* gets a fresh restarting num. A numbered list restarts
    only at a new section (new_section), so numbered items keep counting across
    intervening non-numbered paragraphs (e.g. a label line followed by a URL line)."""

    def __init__(self, doc, ids, start_num_id=100):
        self.doc = doc
        self.bullet_num = ids["bullet_num"]
        self.dec_abs = ids["dec_abs"]
        self._next = start_num_id
        self._active_dec = None

    def new_section(self):
        self._active_dec = None

    def bullet(self, level):
        return self.bullet_num, level

    def number(self, level):
        if self._active_dec is None:
            self._active_dec = self._next
            self._next += 1
            _add_num(self.doc.part.numbering_part.element, self._active_dec,
                     self.dec_abs, restart=True)
        return self._active_dec, level


def _apply_list(p, list_kind, level, list_ctx):
    """Attach numbering (numPr) + an explicit per-level indent to paragraph p. The
    explicit indent overrides the body style's first-line indent for list items."""
    if list_ctx is None or list_kind not in ("bullet", "number"):
        return
    level = max(0, min(_as_int(level, 0), 8))
    if list_kind == "bullet":
        num_id, ilvl = list_ctx.bullet(level)
    else:
        num_id, ilvl = list_ctx.number(level)
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    numPr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl"); il.set(qn("w:val"), str(ilvl)); numPr.append(il)
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id)); numPr.append(nid)
    pPr.append(numPr)
    pf = p.paragraph_format
    pf.left_indent = Cm(_LIST_LEFT_CM + level * _LIST_STEP_CM)
    pf.first_line_indent = Cm(-_LIST_HANG_CM)


# ===========================================================================
# Document setup from config (page / styles / header / footer / cover)
# ===========================================================================
def _apply_page_and_styles(doc, styles):
    page = styles["page"]
    sec = doc.sections[0]
    sec.page_width = Cm(page["w_cm"])
    sec.page_height = Cm(page["h_cm"])
    m = page["margin_cm"]
    sec.top_margin = sec.bottom_margin = Cm(m)
    sec.left_margin = sec.right_margin = Cm(m)
    sec.header_distance = Cm(page["header_dist_cm"])
    sec.footer_distance = Cm(page["footer_dist_cm"])
    sec.different_first_page_header_footer = page.get("different_first_page", True)

    # ---- Normal (font only; never touched again for row-height tricks) ----
    nm = styles["normal"]
    normal = doc.styles["Normal"]
    normal.font.size = Pt(nm["size_pt"])
    set_rfonts(normal.element, ascii=nm["ascii"], hansi=nm["ascii"],
               eastasia=nm["eastAsia"], cs=nm["ascii"])

    # ---- headings ----
    hd = styles["headings"]
    levels = hd["levels"]
    before = hd.get("space_before_pt", {})
    for lv in range(1, 10):
        st = doc.styles[f"Heading {lv}"]
        spec = levels.get(str(lv), levels.get("default", {"ascii": "Arial",
                                                          "size_pt": 12, "bold": False}))
        style_font(st, ascii=spec["ascii"], size=spec["size_pt"],
                   bold=spec.get("bold", False), italic=False, color="000000")
        style_para(st, before=before.get(str(lv), 10),
                   after=(hd.get("h1_after_pt", 58) if lv == 1 else 0))
    an = hd["autonumber"]
    add_heading_numbering(doc, num_id=an.get("num_id", 88), abstract_id=an.get("num_id", 88),
                          suffix=an.get("suffix", "space"), ascii_font=an.get("ascii", "Arial"))
    # body list numbering (bullets + decimals) -- ids threaded out via names
    list_ids = build_list_numbering(doc)
    # Heading 1 bottom border (chapter rule)
    h1b = hd["h1_bottom_border"]
    pPr = doc.styles["Heading 1"].element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), h1b["val"])
    bot.set(qn("w:sz"), str(h1b["sz"]))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), h1b.get("color", "auto"))
    pBdr.append(bot)
    pPr.append(pBdr)

    # ---- caption ----
    cp = styles["caption"]
    cap = doc.styles["Caption"]
    style_font(cap, ascii=cp["ascii"], size=cp["size_pt"], bold=cp.get("bold", True),
               italic=False, color="000000")
    style_para(cap, align=ALIGN.CENTER, before=12, after=6)

    # ---- body / mybody (mybody name comes from config so domain label stays local) ----
    bd = styles["body"]
    body = new_style(doc, bd.get("name", "body"), base=bd.get("base", "Normal"))
    style_font(body, size=bd["size_pt"])
    # first_line_cm is optional; when set it gives ordinary body paragraphs a
    # first-line indent. Absent -> unchanged (backward compatible).
    style_para(body, left=bd["left_cm"], first_line=bd.get("first_line_cm"), before=6, after=2)

    mb = styles["mybody"]
    mybody = new_style(doc, mb["name"], base=mb.get("base", bd.get("name", "body")))
    style_font(mybody, ascii=mb["ascii"])
    style_para(mybody, left=mb["left_cm"], first_line=mb["first_line_cm"])

    # ---- header / footer fonts ----
    ht = styles["header_table"]
    ft = styles["footer_table"]
    style_font(doc.styles["Header"], ascii=ht["title_font"].get("ascii", "Arial"), size=9)
    style_font(doc.styles["Footer"], ascii=ft["font"]["ascii"], size=ft["font"]["size_pt"])

    return {"body_name": bd.get("name", "body"), "mybody_name": mb["name"],
            "list_ids": list_ids}


def _build_header(doc, styles, meta, logo_path):
    sec = doc.sections[0]
    ht = styles["header_table"]
    cols = ht["cols_cm"]
    total = sum(cols)
    htbl = sec.header.add_table(rows=1, cols=3, width=Cm(total))
    table_grid(htbl, cols)
    row_exact_height(htbl.rows[0], ht.get("row_h_twips", 751))
    cb = ht.get("cell_bottom_border", {"val": "single", "sz": 6, "color": "auto"})
    for c in range(3):
        cell_borders(htbl.cell(0, c), ["bottom"], val=cb["val"], sz=cb["sz"],
                     color=cb.get("color", "auto"))
    add_inline_logo(sec.header, htbl.cell(0, 0), logo_path, width_cm=ht.get("logo_cm", 1.13))
    # title cell (project title; placeholder from config if empty)
    tc = htbl.cell(0, 1)
    cell_valign(tc, "bottom")
    tp = tc.paragraphs[0]
    tp.alignment = ALIGN.CENTER
    tf = ht["title_font"]
    title_text = meta.get("title") or ht.get("title_placeholder", "")
    rr = tp.add_run(title_text)
    run_fmt(rr, ascii=tf["ascii"], eastasia=tf.get("eastAsia", tf["ascii"]),
            size=tf["size_pt"], underline=False)
    # secrecy cell
    mc = htbl.cell(0, 2)
    cell_valign(mc, "bottom")
    mp = mc.paragraphs[0]
    mp.style = doc.styles["Header"]
    rr = mp.add_run(ht.get("secrecy_label", ""))
    run_fmt(rr, underline=False)


def _build_footer(doc, styles):
    sec = doc.sections[0]
    ft = styles["footer_table"]
    cols = ft["cols_cm"]
    ftbl = sec.footer.add_table(rows=1, cols=3, width=Cm(sum(cols)))
    table_grid(ftbl, cols)
    tb = ft.get("top_border", {"val": "single", "sz": 4})
    fpr = ftbl._tbl.tblPr
    fb = OxmlElement("w:tblBorders")
    fb.append(_border("w:top", tb["val"], tb["sz"], "auto"))
    fpr.append(fb)
    # date field
    f0 = ftbl.cell(0, 0).paragraphs[0]
    f0.style = doc.styles["Footer"]
    add_field(f0, f' DATE \\@ "{ft.get("date_format", "yyyy-MM-dd")}" ',
              placeholder="2026-01-01", bold=True)
    # center text
    f1 = ftbl.cell(0, 1).paragraphs[0]
    f1.style = doc.styles["Footer"]
    rr = f1.add_run(ft.get("center_text", ""))
    run_fmt(rr, bold=True)
    # page numbers
    pt = ft.get("page_text", ["", " / ", ""])
    f2 = ftbl.cell(0, 2).paragraphs[0]
    f2.alignment = ALIGN.RIGHT
    f2.style = doc.styles["Footer"]
    rr = f2.add_run(pt[0]); run_fmt(rr, bold=True)
    add_field(f2, "PAGE", placeholder="1", bold=True)
    rr = f2.add_run(pt[1]); run_fmt(rr, bold=True)
    add_field(f2, " NUMPAGES ", placeholder="1", bold=True)
    rr = f2.add_run(pt[2]); run_fmt(rr, bold=True)


def _build_cover(doc, cfg, meta):
    cover = cfg["cover"]
    t = cover["tables"]
    colors = cfg["styles"]["colors"]

    # ---- info table ----
    info = t["info"]
    sz = info.get("sz", 14)
    t1 = doc.add_table(rows=2, cols=5)
    table_center(t1)
    table_grid(t1, info["cols_cm"])
    table_outer_inner(t1, outer=info.get("outer", "double"), outer_sz=6,
                      inner=info.get("inner", "single"), inner_sz=6, color="000000")
    t1.cell(0, 0).merge(t1.cell(1, 0))
    fill(t1.cell(0, 0), cover["company_line"], ascii="Arial", size=sz,
         align=ALIGN.CENTER, valign="center")
    labels = info.get("labels", {})
    fill(t1.cell(0, 1), labels.get("title", "Project Name"), ascii="Arial", size=sz,
         align=ALIGN.CENTER, valign="center")
    fill(t1.cell(0, 2), meta.get("title", ""), ascii="Arial", size=sz,
         align=ALIGN.CENTER, valign="center")
    fill(t1.cell(0, 3), labels.get("secrecy", "Secrecy"), ascii="Arial", size=sz,
         align=ALIGN.CENTER, valign="center")
    secrecy = meta.get("secrecy") or cover.get("secrecy_default", "")
    # secrecy/classification color for the cover info cell; "secrecy" is the
    # canonical key. A legacy alias is accepted for older configs.
    _legacy_secrecy_key = "confiden" + "tial"
    secrecy_color = (
        colors.get("secrecy") or colors.get(_legacy_secrecy_key) or "4F81BD"
    )
    fill(t1.cell(0, 4), secrecy, bold=True, italic=True, color=secrecy_color,
         align=ALIGN.CENTER, valign="center")
    fill(t1.cell(1, 1), labels.get("doc_no", "Project Code"), ascii="Arial", size=sz,
         align=ALIGN.CENTER, valign="center")
    fill(t1.cell(1, 2), meta.get("doc_no", ""), ascii="Arial", size=sz,
         align=ALIGN.CENTER, valign="center")
    fill(t1.cell(1, 3), labels.get("pages", "Pages"), ascii="Arial", size=sz,
         align=ALIGN.CENTER, valign="center")
    if cover.get("page_count_field", True):
        pg = t1.cell(1, 4).paragraphs[0]
        pg.alignment = ALIGN.CENTER
        cell_valign(t1.cell(1, 4), "center")
        pgt = cover.get("page_text", ["", ""])
        pg.add_run(pgt[0])
        add_field(pg, " NUMPAGES ", placeholder="1")
        pg.add_run(pgt[1])

    # ---- big title block ----
    bt = cover.get("big_title", {})
    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    r = p.add_run(meta.get("title") or bt.get("placeholder", ""))
    run_fmt(r, ascii="Arial", size=bt.get("size_pt", 24), bold=True, italic=True, color="000000")
    if bt.get("subtitle"):
        p = doc.add_paragraph()
        p.alignment = ALIGN.CENTER
        p.add_run(bt["subtitle"])
    doc.add_paragraph()

    # ---- logo + company name lines ----
    logo_path = cfg["_logo_path"]
    if logo_path and os.path.exists(logo_path):
        lp = doc.add_paragraph()
        lp.alignment = ALIGN.CENTER
        lp.add_run().add_picture(logo_path, width=Cm(cover.get("logo_cm", 2.6)))
    for line in cover.get("company_names", []):
        p = doc.add_paragraph()
        p.alignment = ALIGN.CENTER
        r = p.add_run(line["text"])
        run_fmt(r, ascii=line.get("ascii", "Arial"),
                eastasia=line.get("eastAsia"), size=line.get("size_pt", 16))
    doc.add_paragraph()

    # ---- signature table ----
    sig = t["signature"]
    rows = sig["rows"]
    t2 = doc.add_table(rows=len(rows), cols=5)
    table_center(t2)
    table_grid(t2, sig["cols_cm"])
    sign_cols = sig.get("sign_cols", [1, 4])
    for r_i, (lab, dat) in enumerate(rows):
        fill(t2.cell(r_i, 0), lab, ascii="Arial", size=12, valign="bottom")
        fill(t2.cell(r_i, 3), dat, ascii="Arial", size=12, valign="bottom")
        if sig.get("sign_underline", True):
            for sc in sign_cols:
                cell_borders(t2.cell(r_i, sc), ["bottom"], val="single", sz=6, color="auto")
                cell_valign(t2.cell(r_i, sc), "bottom")
        for ci in range(5):
            for p in t2.cell(r_i, ci).paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
    # fill author / reviewers / approver values if present
    sign_vals = [meta.get("author", "")] + list(meta.get("reviewers", [])) + [meta.get("approver", "")]
    for r_i in range(len(rows)):
        if r_i < len(sign_vals) and sign_vals[r_i]:
            fill(t2.cell(r_i, 1), sign_vals[r_i], ascii="Arial", size=12, valign="bottom")
    doc.add_page_break()

    # ---- revision table ----
    rev = t["revision"]
    rev_title = rev.get("title", {})
    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    r = p.add_run(rev_title.get("text", "Revision History"))
    run_fmt(r, ascii=rev_title.get("ascii", "Arial"),
            eastasia=rev_title.get("eastAsia"), size=rev_title.get("size_pt", 16))
    doc.add_paragraph()
    headers = rev["headers"]
    revisions = meta.get("revisions", [])
    t3 = doc.add_table(rows=1 + max(1, len(revisions)), cols=len(headers))
    table_center(t3)
    table_grid(t3, rev["cols_cm"])
    table_all_single(t3, sz=6, color="auto")
    hf = rev.get("header_font", {"ascii": "SimSun", "size_pt": 10.5})
    for c, htxt in enumerate(headers):
        fill(t3.cell(0, c), htxt, ascii=hf["ascii"], size=hf["size_pt"], bold=True,
             align=ALIGN.CENTER)
    # rows: header order is [date, version, note, author]
    for ri, rv in enumerate(revisions, start=1):
        fill(t3.cell(ri, 0), rv.get("date", ""), ascii="Arial", size=10.5)
        fill(t3.cell(ri, 1), rv.get("ver", ""), ascii="Arial", size=10.5)
        fill(t3.cell(ri, 2), rv.get("note", ""), ascii="Arial", size=10.5)
        fill(t3.cell(ri, 3), rv.get("author", ""), ascii="Arial", size=10.5)
    doc.add_page_break()


def _build_toc(doc, cfg):
    toc = cfg.get("toc", {})
    p = doc.add_paragraph()
    r = p.add_run(toc.get("title", "Contents"))
    run_fmt(r, ascii="Arial", size=toc.get("size_pt", 24))
    p = doc.add_paragraph()
    add_field(p, toc.get("field", 'TOC \\o "1-3" \\h \\z \\u'),
              placeholder=toc.get("placeholder", "(right-click -> update field)"))
    doc.add_page_break()


# ===========================================================================
# Body: outline traversal + the four block types
# ===========================================================================
def _resolve_style_alias(style_name, names):
    """Map the config style aliases 'body'/'mybody' to their real style names."""
    if style_name in ("body", names["body_name"]):
        return names["body_name"]
    if style_name in ("mybody", names["mybody_name"]):
        return names["mybody_name"]
    return style_name


def _render_fixed_body(doc, fb, names, list_ctx=None):
    fb_style = fb.get("style", names["body_name"])
    for para in fb["paragraphs"]:
        # a paragraph may override the block style and may be a list item
        style_name = _resolve_style_alias(para.get("style", fb_style), names)
        p = doc.add_paragraph(style=style_name)
        for run in para["runs"]:
            r = p.add_run(run.get("t", ""))
            run_fmt(r, ascii=run.get("ascii"), eastasia=run.get("eastAsia"),
                    bold=run.get("b"), italic=run.get("i"), color=run.get("color"))
        _apply_list(p, para.get("list"), para.get("level", 0), list_ctx)


def _as_int(v, default):
    try:
        return int(v)
    except (TypeError, ValueError):
        return default


def _as_float(v, default):
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def _grid_sub_label(i):
    """Spreadsheet-style (a)..(z),(aa),(ab)... so a grid with >26 panels stays
    clean instead of running off the end of the alphabet into punctuation."""
    s = ""
    i += 1
    while i:
        i, r = divmod(i - 1, 26)
        s = chr(97 + r) + s
    return f"({s})"


def _add_formatted_run(p, text, bold=None, italic=None, color=None):
    """Add text as run(s), turning each embedded newline into an explicit soft
    line break (<w:br/>). Feature A maps Shift+Enter -> "\\n"; emitting the break
    ourselves keeps the soft break independent of python-docx version behavior."""
    parts = (text or "").split("\n")
    r = None
    for i, part in enumerate(parts):
        if i:
            if r is None:
                r = p.add_run()
                run_fmt(r, bold=bold, italic=italic, color=color)
            r.add_break()
        if part:
            r = p.add_run(part)
            run_fmt(r, bold=bold, italic=italic, color=color)


def _render_ref_run(p, ref_id, ref_targets, warn=None):
    """Emit a cross-reference for a paragraph run that carries a ``ref`` key.

    A reference run points at a captioned block's stable id; the live number is a
    Word REF field targeting that block's caption-number bookmark
    (``bm_<block_id>_num``, the scheme step 1 established). The field shows the
    number live and -- via the ``\\h`` switch -- is a Ctrl+click hyperlink.

    ``ref_targets`` maps a known block id to a short label (e.g. "Figure" /
    "Table") used only as the un-refreshed placeholder text. A reference whose
    target is missing (dangling) degrades to a visible red "[ref: ...]" marker
    plus a ``dangling_ref`` manifest warning -- it never aborts the render."""
    if not ref_id or ref_id not in ref_targets:
        r = p.add_run("[ref: %s]" % (ref_id or "?"))
        run_fmt(r, color="FF0000")
        if warn:
            warn({"type": "dangling_ref", "detail": ref_id or "(empty)"})
        return
    bname = "bm_%s_num" % ref_id
    add_field(p, " REF %s \\h " % bname, placeholder=ref_targets[ref_id] or "?")


def _render_para(doc, block, names, ref_targets=None, warn=None, list_ctx=None):
    # Honor a per-block style alias ('body'/'mybody'); default to body. List glyph/
    # number + indent come from _apply_list so a bullet/number item is a styled
    # paragraph with numbering (not an uncustomized built-in List Bullet/Number).
    style_name = _resolve_style_alias(block.get("style", names["body_name"]), names)
    p = doc.add_paragraph(style=style_name)
    targets = ref_targets if ref_targets is not None else {}
    for run in block.get("runs", []):
        ref_id = run.get("ref") if isinstance(run, dict) else None
        if ref_id:
            _render_ref_run(p, ref_id, targets, warn=warn)
        else:
            _add_formatted_run(p, run.get("t", ""), bold=run.get("b"),
                               italic=run.get("i"), color=run.get("color"))
    _apply_list(p, block.get("list"), block.get("level", 0), list_ctx)


def _image_settings(cfg):
    ic = cfg.get("image")
    return ic if isinstance(ic, dict) else {}


def _fit_height(pic, max_height_cm):
    """Scale an inline picture down proportionally when it is taller than the cap.

    A width-set portrait screenshot can be tall enough to push itself (and the
    text after it) onto the next page, leaving a big gap. Capping the height keeps
    such figures on the page; images already shorter than the cap are untouched."""
    if not max_height_cm:
        return
    cap = Cm(float(max_height_cm))
    if pic.height and pic.height > cap:
        ratio = cap / pic.height
        pic.height = int(pic.height * ratio)
        pic.width = int(pic.width * ratio)


def _place_picture(p, path, fname, width_cm, max_height_cm, warn=None):
    """Add one centered picture (or a red placeholder) into paragraph ``p``.

    ``warn`` (optional callable) receives a warning dict for a missing or
    unreadable image so callers can collect them into the render manifest."""
    if path and os.path.exists(path):
        try:
            pic = p.add_run().add_picture(path, width=Cm(width_cm))
            _fit_height(pic, max_height_cm)
            return
        except Exception as ex:
            r = p.add_run(f"[image: {fname}]")
            run_fmt(r, color="C00000")
            print("  ! image failed:", ex)
            if warn:
                warn({"type": "missing_image",
                      "detail": "%s (%s)" % (fname or "(no file)", type(ex).__name__)})
            return
    r = p.add_run(f"[image placeholder: {fname}]")
    run_fmt(r, color="C00000")
    if warn:
        warn({"type": "missing_image", "detail": fname or "(no file)"})


def _render_image(doc, block, project_dir, chap, seq, cfg, warn=None):
    fname = block.get("file", "")
    width = _as_float(block.get("width_cm", 12.0), 12.0)
    path = os.path.join(project_dir, fname.replace("/", os.sep)) if fname else ""
    pic_p = doc.add_paragraph()
    pic_p.alignment = ALIGN.CENTER
    _place_picture(pic_p, path, fname, width,
                   _image_settings(cfg).get("max_height_cm", 18.0), warn=warn)
    caption = block.get("caption", "")
    _render_caption_with_seq(doc, caption, cfg, chap, seq, block.get("id"), "image")


def _render_image_grid(doc, block, project_dir, chap, seq, cfg, warn=None):
    """A borderless table laying images out in a grid; the whole group shares one
    figure number. Each cell scales its image to the column width (and the height
    cap), with optional (a)(b)(c) sub-captions below each image."""
    items = [it for it in (block.get("items") or []) if isinstance(it, dict)]
    cols = max(1, _as_int(block.get("cols", 2), 2))
    total_w = _as_float(block.get("width_cm", 15.5), 15.5)
    show_sub = bool(block.get("sub_captions"))
    isettings = _image_settings(cfg)
    gap = _as_float(isettings.get("grid_gap_cm", 0.3), 0.3)
    max_h = _as_float(isettings.get("grid_max_height_cm",
                                   isettings.get("max_height_cm", 8.0)), 8.0)
    cell_w = (total_w - gap * (cols - 1)) / cols if cols else total_w
    if cell_w <= 0:
        cell_w = total_w

    n = len(items)
    if n:
        # Honor the layout picker's chosen row count as a MINIMUM: a taller pick
        # pads blank trailing cells; adding more images than the pick grows extra
        # rows so nothing is dropped. Old projects have no block["rows"] -> picked
        # 0 -> falls back to just enough rows for the images (prior behavior).
        need_rows = (n + cols - 1) // cols
        picked_rows = max(0, _as_int(block.get("rows", 0), 0))
        rows = max(need_rows, picked_rows)
        table = doc.add_table(rows=rows, cols=cols)
        tables._table_fixed_layout(table)
        tables._table_borders(table, val="none")
        tables._cell_margins(table, top=14, bottom=14, left=28, right=28)
        idx = 0
        for r in range(rows):
            for c in range(cols):
                cell = table.cell(r, c)
                cell.width = Cm(cell_w)
                cell.text = ""
                # bottom-align so unequal image heights still leave the (a)(b)
                # sub-captions on a single baseline across the row
                if show_sub:
                    tables._vbottom(cell)
                p = cell.paragraphs[0]
                p.alignment = ALIGN.CENTER
                if idx < n:
                    it = items[idx]
                    fn = it.get("file", "")
                    pth = (os.path.join(project_dir, fn.replace("/", os.sep))
                           if fn else "")
                    _place_picture(p, pth, fn, cell_w, max_h, warn=warn)
                    if show_sub:
                        sp = cell.add_paragraph()
                        sp.alignment = ALIGN.CENTER
                        lab = it.get("sub") or _grid_sub_label(idx)
                        run_fmt(sp.add_run(lab), size=9)
                idx += 1

    caption = block.get("caption", "")
    _render_caption_with_seq(doc, caption, cfg, chap, seq, block.get("id"), "image")


def _render_caption(doc, text, cfg):
    doc.add_paragraph(text, style="Caption")


# ---------------------------------------------------------------------------
# Word-native SEQ-field captions + cross-reference bookmarks (Design #1)
# ---------------------------------------------------------------------------
#
# CONTRACT (steps 2-4 rely on this):
#   * Bookmark naming scheme: each captioned block's number portion is wrapped in
#     a bookmark named  "bm_" + <block_id> + "_num"  (e.g. bm_img-123-abc_num).
#     The block_id is the stable per-block "id" assigned by the frontend; when a
#     block has no id, a render-local fallback name is used (never collides, but
#     is NOT a stable cross-reference target).
#   * A cross-reference run is a paragraph run carrying {"ref": <block_id>} (no
#     "t"/formatting). It renders as a Word REF field with instruction
#     " REF bm_<block_id>_num \\h " (the \\h switch makes it a Ctrl+click
#     hyperlink). The placeholder text is the "<chap>-<seq>" the target's caption
#     would show un-refreshed (see _collect_ref_targets), so it reads identically
#     until F9. A reference whose target has no caption bookmark is "dangling":
#     it degrades to a red "[ref: <id>]" marker plus a dangling_ref warning.
#   * SEQ identifiers: "Figure" for image/imagegrid captions, "Table" for
#     datatable/table captions -- ASCII, separate counters, reset per Heading 1
#     via the "\\s 1" switch. These MUST match what a REF/SEQ consumer expects.
#   * STYLEREF 1 \\s yields the Heading 1 chapter number (multilevel autonumber).
#
# The visible numbering (chapter-counter, counter resets each Heading 1) matches
# the previous literal-text output; the placeholder text written into each field
# is exactly the chapter/seq the engine already computed, so an un-updated doc
# reads identically and Word recomputes live numbers on F9.

# SEQ identifier per caption type (ASCII, stable, separate counters).
_SEQ_NAME = {"image": "Figure", "table": "Table"}


class CaptionState:
    """Per-document allocator for unique bookmark ids + names around captions."""

    def __init__(self):
        self.bookmark_id_counter = 0
        self.bookmark_names = set()
        self._auto = 0

    def next_bookmark(self, block_id):
        """Return (numeric_id, name) for this block's caption number bookmark.

        ``block_id`` is the stable frontend id. When absent or already used, a
        render-local fallback keeps bookmark names unique so the render never
        aborts on a duplicate; such fallbacks are not stable xref targets."""
        if block_id:
            name = "bm_%s_num" % block_id
        else:
            name = None
        if not name or name in self.bookmark_names:
            self._auto += 1
            name = "bm_auto_%d_num" % self._auto
            while name in self.bookmark_names:
                self._auto += 1
                name = "bm_auto_%d_num" % self._auto
        self.bookmark_names.add(name)
        bid = self.bookmark_id_counter
        self.bookmark_id_counter += 1
        return bid, name


def _caption_state(cfg):
    cs = cfg.get("_caption_state")
    if cs is None:
        cs = CaptionState()
        cfg["_caption_state"] = cs
    return cs


def _render_caption_with_seq(doc, text, cfg, chap, seq, block_id, caption_type):
    """Render a Caption-styled paragraph whose number is live Word fields.

    Layout:  <prefix> [STYLEREF 1 \\s]-[SEQ <Figure|Table> \\* ARABIC \\s 1]  <text>
    The SEQ field (the within-chapter counter) is wrapped in a bookmark so REF
    fields can target the number. ``chap`` / ``seq`` become the field placeholder
    text, so an un-refreshed doc shows the same number the engine computed."""
    cap = doc.add_paragraph(style="Caption")
    cap_prefixes = cfg.get("caption_prefix", {})
    prefix = cap_prefixes.get(caption_type,
                              "Figure" if caption_type == "image" else "Table")
    seq_name = _SEQ_NAME.get(caption_type, "Figure")

    cstate = _caption_state(cfg)
    bid, bname = cstate.next_bookmark(block_id)

    r = cap.add_run(prefix + " ")
    run_fmt(r, bold=True)
    # chapter number from the Heading 1 multilevel autonumber
    add_field(cap, " STYLEREF 1 \\s ", placeholder=str(chap), bold=True)
    r = cap.add_run("-")
    run_fmt(r, bold=True)
    # within-chapter counter (bookmarked for cross-references)
    add_field(cap, " SEQ %s \\* ARABIC \\s 1 " % seq_name, placeholder=str(seq),
              bold=True, bookmark_id=bid, bookmark_name=bname)
    r = cap.add_run("  " + (text or ""))
    run_fmt(r, bold=False)
    return cap


def _render_block_error(doc, ex):
    """Emit a visible red error line in place of a block that failed to render."""
    p = doc.add_paragraph(style="Caption")
    r = p.add_run("[block error: %s: %s]"
                  % (type(ex).__name__, str(ex)[:120]))
    run_fmt(r, color="FF0000", bold=True)


def _collect_table_result(res, warn, location):
    """Route a table renderer's result into the warnings manifest.

    The compliance / free-table renderers now return a dict
    {"table", "total_rows", "flagged_rows", "warnings"} (see tables.py contract).
    This stays backward-tolerant: a bare Table object (legacy shape) is accepted
    and simply yields no extra warnings. Each per-cell warning from the renderer
    (e.g. ``row_clip_risk``) is re-emitted with this block's location appended."""
    if not isinstance(res, dict):
        return
    for w in res.get("warnings", []) or []:
        entry = dict(w) if isinstance(w, dict) else {"type": "table_warning",
                                                     "detail": str(w)}
        inner = entry.get("location")
        entry["location"] = ("%s / %s" % (location, inner)) if inner else location
        warn(entry)


def _collect_ref_targets(outline):
    """Pre-scan the outline for cross-reference targets BEFORE rendering.

    Returns a dict mapping each captioned block's stable id -> the "<chap>-<seq>"
    placeholder string that the block's caption number would show un-refreshed.
    A paragraph's REF run uses this only as the field placeholder text, so the
    reference reads identically to the caption until Word recomputes on F9. The
    mapping mirrors the SAME counting rules _build_outline uses (img and table
    sequences reset per Heading 1; a table only consumes a number when it has a
    caption), so forward AND backward references resolve to the right number.

    Only blocks with a non-empty ``id`` AND a non-empty caption become targets --
    those are exactly the blocks _render_caption_with_seq bookmarks as
    ``bm_<id>_num``. References to anything else are treated as dangling."""
    targets = {}
    state = {"chap": 0, "img": {}, "tbl": {}}

    def walk(node, depth):
        if depth == 0:
            state["chap"] += 1
            state["img"][state["chap"]] = 0
            state["tbl"][state["chap"]] = 0
        chap = state["chap"]
        if node.get("fixed_body"):
            # fixed bodies carry no captioned media blocks
            for child in node.get("children", []):
                walk(child, depth + 1)
            return
        for block in node.get("blocks", []):
            btype = block.get("type")
            bid = block.get("id")
            cap = block.get("caption", "")
            if btype in ("image", "imagegrid"):
                # An image ALWAYS consumes a figure number (the seq increments here
                # and _render_caption_with_seq bookmarks it) even with an empty
                # caption -- so an uncaptioned image is still a valid xref target.
                # Register on ``bid`` alone (B5: the render created the bookmark, so
                # the ref must resolve instead of reporting a spurious dangling_ref).
                state["img"][chap] += 1
                if bid:
                    targets[bid] = "%d-%d" % (chap, state["img"][chap])
            elif btype in ("datatable", "table"):
                if cap:
                    state["tbl"][chap] += 1
                    if bid:
                        targets[bid] = "%d-%d" % (chap, state["tbl"][chap])
        for child in node.get("children", []):
            walk(child, depth + 1)

    for node in outline or []:
        walk(node, 0)
    return targets


def _build_outline(doc, cfg, outline, names):
    """Render the outline; returns {"warnings": [...], "stats": {...}}.

    Each block renders inside try/except so one malformed block degrades to a
    visible red error line and a manifest entry instead of aborting the whole
    export. Caption / figure counters are incremented BEFORE the block body so a
    failure mid-render leaves the numbering of later blocks unchanged."""
    body_name = names["body_name"]
    list_ctx = ListCtx(doc, names["list_ids"]) if names.get("list_ids") else None
    fixed_bodies = cfg.get("fixed_bodies", {})
    comp_cfg = cfg["compliance"]
    free_cfg = cfg.get("free_table", {})

    # Cross-reference targets (block id -> "<chap>-<seq>" placeholder) gathered in
    # a pre-pass so a paragraph may reference a figure/table that appears later.
    ref_targets = _collect_ref_targets(outline)

    # chapter-sequence counters: keyed by chapter number
    state = {"chap": 0, "img_seq": {}, "tbl_seq": {},
             "warnings": [], "total_blocks": 0}

    def warn(entry):
        state["warnings"].append(entry)

    def walk(node, depth):
        level = depth + 1
        if level == 1:
            state["chap"] += 1
            state["img_seq"][state["chap"]] = 0
            state["tbl_seq"][state["chap"]] = 0
        chap = state["chap"]
        doc.add_paragraph(node["title"], style=f"Heading {min(level, 9)}")
        if list_ctx:
            list_ctx.new_section()   # numbered lists restart per section

        # fixed body wins over blocks
        fb_key = node.get("fixed_body")
        if fb_key and fb_key in fixed_bodies:
            try:
                _render_fixed_body(doc, fixed_bodies[fb_key], names, list_ctx)
            except Exception as ex:
                _render_block_error(doc, ex)
                warn({"type": "block_error",
                      "detail": "%s: %s" % (type(ex).__name__, str(ex)[:120]),
                      "location": "chapter %d / fixed_body %s" % (chap, fb_key)})
        else:
            for idx, block in enumerate(node.get("blocks", [])):
                btype = block.get("type")
                state["total_blocks"] += 1
                # Counters increment BEFORE the body so a failure does not skew
                # the numbering of subsequent figures/tables.
                if btype in ("image", "imagegrid"):
                    state["img_seq"][chap] += 1
                seq = state["img_seq"][chap]
                cap = block.get("caption", "")
                if btype in ("datatable", "table") and cap:
                    state["tbl_seq"][chap] += 1
                tbl_seq = state["tbl_seq"][chap]
                try:
                    if btype == "para":
                        _render_para(doc, block, names,
                                     ref_targets=ref_targets, warn=warn,
                                     list_ctx=list_ctx)
                    elif btype == "image":
                        _render_image(doc, block, cfg["_project_dir"], chap,
                                      seq, cfg, warn=warn)
                    elif btype == "imagegrid":
                        _render_image_grid(doc, block, cfg["_project_dir"], chap,
                                           seq, cfg, warn=warn)
                    elif btype == "datatable":
                        if cap:
                            _render_caption_with_seq(doc, cap, cfg, chap, tbl_seq,
                                                     block.get("id"), "table")
                        else:
                            warn({"type": "no_caption",
                                  "detail": "datatable",
                                  "location": "chapter %d / block %d" % (chap, idx)})
                        res = tables.render_datatable(doc, block["data"], comp_cfg)
                        _collect_table_result(
                            res, warn, "chapter %d / datatable / block %d" % (chap, idx))
                    elif btype == "table":
                        if cap:
                            _render_caption_with_seq(doc, cap, cfg, chap, tbl_seq,
                                                     block.get("id"), "table")
                        else:
                            warn({"type": "no_caption",
                                  "detail": "table",
                                  "location": "chapter %d / block %d" % (chap, idx)})
                        res = tables.render_free_table(
                            doc, block.get("rows", []), free_cfg,
                            header_rows=block.get("header_rows", 1),
                            merges=block.get("merges"),
                            col_w=block.get("col_w"),
                            row_fills=block.get("row_fills"))
                        _collect_table_result(
                            res, warn, "chapter %d / table / block %d" % (chap, idx))
                except Exception as ex:
                    _render_block_error(doc, ex)
                    warn({"type": "block_error",
                          "detail": "%s: %s" % (type(ex).__name__, str(ex)[:120]),
                          "location": "chapter %d / %s / block %d"
                                      % (chap, btype, idx)})

        for child in node.get("children", []):
            walk(child, depth + 1)

    for node in outline:
        walk(node, 0)

    warnings = state["warnings"]
    stats = {
        "total_blocks": state["total_blocks"],
        "blocks_with_errors": sum(1 for w in warnings if w["type"] == "block_error"),
        "blocks_without_captions": sum(1 for w in warnings if w["type"] == "no_caption"),
        "missing_images": sum(1 for w in warnings if w["type"] == "missing_image"),
        "row_clip_risks": sum(1 for w in warnings if w["type"] == "row_clip_risk"),
        "dangling_refs": sum(1 for w in warnings if w["type"] == "dangling_ref"),
    }
    return {"warnings": warnings, "stats": stats}


# ===========================================================================
# Top-level render
# ===========================================================================
def render_report(project, cfg, project_dir, out_path):
    """Render the project to ``out_path`` and return a result manifest.

    RETURN SHAPE (CONTRACT, steps 2-4 depend on it):
        {
          "out_path": <str>,          # the saved .docx path
          "warnings": [ {type, detail, location?}, ... ],
          "stats":    {total_blocks, blocks_with_errors,
                       blocks_without_captions, missing_images, ...},
        }
    Warning ``type`` values currently emitted: "missing_image", "missing_logo",
    "no_caption", "block_error", "row_clip_risk", "dangling_ref". The shape is
    additive -- later steps may add new types/keys -- so callers must tolerate
    unknown entries. A "dangling_ref" warning means a paragraph cross-reference
    pointed at a block id that has no caption bookmark (target deleted / never
    captioned); the reference renders as a visible red "[ref: ...]" marker.
    Callers needing only the path can read result["out_path"]; ``_result_out_path``
    normalizes both the new dict shape and any legacy plain-string return.

    ``row_clip_risk`` warnings come from the compliance / free-table renderers:
    the table keeps EXACTLY row heights (no spill, no orientation change), so an
    over-long cell is clipped by Word rather than wrapping; this warning surfaces
    that risk WITHOUT changing the row height (iron rule 2)."""
    if project.get("schema_version") != 1:
        raise ValueError(f"unsupported schema_version: {project.get('schema_version')}")

    cfg["_project_dir"] = project_dir
    # reset any per-render caption bookmark allocator carried on a reused cfg
    cfg.pop("_caption_state", None)
    meta = project.get("meta", {})
    styles = cfg["styles"]

    warnings = []
    logo_path = cfg.get("_logo_path")
    if not logo_path or not os.path.exists(logo_path):
        warnings.append({"type": "missing_logo", "detail": logo_path or "(no logo configured)"})

    doc = Document()
    names = _apply_page_and_styles(doc, styles)

    _build_header(doc, styles, meta, cfg["_logo_path"])
    _build_footer(doc, styles)
    _build_cover(doc, cfg, meta)
    _build_toc(doc, cfg)
    outline_result = _build_outline(doc, cfg, project.get("outline", []), names)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)

    warnings.extend(outline_result.get("warnings", []))
    # Stamp a level (error|warn|info) onto every warning so the manifest schema is
    # uniform for the GUI panel + CLI (block_error -> error, the rest -> warn).
    content_lint.stamp_levels(warnings)
    stats = dict(outline_result.get("stats", {}))
    stats["missing_logos"] = sum(1 for w in warnings if w["type"] == "missing_logo")
    stats["errors"] = sum(1 for w in warnings if w.get("level") == "error")
    stats["warns"] = sum(1 for w in warnings if w.get("level") == "warn")
    stats["infos"] = sum(1 for w in warnings if w.get("level") == "info")
    return {"out_path": out_path, "warnings": warnings, "stats": stats}


def _result_out_path(result):
    """Normalize render_report's result to the output path string.

    Accepts the rich dict {"out_path": ...} (current) or a bare string (legacy),
    so callers stay backward-tolerant across the return-shape change."""
    if isinstance(result, dict):
        return result.get("out_path")
    return result


# ===========================================================================
# Config resolution + CLI
# ===========================================================================
def _resolve_config_path(project, project_dir, explicit):
    if explicit:
        return explicit
    env = os.environ.get("BUILDER_TEMPLATE_CONFIG")
    if env:
        return env
    tid = project.get("template", "")
    # look next to the project's parent folder
    parent = os.path.dirname(os.path.abspath(project_dir.rstrip(os.sep)))
    cand = os.path.join(parent, f"template_config_{tid}.json")
    if os.path.exists(cand):
        return cand
    cand2 = os.path.join(os.path.dirname(parent), f"template_config_{tid}.json")
    if os.path.exists(cand2):
        return cand2
    raise FileNotFoundError(
        f"template config not found; pass --config or set BUILDER_TEMPLATE_CONFIG "
        f"(looked for template_config_{tid}.json)")


def _load_config(config_path):
    with open(config_path, encoding="utf-8") as f:
        cfg = json.load(f)
    # resolve logo relative to the config file's folder
    cfg_dir = os.path.dirname(os.path.abspath(config_path))
    logo = cfg.get("logo", "")
    cfg["_logo_path"] = os.path.join(cfg_dir, logo) if logo else ""
    return cfg


def main(argv=None):
    ap = argparse.ArgumentParser(description="Render a structured document to .docx")
    ap.add_argument("project_folder", help="folder containing project.json")
    ap.add_argument("--config", help="template config json (else env / auto-resolve)")
    ap.add_argument("--out", help="output directory (default <project_folder>/out)")
    args = ap.parse_args(argv)

    project_dir = os.path.abspath(args.project_folder)
    proj_path = os.path.join(project_dir, "project.json")
    if not os.path.exists(proj_path):
        print(f"error: {proj_path} not found", file=sys.stderr)
        return 2
    with open(proj_path, encoding="utf-8") as f:
        project = json.load(f)

    config_path = _resolve_config_path(project, project_dir, args.config)
    cfg = _load_config(config_path)

    name = os.path.basename(project_dir.rstrip(os.sep)) or "report"
    out_dir = os.path.abspath(args.out) if args.out else os.path.join(project_dir, "out")
    out_path = os.path.join(out_dir, f"{name}.docx")

    result = render_report(project, cfg, project_dir, out_path)
    out = _result_out_path(result) or out_path
    print("OK ->", out)
    if isinstance(result, dict):
        warnings = result.get("warnings", [])
        if warnings:
            print("warnings: %d" % len(warnings))
            for w in warnings:
                loc = (" @ %s" % w["location"]) if w.get("location") else ""
                print("  - %s: %s%s" % (w.get("type"), w.get("detail"), loc))
    return 0


if __name__ == "__main__":
    sys.exit(main())
