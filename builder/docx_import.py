# -*- coding: utf-8 -*-
"""
docx_import.py -- read a designer .docx and produce two best-effort seeds.

This is a pure-Python parsing core (no server / no GUI). It turns a structured
review document into either:

  1. a PROJECT seed  -- ``parse_docx_report(path)`` -> {"meta", "outline"}.
     The document is split at the first level-1 heading into a FRONT-MATTER layer
     (cover / sign-off / revision tables, table of contents) and a BODY layer
     (chapters, prose, data tables, images). Front-matter values are lifted into
     ``meta``; the body becomes a nested outline tree of blocks.

  2. a TEMPLATE seed -- ``derive_template(path)`` -> {"config", "skeleton"}.
     ``config`` is exactly what ``make_config`` extracts (styles, cover layout,
     compliance styling, caption prefixes, logo). ``skeleton`` is the body outline
     tree with each node marked ``fixed`` (boilerplate that repeats across every
     report, carried verbatim) or fillable (placeholder body).

Design notes:
  * Heading detection, title cleaning, custom body-style resolution, cover-table
    shape predicates, unit conversion and the whole template-config extraction are
    REUSED from ``make_config`` -- never reimplemented here.
  * The compliance data-table model is the one produced by
    ``server.parse_xlsx_compliance``; its grid->model core is ported here as
    ``_compliance_model_from_grid`` (source of truth: that function) so this module
    stays standalone with no server import. The function and axis-label set are kept
    identical so a later merge is mechanical.
  * Everything is best-effort: every extractor is guarded so malformed input yields
    a partial result instead of a crash.

The source code, comments and default strings are intentionally English and
neutral. Any company / domain text only ever appears as DATA copied out of the
document being read, never as a literal in this file.

CLI (the test harness)::

    python docx_import.py <docx> [--mode report|template] [--images-dir DIR]

prints the resulting dict as indented JSON to stdout; extracted images are written
to ``--images-dir`` (default a fresh temp dir).
"""
import argparse
import json
import os
import re
import sys
import tempfile
import zipfile

from docx import Document
from docx.oxml.ns import qn

# make_config lives next to this file; mirror server.py's sys.path trick so the
# module imports the same way whether run as a script or imported.
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)
import make_config as mc  # noqa: E402


# ===========================================================================
# Layer split (front matter vs body)
# ===========================================================================
def _split_layers(document, by_id):
    """Split the body element's children at the first level-1 heading.

    Returns ``(front_children, body_children, first_h1_index)``. The body layer
    starts AT the first level-1 heading paragraph (it is the first chapter). If no
    level-1 heading exists, everything is front matter and the body is empty.
    """
    children = list(document.element.body)
    first_h1 = None
    for idx, child in enumerate(children):
        if child.tag == qn("w:p") and mc._heading_level(by_id, child) == 1:
            first_h1 = idx
            break
    if first_h1 is None:
        return children, [], None
    return children[:first_h1], children[first_h1:], first_h1


# ===========================================================================
# Front-matter tables -> meta
# ===========================================================================
def _front_tables(document, by_id):
    """Classify the front-matter tables by shape into (info, signature, revision).

    Reuses ``make_config._tables_before_first_heading`` (which stops at the first
    level-1 heading, so a body data table cannot be grabbed) and the same shape
    predicates ``extract_cover`` uses: info = 2 rows x >=5 cols, signature = 3 rows
    x >=5 cols, revision = 4 cols x >=2 rows.
    """
    candidates = mc._tables_before_first_heading(document, by_id)
    if not candidates:
        # heading-detection miss: fall back to every table rather than blank meta.
        candidates = list(document.tables)
    info_t = sig_t = rev_t = None
    for t in candidates:
        try:
            rows, cols = len(t.rows), len(t.columns)
        except Exception:
            continue
        if info_t is None and rows == 2 and cols >= 5:
            info_t = t
        elif sig_t is None and rows == 3 and cols >= 5:
            sig_t = t
        elif rev_t is None and cols == 4 and rows >= 2:
            rev_t = t
    return info_t, sig_t, rev_t


def _meta_from_info(info_t, meta, warn):
    """Lift project-name / classification / document-no / page-count VALUES.

    Layout (merge-inflated by a banner cell at column 0):
        row0: <banner> | titleLabel  | titleValue | secrecyLabel | secrecyValue
        row1: <banner> | docNoLabel   | docNoValue | pagesLabel   | pagesValue
    make_config reads the LABELS at (.,1)/(.,3); we read the adjacent VALUE cells.
    Every read goes through ``_safe_cell_text`` so a merged/out-of-range access
    yields "" rather than raising.
    """
    try:
        meta["title"] = mc._safe_cell_text(info_t, 0, 2).strip()
        meta["secrecy"] = mc._safe_cell_text(info_t, 0, 4).strip()
        meta["doc_no"] = mc._safe_cell_text(info_t, 1, 2).strip()
        meta["page_count"] = mc._safe_cell_text(info_t, 1, 4).strip()
    except Exception as ex:  # pragma: no cover - defensive
        warn(f"meta: info table read failed ({ex!r})")


# Role keyword aliases for sign-off rows. There is no structural / style signal to
# tell author from reviewer from approver -- only the row label text -- so a small
# alias table is unavoidable. English aliases come first; the trailing entries are
# the localized labels the real fixtures use (kept minimal, here only because the
# label text is the sole discriminator). These are detection heuristics, not output.
# Localized aliases are written as \u escapes so the source carries no non-ASCII
# glyphs; the runtime strings are unchanged.
_AUTHOR_ALIASES = ("author", "\u4f5c\u8005")            # + "author"
_REVIEWER_ALIASES = ("review", "reviewer", "\u8bc4\u5ba1")  # + "review"
_APPROVER_ALIASES = ("approve", "approver", "\u6279\u51c6")  # + "approve"


def _classify_signoff_label(label):
    """-> "author" | "reviewer" | "approver" | None for a sign-off row label."""
    low = (label or "").strip().lower()
    if not low:
        return None
    for a in _AUTHOR_ALIASES:
        if a in low:
            return "author"
    for a in _REVIEWER_ALIASES:
        if a in low:
            return "reviewer"
    for a in _APPROVER_ALIASES:
        if a in low:
            return "approver"
    return None


# Date-label aliases used in the sign-off table (English + the CJK label real
# fixtures use). Detection-only: the VALUE is the cell to the right of the label.
_DATE_LABEL = ("date", "\u65e5\u671f")  # CJK "date"


def _signoff_date_value(sig_t, r, n_cols):
    """Return the date VALUE in row ``r``: the cell right of the date label.

    Falls back to "" when no date label is found, so a label cell ("date:") is
    never mistaken for the value.
    """
    for c in range(2, n_cols):
        try:
            cell = mc._safe_cell_text(sig_t, r, c).strip()
        except Exception:
            continue
        cl = cell.lower()
        if any(a in cl for a in _DATE_LABEL):
            try:
                return mc._safe_cell_text(sig_t, r, c + 1).strip()
            except Exception:
                return ""
    return ""


def _meta_from_signature(sig_t, meta, warn):
    """Lift author / reviewers[] / approver names from the sign-off table.

    Each row reads ``roleLabel | nameValue | ... | dateLabel | dateValue``; the
    name value sits at column 1, the date VALUE in the column just right of the
    date LABEL. Rows are classified by the role label keyword (language-neutral
    via the alias table above).
    """
    try:
        n_rows = len(sig_t.rows)
        n_cols = len(sig_t.columns)
    except Exception:
        n_rows = n_cols = 0
    for r in range(n_rows):
        try:
            label = mc._safe_cell_text(sig_t, r, 0).strip()
            name = mc._safe_cell_text(sig_t, r, 1).strip()
            date = _signoff_date_value(sig_t, r, n_cols)
        except Exception:
            continue
        role = _classify_signoff_label(label)
        if role == "author":
            meta["author"] = name
            if date:
                meta["author_date"] = date
        elif role == "reviewer":
            if name:
                meta["reviewers"].append(name)
        elif role == "approver":
            meta["approver"] = name
            if date:
                meta["approver_date"] = date


# Revision header keyword aliases (English + the CJK labels real fixtures use),
# for inferring which column holds the date / version / description / author.
_REV_DATE = ("date", "\u65e5\u671f")
_REV_VER = ("version", "ver", "rev", "\u4fee\u8ba2\u7248\u672c", "\u7248\u672c")
_REV_DESC = ("description", "desc", "note", "change", "\u63cf\u8ff0", "\u8bf4\u660e")
_REV_AUTHOR = ("author", "by", "\u4f5c\u8005")


def _infer_revision_columns(headers):
    """Map header cells to column indices for date/ver/note/author.

    Falls back to the de-facto positional layout (0=date, 1=ver, 2=note, 3=author)
    for any role a header keyword does not resolve.
    """
    cols = {"date": None, "ver": None, "note": None, "author": None}
    for ci, h in enumerate(headers):
        low = (h or "").strip().lower()
        if not low:
            continue
        if cols["date"] is None and any(a in low for a in _REV_DATE):
            cols["date"] = ci
        elif cols["ver"] is None and any(a in low for a in _REV_VER):
            cols["ver"] = ci
        elif cols["note"] is None and any(a in low for a in _REV_DESC):
            cols["note"] = ci
        elif cols["author"] is None and any(a in low for a in _REV_AUTHOR):
            cols["author"] = ci
    fallback = {"date": 0, "ver": 1, "note": 2, "author": 3}
    for k, v in cols.items():
        if v is None:
            cols[k] = fallback[k]
    return cols


def _meta_from_revision(rev_t, meta, warn):
    """Lift revision rows {ver, date, author, note}; row 0 is the header."""
    try:
        n_rows = len(rev_t.rows)
        n_cols = len(rev_t.columns)
    except Exception:
        return
    headers = [mc._safe_cell_text(rev_t, 0, c).strip() for c in range(n_cols)]
    cols = _infer_revision_columns(headers)
    for r in range(1, n_rows):
        try:
            cells = [mc._safe_cell_text(rev_t, r, c).strip() for c in range(n_cols)]
        except Exception:
            continue
        if not any(cells):
            continue  # skip fully-empty (trailing) rows

        def _at(key):
            ci = cols[key]
            return cells[ci] if ci is not None and ci < len(cells) else ""
        meta["revisions"].append({
            "ver": _at("ver"),
            "date": _at("date"),
            "author": _at("author"),
            "note": _at("note"),
        })


def extract_meta(document, by_id, warn):
    """Best-effort project meta from the front-matter tables. Never raises."""
    meta = {
        "title": "", "secrecy": "", "doc_no": "", "page_count": "",
        "author": "", "reviewers": [], "approver": "", "revisions": [],
    }
    try:
        info_t, sig_t, rev_t = _front_tables(document, by_id)
    except Exception as ex:
        warn(f"meta: front table detection failed ({ex!r})")
        return meta
    if info_t is not None:
        _meta_from_info(info_t, meta, warn)
    else:
        warn("meta: cover info table (2x>=5) not found")
    if sig_t is not None:
        _meta_from_signature(sig_t, meta, warn)
    else:
        warn("meta: sign-off table (3x>=5) not found")
    if rev_t is not None:
        _meta_from_revision(rev_t, meta, warn)
    else:
        warn("meta: revision table (4-col) not found")
    return meta


# ===========================================================================
# Table classification
# ===========================================================================
# The axis-label set is identical to server.parse_xlsx_compliance's ``AX``.
_AX = {"MIN", "TYP", "MAX", "NTWC", "NT WC", "NT-WC"}


def _cell_fill_hex(tc):
    """The w:shd fill hex of a cell tc element, upper-cased, or None."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill or fill in ("auto",):
        return None
    return fill.lstrip("#").upper()


def _is_yellowish(hex6):
    """True for a yellow-ish fill. Accepts exact FFFF00 plus a tolerance band
    (high red+green, low blue) so FFFF66 / theme-yellow variants also match."""
    if not hex6:
        return False
    h = hex6.lstrip("#").upper()
    if h == "FFFF00":
        return True
    if len(h) != 6:
        return False
    try:
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    except ValueError:
        return False
    return r >= 0xE0 and g >= 0xE0 and b <= 0x40


def classify_table(tbl):
    """-> "datatable" | "table".

    A table is a compliance ``datatable`` iff BOTH hold:
      (A) any cell in the first row has a yellow-ish shading fill, AND
      (B) any cell text in the first ~3 rows is an axis label (MIN/TYP/MAX/NTWC).
    Otherwise it is a free ``table``.
    """
    try:
        rows = tbl.rows
        if not len(rows):
            return "table"
    except Exception:
        return "table"

    yellow = False
    try:
        for cell in rows[0].cells:
            if _is_yellowish(_cell_fill_hex(cell._tc)):
                yellow = True
                break
    except Exception:
        yellow = False
    if not yellow:
        return "table"

    try:
        for r in range(min(3, len(rows))):
            for cell in rows[r].cells:
                if (cell.text or "").strip().upper() in _AX:
                    return "datatable"
    except Exception:
        return "table"
    return "table"


def _docx_table_to_grid(tbl):
    """Convert a docx table to an xlsx-style grid (list[list[str|None]]).

    python-docx returns the SAME merged tc text repeatedly across its span, which
    is exactly what the band logic wants (a merged group title reads identically
    across its columns, like a merged xlsx cell). Whitespace-only cells -> None.
    """
    grid = []
    try:
        n_rows = len(tbl.rows)
        n_cols = len(tbl.columns)
    except Exception:
        return grid
    for r in range(n_rows):
        row = []
        for c in range(n_cols):
            try:
                txt = tbl.cell(r, c).text
            except Exception:
                txt = ""
            txt = (txt or "").strip()
            row.append(txt if txt else None)
        grid.append(row)
    return grid


def _docx_table_row_fills(tbl):
    """Per-row dominant cell shading hex (upper-case) or None -> lets fills survive
    the round-trip (plain-table row_fills; datatable setting-row detection)."""
    out = []
    try:
        n_rows, n_cols = len(tbl.rows), len(tbl.columns)
    except Exception:
        return out
    for r in range(n_rows):
        hexes = []
        for c in range(n_cols):
            try:
                h = _cell_fill_hex(tbl.cell(r, c)._tc)
            except Exception:
                h = None
            if h:
                hexes.append(h.upper())
        out.append(max(set(hexes), key=hexes.count) if hexes else None)
    return out


def _is_setting_fill(hexv):
    """A non-white, non-header shading marks a datatable setting / condition row."""
    if not hexv:
        return False
    h = hexv.upper()
    return h not in ("FFFFFF", "FFFFFE", "FFFF00", "AUTO")


# ---------------------------------------------------------------------------
# Compliance grid -> data model.
# SOURCE OF TRUTH: server.parse_xlsx_compliance (builder/server.py). This is a
# faithful port of its grid->model core so this module needs no server import;
# function name and AX set are kept identical for a mechanical later merge.
# ---------------------------------------------------------------------------
_CN_NAME_RE = re.compile(r"[^A-Za-z0-9_-]+")


def _to_num_or_str(v):
    """Coerce a cell value: numbers stay numeric, blanks -> None, else str."""
    if v is None or v == "":
        return None
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        return v
    s = str(v).strip()
    if s == "":
        return None
    try:
        f = float(s)
        return int(f) if f.is_integer() else f
    except ValueError:
        return s


def _first_str(row, run):
    """First non-empty string within the given run's columns."""
    for ci, _ in run:
        if ci < len(row) and isinstance(row[ci], str) and row[ci].strip():
            return row[ci].strip()
    return None


def _sanitize_name(name):
    name = _CN_NAME_RE.sub("_", name).strip("_")
    return name[:64] if name else ""


def _compliance_model_from_grid(grid, fills=None):
    """Port of server.parse_xlsx_compliance's core. -> {spec_name, sims, rows}.

    ``fills`` (optional per-row shading hex) lets condition/setting rows be told
    apart from result rows. NOTE: a row's ``limit`` (le/ge) and ``sim_span`` are
    NOT encoded anywhere in the rendered Word table, so they cannot be recovered
    on import -- they come back as None/False and must be re-set in the editor."""
    if not grid:
        return {"spec_name": "", "sims": [], "rows": []}
    ncols = max(len(r) for r in grid)
    grid = [list(r) + [None] * (ncols - len(r)) for r in grid]

    header_axis_row = None
    for ri, r in enumerate(grid[:10]):
        if any(isinstance(c, str) and c.strip().upper() in _AX for c in r):
            header_axis_row = ri
            break
    if header_axis_row is None:
        return {"spec_name": "", "sims": [], "rows": []}

    axis_cols = []
    cat_col = item_col = spec_col = unit_col = None
    band = grid[header_axis_row]
    for ci, c in enumerate(band):
        if not isinstance(c, str):
            continue
        cu = c.strip().upper()
        if cu in _AX:
            axis_cols.append(
                (ci, "NTWC" if cu.replace(" ", "").replace("-", "") == "NTWC" else cu))

    for ri in range(0, header_axis_row + 1):
        for ci, c in enumerate(grid[ri]):
            if not isinstance(c, str):
                continue
            cl = c.strip().lower()
            if cl == "category" and cat_col is None:
                cat_col = ci
            elif cl == "item" and item_col is None:
                item_col = ci
            elif cl == "spec" and spec_col is None:
                spec_col = ci
            elif cl == "unit" and unit_col is None:
                unit_col = ci

    spec_name = ""
    sims = []
    spec_axis_cols = None
    groups = []
    if axis_cols:
        runs = []
        cur = [axis_cols[0]]
        seen = {axis_cols[0][1]}
        for prev, nxt in zip(axis_cols, axis_cols[1:]):
            adjacent = (nxt[0] - prev[0] == 1)
            repeats = nxt[1] in seen
            if adjacent and not repeats:
                cur.append(nxt)
                seen.add(nxt[1])
            else:
                runs.append(cur)
                cur = [nxt]
                seen = {nxt[1]}
        runs.append(cur)

        title_row = grid[header_axis_row - 2] if header_axis_row >= 2 else grid[0]
        stage_row = grid[header_axis_row - 1] if header_axis_row >= 1 else grid[0]
        for gi, run in enumerate(runs):
            title = _first_str(title_row, run) or ""
            stage = _first_str(stage_row, run)
            cols = [rc[0] for rc in run]
            if gi == 0 and "spec" in (title or "").lower():
                spec_name = title
                spec_axis_cols = cols
            elif gi == 0 and stage is None and "spec" in (title or "").lower():
                spec_name = title
                spec_axis_cols = cols
            else:
                groups.append((title, stage, cols))
        if spec_axis_cols is None and runs:
            spec_axis_cols = [rc[0] for rc in runs[0]]
            spec_name = _first_str(title_row, runs[0]) or spec_name
            groups = []
            for gi, run in enumerate(runs[1:]):
                title = _first_str(title_row, run) or "Sim%d" % (gi + 1)
                stage = _first_str(stage_row, run)
                groups.append((title, stage, [rc[0] for rc in run]))

    for gi, (title, stage, cols) in enumerate(groups):
        sims.append({
            "key": _sanitize_name(title).lower() or ("sim%d" % (gi + 1)),
            "title": title or ("Sim%d" % (gi + 1)),
            "stage": stage,
        })

    rows = []
    last_cat = ""
    for ri in range(header_axis_row + 1, len(grid)):
        r = grid[ri]
        item = _to_num_or_str(r[item_col]) if item_col is not None else None
        cat = _to_num_or_str(r[cat_col]) if cat_col is not None else None
        if cat:
            last_cat = str(cat)
        if item is None or item == "":
            continue
        unit = _to_num_or_str(r[unit_col]) if unit_col is not None else None
        spec = _to_num_or_str(r[spec_col]) if spec_col is not None else None

        spec_mtm = [None, None, None]
        spec_ntwc = None
        if spec_axis_cols:
            for ai, ci in enumerate(spec_axis_cols[:4]):
                val = _to_num_or_str(r[ci]) if ci < len(r) else None
                if ai < 3:
                    spec_mtm[ai] = val
                else:
                    spec_ntwc = val

        # primary sim -> flat sim_mtm; every extra sim (e.g. a PDR compare column)
        # -> row["sims"][key] so its values are NOT dropped on import.
        sim_mtm = [None, None, None]
        sim_ntwc = None
        per_sim = {}
        for gi, (_gt, _gs, gcols) in enumerate(groups):
            mtm = [None, None, None]
            ntwc = None
            for ai, ci in enumerate(gcols[:4]):
                val = _to_num_or_str(r[ci]) if ci < len(r) else None
                if ai < 3:
                    mtm[ai] = val
                else:
                    ntwc = val
            if gi == 0:
                sim_mtm, sim_ntwc = mtm, ntwc
            elif gi < len(sims):
                per_sim[sims[gi]["key"]] = {"mtm": mtm, "ntwc": ntwc}

        is_setting = (fills is not None and ri < len(fills)
                      and _is_setting_fill(fills[ri]))
        row_obj = {
            "cat": last_cat,
            "item": str(item),
            "unit": "" if unit is None else str(unit),
            "kind": "common_setting" if is_setting else "result",
            "spec": spec,
            "spec_mtm": spec_mtm,
            "sim_mtm": sim_mtm,
            "spec_ntwc": spec_ntwc,
            "sim_ntwc": sim_ntwc,
            "limit": None,
            "sim_span": False,
        }
        if per_sim:
            row_obj["sims"] = per_sim
        rows.append(row_obj)

    return {"spec_name": spec_name, "sims": sims, "rows": rows}


# ---------------------------------------------------------------------------
# Free-table model.
# ---------------------------------------------------------------------------
def _free_table_model(tbl, warn):
    """-> {rows:[[str]], header_rows:int, merges:[{r,c,rs,cs}], col_w:[cm]}.

    Merge-continuation cells are emitted as "" (so a merged value is not
    duplicated); the merge spans are recovered by tracking first-seen tc ids,
    exactly the way inspect_word does. col_w comes from the table grid in cm.
    """
    rows = []
    merges = []
    try:
        n_rows = len(tbl.rows)
        n_cols = len(tbl.columns)
    except Exception:
        return {"rows": [], "header_rows": 1, "merges": [], "col_w": []}

    # Collect every cell's <w:tc> element into a grid FIRST and keep those
    # references alive for the rest of the function. python-docx returns a fresh
    # _Cell wrapper per tbl.cell() call; if a wrapper is GC'd between reads,
    # CPython recycles its id() and a later DISTINCT cell can collide with an
    # earlier one -> it is misread as a merge continuation and BLANKED (the
    # value silently disappears on import). Holding all tc elements at once makes
    # id() a stable per-element identity again.
    tc_grid, txt_grid = [], []
    for r in range(n_rows):
        tc_row, txt_row = [], []
        for c in range(n_cols):
            try:
                cell = tbl.cell(r, c)
                tc_row.append(cell._tc)
                txt_row.append((cell.text or "").strip())
            except Exception:
                tc_row.append(None)
                txt_row.append("")
        tc_grid.append(tc_row)
        txt_grid.append(txt_row)

    # tc id -> top-left (r, c) and accumulated span coverage
    seen = {}
    spans = {}  # tc_id -> {"r":..,"c":..,"rows":set,"cols":set}
    for r in range(n_rows):
        row_out = []
        for c in range(n_cols):
            tc = tc_grid[r][c]
            if tc is None:
                row_out.append("")
                continue
            tc_id = id(tc)
            if tc_id in seen:
                # merge continuation: blank cell, extend the recorded span
                row_out.append("")
            else:
                seen[tc_id] = (r, c)
                spans[tc_id] = {"r": r, "c": c, "rows": set(), "cols": set()}
                row_out.append(txt_grid[r][c])
            spans[tc_id]["rows"].add(r)
            spans[tc_id]["cols"].add(c)
        rows.append(row_out)

    try:
        for tc_id, sp in spans.items():
            rs = len(sp["rows"])
            cs = len(sp["cols"])
            if rs > 1 or cs > 1:
                merges.append({"r": sp["r"], "c": sp["c"], "rs": rs, "cs": cs})
    except Exception:
        merges = []

    try:
        col_w = mc._grid_cols_cm(tbl._tbl)
    except Exception:
        col_w = []
    row_fills = {}
    for ri, hexv in enumerate(_docx_table_row_fills(tbl)):
        if _is_setting_fill(hexv):
            row_fills[str(ri)] = hexv
    return {"rows": rows, "header_rows": 1, "merges": merges,
            "col_w": col_w, "row_fills": row_fills}


def _table_has_images(tbl):
    """True if a table embeds any picture -> it is an engine-rendered imagegrid
    (plain / compliance tables never carry pictures)."""
    try:
        return bool(tbl._tbl.findall(".//" + qn("a:blip")))
    except Exception:
        return False


def _imagegrid_from_table(tbl, document, docx_path, images_dir, img_seq, warn):
    """Reconstruct an imagegrid from a borderless table of pictures: extract each
    cell's image in reading order so nothing is lost. sub_captions=True when a cell
    also carries caption text beneath its picture (the (a)(b) sub-labels)."""
    try:
        n_rows, n_cols = len(tbl.rows), len(tbl.columns)
    except Exception:
        n_rows, n_cols = 0, 0
    items = []
    has_sub = False
    for r in range(n_rows):
        for c in range(n_cols):
            try:
                cell = tbl.cell(r, c)
            except Exception:
                continue
            cell_imgs = []
            for para in cell.paragraphs:
                cell_imgs += _extract_inline_images(
                    para._p, document, docx_path, images_dir, img_seq, warn)
            for ib in cell_imgs:
                items.append({"file": ib["file"]})
            if cell_imgs and (cell.text or "").strip():
                has_sub = True
    return {"type": "imagegrid", "cols": max(1, n_cols), "caption": "",
            "items": items, "sub_captions": has_sub}


# ===========================================================================
# Captions
# ===========================================================================
# Strip a leading auto-prefix the designer typed (the builder re-numbers). The
# CJK pair is included because real fixtures localize captions; it is matched
# literally only as a prefix and is detection-only, never emitted.
_CAPTION_PREFIX_RE = re.compile(
    r"^\s*(?:Table|Figure|\u8868|\u56fe)\s*\d+\s*[:.\-\u3001\uff1a]?\s*",
    re.IGNORECASE,
)
# The engine numbers captions chapter-scoped as "Figure <chap>-<seq>"; the base
# regex strips "Figure <chap>-", leaving the <seq>. This ASCII-only pass drops
# that leftover leading number so the caption text survives clean.
_CAPTION_LEFTOVER_SEQ_RE = re.compile(r"^\s*\d+\s+")


def _strip_caption_prefix(text):
    """Strip a leading Table N / Figure N / localized figure prefix, including
    the engine chapter-scoped Figure <chap>-<seq> numbering."""
    if not text:
        return ""
    t = _CAPTION_PREFIX_RE.sub("", text, count=1)
    t = _CAPTION_LEFTOVER_SEQ_RE.sub("", t, count=1)
    return t.strip()


def _is_caption_para(p, by_id):
    """True if a paragraph carries the Caption style (by id or resolved name)."""
    sid = mc._para_style_id(p)
    if sid is None:
        return False
    if sid == "Caption":
        return True
    st = by_id.get(sid)
    name = mc._style_name(st) if st is not None else sid
    return (name or "").strip().lower() == "caption"


# ===========================================================================
# Images
# ===========================================================================
_IMG_EXT_BY_CT = {
    "image/png": "png", "image/jpeg": "jpg", "image/jpg": "jpg",
    "image/gif": "gif", "image/bmp": "bmp", "image/tiff": "tiff",
    "image/x-emf": "emf", "image/x-wmf": "wmf",
}


def _resolve_image_blob(document, docx_path, rid):
    """-> (blob_bytes, ext) for an embedded image relationship id, or (None, None)."""
    # Preferred: resolve through the document part's related parts.
    try:
        part = document.part.related_parts.get(rid)
        if part is not None:
            blob = part.blob
            ct = getattr(part, "content_type", "") or ""
            ext = _IMG_EXT_BY_CT.get(ct.lower())
            if ext is None:
                pn = str(getattr(part, "partname", ""))
                ext = os.path.splitext(pn)[1].lstrip(".").lower() or "png"
            return blob, ext
    except Exception:
        pass
    # Fallback: resolve the target_ref against the zip (like make_config.extract_logo).
    try:
        target_ref = document.part.rels[rid].target_ref
        with zipfile.ZipFile(docx_path) as z:
            names = z.namelist()
            cand = "word/" + target_ref.lstrip("/")
            chosen = cand if cand in names else None
            if chosen is None:
                base = os.path.basename(target_ref)
                chosen = next((n for n in names if n.endswith("media/" + base)), None)
            if chosen is not None:
                ext = os.path.splitext(chosen)[1].lstrip(".").lower() or "png"
                return z.read(chosen), ext
    except Exception:
        pass
    return None, None


def _extract_inline_images(p, document, docx_path, images_dir, seq_ref, warn):
    """Emit an image block for every embedded picture (a:blip) in a paragraph.

    Writes ``images/img_<n>.<ext>`` under ``images_dir`` and returns a list of
    ``{"type":"image","src","file","caption":"","width_cm"?}`` blocks. The
    designer file embeds real pictures; there is NO ``{{IMG:...}}`` token logic.
    """
    blocks = []
    blips = p.findall(".//" + qn("a:blip"))
    if not blips:
        return blocks
    # one extent per drawing (width); align positionally with blips, best-effort.
    extents = p.findall(".//" + qn("wp:extent"))
    for i, blip in enumerate(blips):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid:
            continue
        blob, ext = _resolve_image_blob(document, docx_path, rid)
        if blob is None:
            warn(f"image relationship '{rid}' could not be resolved")
            continue
        seq_ref[0] += 1
        fname = f"img_{seq_ref[0]}.{ext or 'png'}"
        rel = "images/" + fname
        try:
            if images_dir:
                os.makedirs(images_dir, exist_ok=True)
                with open(os.path.join(images_dir, fname), "wb") as f:
                    f.write(blob)
        except Exception as ex:
            warn(f"image '{fname}' could not be written ({ex!r})")
        block = {"type": "image", "src": rel, "file": rel, "caption": ""}
        if i < len(extents):
            cx = extents[i].get("cx")
            if cx:
                try:
                    block["width_cm"] = mc._emu_to_cm(int(cx))
                except (TypeError, ValueError):
                    pass
        blocks.append(block)
    return blocks


# ===========================================================================
# Body paragraph blocks
# ===========================================================================
def _run_bool(rPr, tag):
    """Read a boolean run property (e.g. ``w:b`` / ``w:i``) honoring its toggle.

    A missing element is False; a present element with no (or a truthy) ``w:val``
    is True; ``val`` of 0 / false / off turns the toggle back off.
    """
    if rPr is None:
        return False
    el = rPr.find(qn(tag))
    if el is None:
        return False
    val = el.get(qn("w:val"))
    return val is None or val not in ("0", "false", "off")


def _run_color(rPr):
    """Direct run color as an uppercase 6-hex string, or None (auto / unset)."""
    if rPr is None:
        return None
    c = rPr.find(qn("w:color"))
    if c is None:
        return None
    val = c.get(qn("w:val"))
    if not val or val == "auto":
        return None
    return val.upper()


# Office default hyperlink color, applied to link runs that carry no direct color
# so imported links stay visibly blue. A neutral constant, never document data.
_HYPERLINK_COLOR = "0563C1"


def _run_text(r):
    """A run's visible text, mapping tabs / line breaks to whitespace."""
    parts = []
    for ch in r:
        tag = ch.tag
        if tag == qn("w:t"):
            parts.append(ch.text or "")
        elif tag == qn("w:tab"):
            parts.append("\t")
        elif tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
    return "".join(parts)


def _emit_run(r, runs, default_color=None):
    """Append ``r`` to ``runs`` as a canonical ``{t, b?, i?, color?}`` dict."""
    t = _run_text(r)
    if not t:
        return
    rPr = r.find(qn("w:rPr"))
    run = {"t": t}
    if _run_bool(rPr, "w:b"):
        run["b"] = True
    if _run_bool(rPr, "w:i"):
        run["i"] = True
    color = _run_color(rPr) or default_color
    if color:
        run["color"] = color
    runs.append(run)


def _para_runs(p):
    """Capture a paragraph as canonical rich runs: a list of ``{t, b, i, color}``.

    Walks runs in document order (including those inside hyperlinks) so inline
    bold / italic / color survive into the seed. This is the schema both the engine
    (``_render_para``) and the editor (``runsToHtml``) consume; a bare ``text`` is
    understood by neither, so it must never be emitted on its own.
    """
    runs = []
    for child in p:
        tag = child.tag
        if tag == qn("w:r"):
            _emit_run(child, runs)
        elif tag == qn("w:hyperlink"):
            for r in child.findall(qn("w:r")):
                _emit_run(r, runs, default_color=_HYPERLINK_COLOR)
    if runs:
        runs[0]["t"] = runs[0]["t"].lstrip()
        runs[-1]["t"] = runs[-1]["t"].rstrip()
        runs = [r for r in runs if r["t"]]
    if not runs:
        # Safety net: never drop body text a structured walk might have missed
        # (e.g. runs nested in revision marks) -- keep it as one plain run.
        txt = mc._para_text(p).strip()
        if txt:
            runs = [{"t": txt}]
    return runs


def _runs_text(runs):
    """Plain concatenated text of a runs list (for placeholder detection)."""
    return "".join(r.get("t", "") for r in (runs or []))


def _para_list_kind(p, document):
    """-> 'bullet' | 'number' | None for a paragraph, resolving its w:numPr through
    the numbering part (bullet abstractNum uses numFmt='bullet', decimal otherwise).
    Mirrors engine.build_list_numbering so a list survives the round-trip."""
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        return None
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return None
    nid_el = numpr.find(qn("w:numId"))
    if nid_el is None or nid_el.get(qn("w:val")) in (None, "0"):
        return None
    numid = nid_el.get(qn("w:val"))
    ilvl_el = numpr.find(qn("w:ilvl"))
    try:
        ilvl = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
    except (TypeError, ValueError):
        ilvl = 0
    try:
        numbering = document.part.numbering_part.element
    except Exception:
        return "bullet"  # numbered somehow but no numbering part -> treat as a list
    abs_id = None
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(numid):
            ael = num.find(qn("w:abstractNumId"))
            abs_id = ael.get(qn("w:val")) if ael is not None else None
            break
    if abs_id is None:
        return "bullet"
    for anum in numbering.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) != str(abs_id):
            continue
        lvl = None
        for lv in anum.findall(qn("w:lvl")):
            if lv.get(qn("w:ilvl")) == str(ilvl):
                lvl = lv
                break
        if lvl is None:
            lvl = anum.find(qn("w:lvl"))
        if lvl is not None:
            nf = lvl.find(qn("w:numFmt"))
            if nf is not None and nf.get(qn("w:val")) == "bullet":
                return "bullet"
        return "number"
    return "bullet"


def _paragraph_block(p, by_id, body_id, mybody_id, warn, document=None):
    """-> a ``para`` block for a non-empty body paragraph, or None to skip."""
    runs = _para_runs(p)
    if not _runs_text(runs).strip():
        return None
    sid = mc._para_style_id(p)
    alias = mc._style_alias(by_id, sid, body_id, mybody_id, warn)
    block = {"type": "para", "runs": runs, "style": alias}
    if document is not None:
        lk = _para_list_kind(p, document)
        if lk:
            block["list"] = lk
    return block


def _table_block(tbl, warn):
    """Classify a body table -> a datatable or free-table block (no caption yet)."""
    kind = classify_table(tbl)
    if kind == "datatable":
        grid = _docx_table_to_grid(tbl)
        fills = _docx_table_row_fills(tbl)
        model = _compliance_model_from_grid(grid, fills)
        return {"type": "datatable", "caption": "", "data": model}
    ft = _free_table_model(tbl, warn)
    block = {
        "type": "table", "caption": "",
        "rows": ft["rows"], "header_rows": ft["header_rows"],
        "merges": ft["merges"], "col_w": ft["col_w"],
    }
    if ft.get("row_fills"):
        block["row_fills"] = ft["row_fills"]
    return block


def build_outline(document, by_id, body_id, mybody_id, body_children,
                  images_dir, docx_path, warn):
    """Walk the body layer into a nested outline tree.

    Each node is ``{title, level, blocks[], children[]}``. Headings open nodes
    (nested by level); non-heading paragraphs, tables and images become blocks on
    the current node. A pending Caption-styled paragraph is attached to the next
    image / table block (captions sit just before their figure in these fixtures).
    """
    root = []
    stack = []  # list of (level, node)
    current = None
    tbl_to_obj = {t._tbl: t for t in document.tables}
    img_seq = [0]
    pending_caption = None  # text waiting to attach to the next image/table

    def _attach_caption(block):
        nonlocal pending_caption
        if pending_caption is not None and not block.get("caption"):
            block["caption"] = pending_caption
        pending_caption = None

    for child in body_children:
        tag = child.tag
        if tag == qn("w:p"):
            lv = mc._heading_level(by_id, child)
            if lv is not None:
                title = mc._strip_autonumber(mc._para_text(child)).strip()
                node = {"title": title, "level": lv, "blocks": [], "children": []}
                while stack and stack[-1][0] >= lv:
                    stack.pop()
                if stack:
                    stack[-1][1]["children"].append(node)
                else:
                    root.append(node)
                stack.append((lv, node))
                current = node
                pending_caption = None
                continue
            # non-heading paragraph
            if _is_caption_para(child, by_id):
                cap = _strip_caption_prefix(mc._para_text(child))
                # The engine emits a caption AFTER its figure; company masters put it
                # BEFORE. Attach backward to an immediately-preceding captionless
                # figure/table, else hold it to attach forward to the next one.
                prev = current["blocks"][-1] if (current and current["blocks"]) else None
                if prev is not None and prev.get("type") in \
                        ("image", "imagegrid", "table", "datatable") \
                        and not prev.get("caption"):
                    prev["caption"] = cap
                else:
                    pending_caption = cap
                continue
            # images first (a paragraph may carry both a picture and stray text)
            imgs = _extract_inline_images(
                child, document, docx_path, images_dir, img_seq, warn)
            if imgs:
                if current is not None:
                    for ib in imgs:
                        _attach_caption(ib)
                        current["blocks"].append(ib)
                continue
            blk = _paragraph_block(child, by_id, body_id, mybody_id, warn, document)
            if blk is not None and current is not None:
                current["blocks"].append(blk)
        elif tag == qn("w:tbl"):
            tbl = tbl_to_obj.get(child)
            if tbl is None or current is None:
                continue
            try:
                if _table_has_images(tbl):
                    # engine renders an imagegrid as a borderless table of pictures;
                    # plain / compliance tables never embed images, so a picture in a
                    # cell is a reliable imagegrid signal.
                    block = _imagegrid_from_table(
                        tbl, document, docx_path, images_dir, img_seq, warn)
                else:
                    block = _table_block(tbl, warn)
            except Exception as ex:
                warn(f"table classification failed ({ex!r}); emitted empty table")
                block = {"type": "table", "caption": "", "rows": [],
                         "header_rows": 1, "merges": [], "col_w": []}
            _attach_caption(block)
            current["blocks"].append(block)
    return root


# ===========================================================================
# PROJECT seed
# ===========================================================================
def parse_docx_report(path, images_dir=None, warn=None):
    """Read a .docx into a PROJECT seed: {"meta", "outline"}. Never raises."""
    warnings = []
    warn = warn or warnings.append
    meta = {
        "title": "", "secrecy": "", "doc_no": "", "page_count": "",
        "author": "", "reviewers": [], "approver": "", "revisions": [],
    }
    outline = []
    try:
        document = Document(path)
        by_id = mc._styles_by_id(document)
    except Exception as ex:
        warn(f"document could not be opened ({ex!r})")
        meta["_warnings"] = warnings or None
        return {"meta": meta, "outline": outline}

    try:
        meta = extract_meta(document, by_id, warn)
    except Exception as ex:
        warn(f"meta extraction failed ({ex!r})")

    try:
        theme = mc._theme_fonts(document)
        body_id, mybody_id = mc.find_custom_body_styles(document, by_id, theme, warn)
    except Exception as ex:
        warn(f"body style detection failed ({ex!r})")
        body_id = mybody_id = None

    try:
        _front, body_children, _ = _split_layers(document, by_id)
        outline = build_outline(document, by_id, body_id, mybody_id,
                                body_children, images_dir, path, warn)
    except Exception as ex:
        warn(f"outline build failed ({ex!r})")

    if warnings:
        meta["_warnings"] = warnings
    return {"meta": meta, "outline": outline}


# ===========================================================================
# Fixed / fillable marking (template skeleton)
# ===========================================================================
# Placeholder marker the fillable bodies use ("to fill", full-width parens). There
# is no structural signal distinguishing a fixed subsection from a fillable one --
# both sit at the same outline level -- so the body text is the only discriminator.
# Detection-only; never emitted.
_PLACEHOLDER_LITERALS = ("\uff08\u5f85\u586b\uff09",)  # "(to fill)"
_TOKEN_RE = re.compile(r"^\s*\{\{.*\}\}.*$")  # a pure {{...}} placeholder line


def _is_placeholder_text(text):
    """True if a paragraph's text is empty / a known placeholder / a pure token."""
    s = (text or "").strip()
    if not s:
        return True
    if s in _PLACEHOLDER_LITERALS:
        return True
    if _TOKEN_RE.match(s):
        return True
    return False


def _node_has_real_prose(node):
    """True if a node is repeating boilerplate to inherit verbatim.

    Only prose paragraphs (and link lists, which are paragraphs) mark a node fixed.
    A node that carries chip-specific DATA -- a datatable, free table or image -- is
    a fillable data section even if it has a one-line prose intro above the data, so
    the presence of any such data block disqualifies it from being fixed. This keeps
    the compliance / results chapters fillable rather than flipping them to fixed on
    an introductory sentence.
    """
    has_prose = False
    for b in node.get("blocks", []):
        bt = b.get("type")
        if bt in ("datatable", "table", "image"):
            return False
        if bt == "para" and not _is_placeholder_text(_runs_text(b.get("runs"))):
            has_prose = True
    return has_prose


def _mark_fixed(node):
    """Post-order: set node['fixed']; fixed nodes keep real blocks, fillable nodes
    get a single placeholder block. Returns the (chapter,sub) fixed-pair recorder
    list via the caller-supplied accumulation in _strip_to_skeleton."""
    fixed = _node_has_real_prose(node)
    node["fixed"] = fixed
    if not fixed:
        # a template must not carry a specific chip's data: blank to a placeholder.
        node["blocks"] = [{"type": "para", "runs": [], "placeholder": True}]
    for child in node.get("children", []):
        _mark_fixed(child)
    return node


def _collect_fixed_pairs(skeleton):
    """List the (chapter, sub) index pairs whose node came out fixed:true.

    Chapter = 1-based level-1 index; sub = 1-based level-2 index within it (a fixed
    level-1 node itself reports sub=0). Mirrors make_config's node_index keying so
    the result can be compared to ``fixed_body_sections``.
    """
    pairs = []
    for ci, chap in enumerate(skeleton, start=1):
        if chap.get("fixed"):
            pairs.append([ci, 0])
        for si, sub in enumerate(chap.get("children", []), start=1):
            if sub.get("fixed"):
                pairs.append([ci, si])
            # deeper levels are rare in these fixtures; record level-2 only.
    return pairs


def _strip_to_skeleton(outline, warn):
    """Mark fixed/fillable across the outline tree; warn if the resulting fixed set
    diverges from the conventional [[1,1],[1,3]] so drift surfaces."""
    for node in outline:
        _mark_fixed(node)
    pairs = _collect_fixed_pairs(outline)
    conventional = [[1, 1], [1, 3]]
    if sorted(pairs) != sorted(conventional):
        warn(f"fixed node set {pairs} differs from the conventional "
             f"{conventional}; reconcile with fixed_body_sections")
    return outline


# ===========================================================================
# TEMPLATE seed
# ===========================================================================
def derive_template(path, images_dir=None, logo_dir=None, warn=None):
    """Read a .docx into a TEMPLATE seed: {"config", "skeleton"}. Never raises."""
    warnings = []
    warn = warn or warnings.append
    config = {}
    skeleton = []

    # ---- config: reuse make_config end to end ----
    logo_filename = None
    try:
        dest_dir = logo_dir or images_dir or tempfile.mkdtemp(prefix="docx_import_logo_")
        logo_path = os.path.join(dest_dir, "logo.png")
        logo_filename = mc.extract_logo(path, Document(path), logo_path, warn)
    except Exception as ex:
        warn(f"logo extraction failed ({ex!r})")
    try:
        extracted = mc.build_extracted(path, logo_filename, overrides={}, warn=warn)
        config = mc._strip_private(mc.deep_merge(mc.neutral_defaults(), extracted))
    except Exception as ex:
        warn(f"config extraction failed ({ex!r})")
        config = mc._strip_private(mc.neutral_defaults())

    # ---- skeleton: our own outline tree + fixed marking ----
    try:
        document = Document(path)
        by_id = mc._styles_by_id(document)
        theme = mc._theme_fonts(document)
        body_id, mybody_id = mc.find_custom_body_styles(document, by_id, theme, warn)
        _front, body_children, _ = _split_layers(document, by_id)
        outline = build_outline(document, by_id, body_id, mybody_id,
                                body_children, images_dir, path, warn)
        skeleton = _strip_to_skeleton(outline, warn)
    except Exception as ex:
        warn(f"skeleton build failed ({ex!r})")

    result = {"config": config, "skeleton": skeleton}
    if warnings:
        result["_warnings"] = warnings
    return result


# ===========================================================================
# CLI test harness
# ===========================================================================
def main(argv=None):
    ap = argparse.ArgumentParser(
        description="Parse a .docx into a project seed or derive a template seed.")
    ap.add_argument("docx", help="the .docx to read")
    ap.add_argument("--mode", choices=("report", "template"), default="report",
                    help="report -> {meta, outline}; template -> {config, skeleton}")
    ap.add_argument("--images-dir", default=None,
                    help="directory for extracted images (default: a temp dir)")
    args = ap.parse_args(argv)

    if not os.path.isfile(args.docx):
        print(f"error: file not found: {args.docx}", file=sys.stderr)
        return 2

    images_dir = args.images_dir or tempfile.mkdtemp(prefix="docx_import_")
    warnings = []
    warn = warnings.append

    try:
        if args.mode == "template":
            result = derive_template(args.docx, images_dir=images_dir, warn=warn)
        else:
            result = parse_docx_report(args.docx, images_dir=images_dir, warn=warn)
    except Exception as ex:
        # Never traceback-crash on malformed input: emit a partial result + warning.
        print(f"error: parse failed: {ex!r}", file=sys.stderr)
        print(json.dumps({"meta": {}, "outline": [], "_error": repr(ex)},
                         ensure_ascii=False, indent=2))
        return 1

    print(json.dumps(result, ensure_ascii=False, indent=2))
    if warnings:
        print("\nwarnings:", file=sys.stderr)
        for w in warnings:
            print("  -", w, file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
