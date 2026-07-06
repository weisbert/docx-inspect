#!/usr/bin/env python3
"""Golden render regression test -- locks the document "moat" invariants.

This test renders a small, self-contained, ASCII-only project through the real
engine (engine.render_report) and then inspects the produced .docx (via
python-docx + the underlying XML) to assert the guarantees that make this tool
worth using -- the ones a future refactor must never silently break:

  1. Auto red-flagging: a simulation value that is out of spec is rendered in the
     flag color (red) AND bold (the B&W-safe marker); an in-spec value is NOT.
  2. NTWC corner flagging: an out-of-spec NTWC value (axis index 3) is flagged.
  3. Word-native numbering: figure/table captions contain live SEQ + STYLEREF
     fields (not frozen literal text) with a bookmark around the number so
     cross-references can target it.
  4. Compliance header band: the 3 header rows carry w:tblHeader so they repeat
     on a page break.
  5. Footer + TOC fields: DATE / PAGE / NUMPAGES fields exist in the footer and a
     TOC field exists.
  6. Warnings manifest: render_report reports the missing image and the
     out-of-spec rows it found (flagged_rows via the datatable result is exercised
     indirectly through the rendered red cells; the manifest is asserted for the
     missing image).
  7. No-spill guarantee: every compliance row uses an EXACT row-height rule.

The fixture config and project contain ZERO CJK characters and zero company
terms (iron rule 1). The test is standalone: run it directly, exit 0 on pass and
nonzero on any failed assertion.

Run:
    python builder/test_render_golden.py
"""

import os
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))
if HERE not in sys.path:
    sys.path.insert(0, HERE)

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

import engine  # noqa: E402
import tables  # noqa: E402


# ---------------------------------------------------------------------------
# Fixture: a complete, fully renderable, ASCII-only template config.
# Built inline (NOT copied from any local config) so it stays company-clean.
# ---------------------------------------------------------------------------
def golden_config():
    return {
        "id": "golden_tpl_v1",
        "caption_prefix": {"image": "Figure", "table": "Table"},
        "toc": {
            "title": "Contents",
            "field": 'TOC \\o "1-3" \\h \\z \\u',
            "placeholder": "(update field)",
            "size_pt": 20,
        },
        "logo": "",  # no logo file -> engine records a missing_logo warning
        "styles": {
            "page": {
                "w_cm": 21.0, "h_cm": 29.7, "margin_cm": 2.5,
                "header_dist_cm": 1.2, "footer_dist_cm": 1.2,
                "different_first_page": True,
            },
            "normal": {"ascii": "Arial", "eastAsia": "Arial", "size_pt": 10.5},
            "headings": {
                "levels": {
                    "1": {"ascii": "Arial", "size_pt": 16, "bold": True},
                    "2": {"ascii": "Arial", "size_pt": 14, "bold": True},
                    "3": {"ascii": "Arial", "size_pt": 12, "bold": True},
                    "default": {"ascii": "Arial", "size_pt": 12, "bold": False},
                },
                "space_before_pt": {"1": 12, "2": 10, "3": 8, "4": 6, "5": 6},
                "h1_after_pt": 24,
                "h1_bottom_border": {"val": "single", "sz": 6, "color": "auto"},
                "autonumber": {"num_id": 88, "suffix": "space", "ascii": "Arial"},
            },
            "caption": {"ascii": "Arial", "size_pt": 9, "bold": True, "align": "center"},
            "body": {"name": "ReportBody", "base": "Normal", "size_pt": 10.5,
                     "left_cm": 0.0, "first_line_cm": 0.74},
            "mybody": {"name": "ReportBodyIndent", "base": "ReportBody",
                       "ascii": "Arial", "left_cm": 0.0, "first_line_cm": 0.74},
            "header_table": {
                "cols_cm": [1.5, 12.0, 2.5],
                "row_h_twips": 751,
                "cell_bottom_border": {"val": "single", "sz": 6, "color": "auto"},
                "logo_cm": 1.13,
                "title_font": {"ascii": "Arial", "eastAsia": "Arial", "size_pt": 9},
                "title_placeholder": "Report Title",
                "secrecy_label": "Internal",
            },
            "footer_table": {
                "cols_cm": [5.0, 6.0, 5.0],
                "top_border": {"val": "single", "sz": 4},
                "date_format": "yyyy-MM-dd",
                "center_text": "",
                "page_text": ["", " / ", ""],
                "font": {"ascii": "Arial", "size_pt": 9},
            },
            "colors": {"red": "FF0000", "secrecy": "4F81BD"},
        },
        "cover": {
            "company_line": "ACME",
            "secrecy_default": "Internal",
            "page_count_field": True,
            "page_text": ["", " pages"],
            "company_names": [{"text": "ACME Corp", "ascii": "Arial", "size_pt": 16}],
            "logo_cm": 2.6,
            "big_title": {"placeholder": "Report Title", "subtitle": "", "size_pt": 24},
            "fields": [
                {"key": "title", "label": "Title", "table": "info", "required": True},
                {"key": "author", "label": "Author", "table": "signature", "required": False},
            ],
            "tables": {
                "info": {
                    "cols_cm": [3.0, 3.0, 4.0, 3.0, 3.0],
                    "outer": "double", "inner": "single", "sz": 14,
                    "labels": {"title": "Project", "doc_no": "Code",
                               "secrecy": "Secrecy", "pages": "Pages"},
                },
                "signature": {
                    "cols_cm": [3.0, 3.0, 1.0, 3.0, 3.0],
                    "rows": [["Author", ""], ["Reviewer", ""], ["Approver", ""]],
                    "sign_underline": True, "sign_cols": [1, 4],
                },
                "revision": {
                    "cols_cm": [4.0, 3.0, 6.0, 3.0],
                    "headers": ["Date", "Version", "Note", "Author"],
                    "header_font": {"ascii": "Arial", "size_pt": 10.5},
                    "border": "single",
                    "title": {"text": "Revision History", "ascii": "Arial",
                              "eastAsia": "Arial", "size_pt": 16},
                },
            },
        },
        "compliance": {
            "col_w_cm": {"cat": 2.0, "item": 3.0, "spec": 1.5, "axis": 1.4,
                         "spacer": 0.2, "unit": 1.2},
            "font_pt": 7,
            "row_h_pt": {"header": 12, "data": 10},
            "axis_labels": ["MIN", "TYP", "MAX", "NTWC"],
            "fills": {"header": "FFF2CC", "setting": "DDEBF7",
                      "result": "FFFFFF", "separator": "BFBFBF"},
            "setting_kinds": ["common_setting", "module_setting", "tb"],
            "default_limit": {"le": "<= upper", "ge": ">= target", "range": "within"},
            "flag_color": "FF0000",
            "borders": {"val": "single", "sz": 4, "color": "000000"},
        },
        "free_table": {
            "header_fill": "D9D9D9",
            "border": {"val": "single", "sz": 4, "color": "000000"},
        },
        "fixed_bodies": {},
        "ui_strings": {},
    }


# Sentinel strings planted in `notes` fields at every level; the render must
# never emit any of them (notes are a side channel, not document content).
NOTE_SENTINEL_REPORT = "ZZNOTESENTINELREPORTZZ"
NOTE_SENTINEL_SECTION = "ZZNOTESENTINELSECTIONZZ"
NOTE_SENTINEL_BLOCK = "ZZNOTESENTINELBLOCKZZ"


def golden_project():
    """One chapter with a (missing) captioned image and a compliance datatable.

    Datatable rows (limit 'le' = simulated value must be <= the spec MAX):
      * I_total : sim MAX 758 > spec MAX 500          -> MAX flagged (red+bold)
      * P_static: sim NTWC 1200 > spec NTWC 1000      -> NTWC flagged (red+bold)
      * P_total : all sims <= spec MAX 2000           -> nothing flagged
    """
    data = {
        "spec_name": "Spec",
        "sims": [{"key": "pilot", "title": "Pilot", "stage": "Pre"}],
        "rows": [
            {
                "cat": "Supply", "item": "I_total", "unit": "uA", "kind": "result",
                "spec": 500, "spec_mtm": [None, 500, None],
                "sim_mtm": [279, 490, 758], "spec_ntwc": None, "sim_ntwc": None,
                "limit": "le", "sim_span": False,
            },
            {
                "cat": "Power", "item": "P_static", "unit": "mW", "kind": "result",
                "spec": 1000, "spec_mtm": [None, 1000, None],
                "sim_mtm": [800, 950, 980], "spec_ntwc": 1000, "sim_ntwc": 1200,
                "limit": "le", "sim_span": False,
            },
            {
                "cat": "Power", "item": "P_total", "unit": "mW", "kind": "result",
                "spec": 2000, "spec_mtm": [None, 2000, None],
                "sim_mtm": [500, 1200, 1500], "spec_ntwc": None, "sim_ntwc": None,
                "limit": "le", "sim_span": False,
            },
        ],
    }
    return {
        "schema_version": 1,
        "template": "golden_tpl_v1",
        # meta carries a report-level note (the revision-remark channel). The
        # engine must NEVER render it into the Word body (asserted below).
        "meta": {"title": "Golden Render Test", "author": "Tester",
                 "reviewers": [], "revisions": [],
                 "notes": [{"by": "user", "at": "2026-07-03", "status": "open",
                            "text": NOTE_SENTINEL_REPORT}]},
        "outline": [
            {
                "title": "Results",
                "level": 1,
                # section-level note -- also must not reach the body.
                "notes": [{"by": "claude", "at": "2026-07-03", "status": "done",
                           "text": NOTE_SENTINEL_SECTION}],
                "blocks": [
                    {"type": "para", "list": None,
                     # block-level note on a real block -- must not reach the body.
                     "notes": [{"by": "user", "at": "2026-07-03", "status": "open",
                                "text": NOTE_SENTINEL_BLOCK}],
                     "runs": [{"t": "See "}, {"ref": "img-gold-1"},
                              {"t": " and "}, {"ref": "dt-gold-1"}, {"t": "."}]},
                    {"type": "image", "id": "img-gold-1",
                     "file": "images/missing.png", "caption": "A missing figure",
                     "width_cm": 12.0, "size": "full"},
                    {"type": "datatable", "id": "dt-gold-1", "kind": "compliance",
                     "caption": "Compliance results", "data": data},
                    # two (missing) side-by-side images with (a)(b) sub-captions:
                    # cells must be bottom-aligned so unequal heights keep the
                    # sub-labels on one baseline. rows=2 with 2 images exercises
                    # the layout picker's min-rows padding: the grid must render 2
                    # rows (top filled, bottom row blank), not just 1.
                    {"type": "imagegrid", "id": "grid-gold-1", "cols": 2, "rows": 2,
                     "caption": "Side-by-side comparison", "sub_captions": True,
                     "width_cm": 15.5,
                     "items": [{"file": "images/left.png"},
                               {"file": "images/right.png"}]},
                ],
                "children": [],
            },
            {
                "title": "Outline lists",
                "level": 1,
                "blocks": [
                    # nested bullets: solid circle (lvl0) with hollow-square sub (lvl1)
                    {"type": "para", "list": "bullet", "level": 0,
                     "runs": [{"t": "PDR"}]},
                    {"type": "para", "list": "bullet", "level": 1,
                     "runs": [{"t": "sub point one"}]},
                    {"type": "para", "list": "bullet", "level": 1,
                     "runs": [{"t": "sub point two"}]},
                    # numbered list whose items are separated by a plain paragraph:
                    # numbering must CONTINUE (restart only per section), so these
                    # two number items share one numId.
                    {"type": "para", "list": "number", "level": 0,
                     "runs": [{"t": "first label"}]},
                    {"type": "para", "list": None,
                     "runs": [{"t": "an intervening address line"}]},
                    {"type": "para", "list": "number", "level": 0,
                     "runs": [{"t": "second label"}]},
                    # free table with styled cells (rich runs: bold / red / italic)
                    {"type": "table", "caption": "Styled cells", "header_rows": 1,
                     "col_w": None,
                     "rows": [
                         ["Style", "Example"],
                         ["bold+red", {"runs": [{"t": "Bold", "b": True},
                                                {"t": " red", "color": "FF0000"}]}],
                         ["italic", {"runs": [{"t": "slanted", "i": True}]}],
                     ]},
                ],
                "children": [],
            },
        ],
    }


# ---------------------------------------------------------------------------
# Assertion harness.
# ---------------------------------------------------------------------------
PASS = 0
FAIL = 0


def check(cond, name, detail=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print("  [PASS] %s" % name)
    else:
        FAIL += 1
        print("  [FAIL] %s%s" % (name, ("  -> " + detail) if detail else ""))


# ---------------------------------------------------------------------------
# XML inspection helpers (public python-docx objects + qn() for qualified names).
# ---------------------------------------------------------------------------
def cell_runs(cell):
    runs = []
    for p in cell.paragraphs:
        runs.extend(p.runs)
    return runs


def run_is_red(run):
    c = run.font.color
    try:
        return c is not None and c.rgb is not None and str(c.rgb) == "FF0000"
    except Exception:
        return False


def run_is_bold(run):
    # Read the explicit <w:b/> too, since font.bold may be None when set via XML.
    if run.font.bold:
        return True
    rpr = run._r.find(qn("w:rPr"))
    return rpr is not None and rpr.find(qn("w:b")) is not None


def run_is_italic(run):
    if run.font.italic:
        return True
    rpr = run._r.find(qn("w:rPr"))
    return rpr is not None and rpr.find(qn("w:i")) is not None


def find_free_table(doc):
    """The free table is the one carrying the styled 'Bold' cell (compliance/grid
    tables don't contain that text)."""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "Bold" in cell_text(cell):
                    return t
    return None


def cell_text(cell):
    return "".join(r.text for r in cell_runs(cell)).strip()


def para_by_text(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def para_numpr(p):
    """Return (ilvl, numId) ints for a paragraph's numbering, or (None, None)."""
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return (None, None)
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return (None, None)
    il = numPr.find(qn("w:ilvl"))
    nid = numPr.find(qn("w:numId"))
    iv = int(il.get(qn("w:val"))) if il is not None else None
    nv = int(nid.get(qn("w:val"))) if nid is not None else None
    return (iv, nv)


def abstract_lvl_text(doc, abstract_id, ilvl):
    """lvlText string for a given abstractNum id + level, or None."""
    numbering = doc.part.numbering_part.element
    for a in numbering.findall(qn("w:abstractNum")):
        if a.get(qn("w:abstractNumId")) == str(abstract_id):
            for lvl in a.findall(qn("w:lvl")):
                if lvl.get(qn("w:ilvl")) == str(ilvl):
                    lt = lvl.find(qn("w:lvlText"))
                    return lt.get(qn("w:val")) if lt is not None else None
    return None


def num_to_abstract(doc, num_id):
    """Map a numId to its abstractNumId (int), or None."""
    numbering = doc.part.numbering_part.element
    for n in numbering.findall(qn("w:num")):
        if n.get(qn("w:numId")) == str(num_id):
            aid = n.find(qn("w:abstractNumId"))
            return int(aid.get(qn("w:val"))) if aid is not None else None
    return None


def find_imagegrid_table(doc):
    """The image grid is the borderless table that carries the "(a)" sub-caption
    (cover/header tables also use bottom vAlign, so match on sub-caption text)."""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "(a)" in cell_text(cell):
                    return t
    return None


def find_compliance_table(doc):
    """The compliance table is the only one with FFF2CC-shaded header cells and a
    3-row header band carrying w:tblHeader; pick the table whose first row cells
    carry tblHeader."""
    for t in doc.tables:
        if not t.rows:
            continue
        trPr = t.rows[0]._tr.find(qn("w:trPr"))
        if trPr is not None and trPr.find(qn("w:tblHeader")) is not None:
            return t
    return None


def paragraph_has_field(p, instr_substr):
    """True when the paragraph contains a w:instrText run holding instr_substr."""
    for it in p._p.iter(qn("w:instrText")):
        if it.text and instr_substr in it.text:
            return True
    return False


def doc_has_field_anywhere(doc, instr_substr):
    """Search the whole document part (body) for a field instruction substring."""
    for it in doc.element.body.iter(qn("w:instrText")):
        if it.text and instr_substr in it.text:
            return True
    return False


def footer_has_field(doc, instr_substr):
    sec = doc.sections[0]
    for it in sec.footer._element.iter(qn("w:instrText")):
        if it.text and instr_substr in it.text:
            return True
    return False


def caption_paragraphs(doc):
    out = []
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Caption":
            out.append(p)
    return out


def bookmark_names(doc):
    names = set()
    for bs in doc.element.body.iter(qn("w:bookmarkStart")):
        nm = bs.get(qn("w:name"))
        if nm:
            names.add(nm)
    return names


# ---------------------------------------------------------------------------
# Locate sim-axis cells by reusing the engine's own column plan (robust to a
# layout change: we ask tables.py where the sim axes live rather than hardcoding).
# ---------------------------------------------------------------------------
def sim_axis_columns(data, cfg):
    """Return {axis_index: column_index} for the (single) sim group."""
    groups = tables.make_groups(data, cfg)
    show_spec_col = not any(g["role"] == "spec" for g in groups)
    plan = tables._plan_columns(groups, show_spec_col, cfg["col_w_cm"])
    sim_key = None
    for g in groups:
        if g["role"] == "sim":
            sim_key = g["key"]
            break
    cols = {}
    for idx, p in enumerate(plan):
        if p["kind"] == "axis" and p.get("group") == sim_key:
            cols[p["axis"]] = idx
    return cols


def main():
    cfg = golden_config()
    cfg["_logo_path"] = ""  # no logo
    project = golden_project()

    tmp = tempfile.mkdtemp(prefix="golden_render_")
    os.makedirs(os.path.join(tmp, "images"), exist_ok=True)  # intentionally empty
    out_path = os.path.join(tmp, "out", "golden.docx")

    result = engine.render_report(project, cfg, tmp, out_path)

    # --- result manifest shape ---
    check(isinstance(result, dict), "render_report returns a dict manifest")
    rp = engine._result_out_path(result)
    check(rp and os.path.isfile(rp), "output .docx written", str(rp))

    warnings = result.get("warnings", []) if isinstance(result, dict) else []
    stats = result.get("stats", {}) if isinstance(result, dict) else {}
    wtypes = [w.get("type") for w in warnings]
    check("missing_image" in wtypes, "manifest reports missing image",
          "warnings=%r" % wtypes)
    check(stats.get("missing_images", 0) >= 1, "stats counts the missing image",
          "stats=%r" % stats)

    doc = Document(rp)

    # --- notes channel never leaks into the rendered document ---
    # Read the raw document.xml (+ headers/footers) and assert no sentinel from
    # any meta/section/block `notes` field appears. This locks the iron rule that
    # the revision-remark side channel is stripped from the Word body.
    import zipfile
    xml_blob = ""
    with zipfile.ZipFile(rp) as zf:
        for name in zf.namelist():
            if name.endswith(".xml"):
                xml_blob += zf.read(name).decode("utf-8", "replace")
    for sentinel, where in ((NOTE_SENTINEL_REPORT, "meta"),
                            (NOTE_SENTINEL_SECTION, "section"),
                            (NOTE_SENTINEL_BLOCK, "block")):
        check(sentinel not in xml_blob,
              "notes are NOT rendered (%s-level note absent from docx)" % where,
              "sentinel %r leaked into the document" % sentinel)

    # --- each new top-level chapter starts on a new page ---
    # "Results" is chapter 1: the TOC already broke the page, so NO break. "Outline
    # lists" is chapter 2: page-break-before set on its Heading 1 (Word/PDF native,
    # no blank page). Uses the paragraph property, not an explicit break run.
    def _h1_page_break(title):
        for p in doc.paragraphs:
            if p.style and p.style.name == "Heading 1" and (p.text or "").strip() == title:
                pPr = p._p.find(qn("w:pPr"))
                return pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None
        return None
    check(_h1_page_break("Results") is False,
          "chapter 1 heading has NO page-break-before (TOC already broke)")
    check(_h1_page_break("Outline lists") is True,
          "chapter 2 heading HAS page-break-before")

    # --- compliance table located + header band ---
    ctbl = find_compliance_table(doc)
    check(ctbl is not None, "compliance table present")
    if ctbl is None:
        return _finish()

    header_marked = 0
    for hr in range(3):
        trPr = ctbl.rows[hr]._tr.find(qn("w:trPr"))
        if trPr is not None and trPr.find(qn("w:tblHeader")) is not None:
            header_marked += 1
    check(header_marked == 3, "3 header rows carry w:tblHeader",
          "marked=%d" % header_marked)

    # --- EXACT row heights (no-spill guarantee) ---
    exact = 0
    for row in ctbl.rows:
        trPr = row._tr.find(qn("w:trPr"))
        th = trPr.find(qn("w:trHeight")) if trPr is not None else None
        if th is not None and th.get(qn("w:hRule")) == "exact":
            exact += 1
    check(exact == len(ctbl.rows), "every compliance row uses EXACT height",
          "exact=%d of %d" % (exact, len(ctbl.rows)))

    # --- red+bold flag invariants: locate sim axis cells, data rows start at 3 ---
    data = project["outline"][0]["blocks"][2]["data"]
    sim_cols = sim_axis_columns(data, cfg["compliance"])
    check(set(sim_cols.keys()) == {0, 1, 2, 3}, "sim group has 4 axis columns",
          "sim_cols=%r" % sim_cols)

    # row 0 (I_total) is table row 3; MAX axis index 2 (758) must be red+bold.
    r_itotal = 3
    max_cell = ctbl.cell(r_itotal, sim_cols[2])
    typ_cell = ctbl.cell(r_itotal, sim_cols[1])
    max_runs = cell_runs(max_cell)
    check(cell_text(max_cell) == "758", "I_total MAX cell shows 758",
          "text=%r" % cell_text(max_cell))
    check(any(run_is_red(r) for r in max_runs), "out-of-spec MAX is RED")
    check(any(run_is_red(r) and run_is_bold(r) for r in max_runs),
          "out-of-spec MAX is RED and BOLD (B&W-safe)")
    typ_runs = cell_runs(typ_cell)
    check(not any(run_is_red(r) for r in typ_runs),
          "in-spec TYP (490) is NOT red", "text=%r" % cell_text(typ_cell))

    # row 1 (P_static) is table row 4; NTWC axis index 3 (1200) must be red+bold.
    r_pstatic = 4
    ntwc_cell = ctbl.cell(r_pstatic, sim_cols[3])
    ntwc_runs = cell_runs(ntwc_cell)
    check(cell_text(ntwc_cell) == "1200", "P_static NTWC cell shows 1200",
          "text=%r" % cell_text(ntwc_cell))
    check(any(run_is_red(r) for r in ntwc_runs), "out-of-spec NTWC is RED")
    check(any(run_is_red(r) and run_is_bold(r) for r in ntwc_runs),
          "out-of-spec NTWC is RED and BOLD")

    # row 2 (P_total) is table row 5; MAX (1500) within 2000 -> NOT red.
    r_ptotal = 5
    ptotal_max = ctbl.cell(r_ptotal, sim_cols[2])
    ptotal_runs = cell_runs(ptotal_max)
    check(cell_text(ptotal_max) == "1500", "P_total MAX cell shows 1500",
          "text=%r" % cell_text(ptotal_max))
    check(not any(run_is_red(r) for r in ptotal_runs),
          "in-spec P_total MAX (1500) is NOT red")

    # --- image-grid sub-caption alignment (bottom-aligned cells) ---
    grid_tbl = find_imagegrid_table(doc)
    check(grid_tbl is not None, "image-grid table present (bottom-aligned cells)")
    if grid_tbl is not None:
        # min-rows padding: picker chose rows=2 with 2 images -> 2 rows x 2 cols
        check(len(grid_tbl.rows) == 2,
              "image-grid honors picked rows as a minimum (2 rows, bottom blank)",
              "rows=%d" % len(grid_tbl.rows))
        cells = [c for row in grid_tbl.rows for c in row.cells]
        bottom = 0
        for cell in cells:
            tcPr = cell._tc.find(qn("w:tcPr"))
            va = tcPr.find(qn("w:vAlign")) if tcPr is not None else None
            if va is not None and va.get(qn("w:val")) == "bottom":
                bottom += 1
        check(bottom == len(cells),
              "every image-grid cell is bottom-aligned (sub-captions share a baseline)",
              "bottom=%d of %d" % (bottom, len(cells)))

    # --- nested body lists (bullets ● / □, continuing numbered list) ---
    il0, nid0 = para_numpr(para_by_text(doc, "PDR"))
    il1, nid1 = para_numpr(para_by_text(doc, "sub point one"))
    check(nid0 is not None and il0 == 0, "bullet lvl0 paragraph has numbering at ilvl 0",
          "ilvl=%r numId=%r" % (il0, nid0))
    check(il1 == 1 and nid1 == nid0, "bullet sub-item is ilvl 1 on the same bullet numId",
          "ilvl=%r numId=%r" % (il1, nid1))
    bul_abs = num_to_abstract(doc, nid0) if nid0 is not None else None
    check(abstract_lvl_text(doc, bul_abs, 0) == "●", "bullet level 0 glyph is ● (solid circle)",
          "lvlText=%r" % abstract_lvl_text(doc, bul_abs, 0))
    check(abstract_lvl_text(doc, bul_abs, 1) == "□", "bullet level 1 glyph is □ (hollow square)",
          "lvlText=%r" % abstract_lvl_text(doc, bul_abs, 1))

    ila, nida = para_numpr(para_by_text(doc, "first label"))
    ilb, nidb = para_numpr(para_by_text(doc, "second label"))
    _, nidmid = para_numpr(para_by_text(doc, "an intervening address line"))
    check(nida is not None and nida == nidb,
          "numbered items keep one numId across an intervening plain paragraph",
          "numIds=%r,%r" % (nida, nidb))
    check(nidmid is None, "the intervening plain paragraph has no numbering", "numId=%r" % nidmid)
    dec_abs = num_to_abstract(doc, nida) if nida is not None else None
    check(abstract_lvl_text(doc, dec_abs, 0) == "%1.", "numbered level 0 renders as '1.'",
          "lvlText=%r" % abstract_lvl_text(doc, dec_abs, 0))
    check(nida != nid0, "bullets and numbers use different numbering definitions",
          "bulletNum=%r numberNum=%r" % (nid0, nida))

    # --- free table with styled (rich) cells ---
    ft = find_free_table(doc)
    check(ft is not None, "free table with styled cells present")
    if ft is not None:
        styled = ital = None
        for row in ft.rows:
            for cell in row.cells:
                if "Bold" in cell_text(cell):
                    styled = cell
                if cell_text(cell) == "slanted":
                    ital = cell
        sruns = cell_runs(styled) if styled is not None else []
        check(any(run_is_bold(r) for r in sruns), "styled cell has a bold run")
        check(any(run_is_red(r) and not run_is_bold(r) for r in sruns),
              "styled cell has a red, non-bold run (per-run styling, not header bold)")
        iruns = cell_runs(ital) if ital is not None else []
        check(any(run_is_italic(r) for r in iruns), "styled cell has an italic run")

    # --- body style carries a (config-driven) first-line indent ---
    try:
        body_style = doc.styles["ReportBody"]
        fli = body_style.paragraph_format.first_line_indent
        check(fli is not None and fli > 0,
              "body style has a positive first-line indent",
              "first_line_indent=%r" % (fli,))
    except KeyError:
        check(False, "body style 'ReportBody' exists")

    # --- Word-native caption numbering (SEQ + STYLEREF) + bookmark ---
    caps = caption_paragraphs(doc)
    check(len(caps) >= 2, "at least 2 caption paragraphs (figure + table)",
          "found=%d" % len(caps))
    has_seq = any(paragraph_has_field(p, "SEQ ") for p in caps)
    has_styleref = any(paragraph_has_field(p, "STYLEREF") for p in caps)
    check(has_seq, "a caption contains a SEQ field (live numbering)")
    check(has_styleref, "a caption contains a STYLEREF field (chapter number)")
    has_fig = any(paragraph_has_field(p, "SEQ Figure") for p in caps)
    has_tab = any(paragraph_has_field(p, "SEQ Table") for p in caps)
    check(has_fig, "figure caption uses 'SEQ Figure'")
    check(has_tab, "table caption uses 'SEQ Table'")

    bms = bookmark_names(doc)
    check("bm_img-gold-1_num" in bms, "image caption number is bookmarked",
          "bookmarks=%r" % sorted(bms))
    check("bm_dt-gold-1_num" in bms, "table caption number is bookmarked",
          "bookmarks=%r" % sorted(bms))

    # --- cross-reference REF fields target those bookmarks ---
    check(doc_has_field_anywhere(doc, "REF bm_img-gold-1_num"),
          "paragraph REF field targets the figure bookmark")
    check(doc_has_field_anywhere(doc, "REF bm_dt-gold-1_num"),
          "paragraph REF field targets the table bookmark")

    # --- footer DATE / PAGE / NUMPAGES + TOC field ---
    check(footer_has_field(doc, "DATE"), "footer has a DATE field")
    check(footer_has_field(doc, "PAGE"), "footer has a PAGE field")
    check(footer_has_field(doc, "NUMPAGES"), "footer has a NUMPAGES field")
    check(doc_has_field_anywhere(doc, "TOC"), "document has a TOC field")

    # --- multi-simulation datatable: per-sim row values + hideable Spec group ---
    comp_cfg = golden_config()["compliance"]
    ms_data = {
        "show_spec": False,
        "sims": [
            {"key": "post", "title": "Post", "axes": ["MIN", "TYP", "MAX", "NTWC"]},
            {"key": "pre", "title": "Pre", "axes": ["MIN", "TYP", "MAX", "NTWC"]},
        ],
        "rows": [
            {"cat": "Current", "item": "I_x", "kind": "result", "unit": "uA",
             "limit": None, "spec_mtm": [None, None, None],
             "sims": {"post": {"mtm": [160.6, 173.6, 198.7], "ntwc": 186.7},
                      "pre": {"mtm": [212.3, 227.8, 252.6], "ntwc": None}}},
        ],
    }
    ms_groups = tables.make_groups(ms_data, comp_cfg)
    check([g["key"] for g in ms_groups] == ["post", "pre"],
          "show_spec=False suppresses Spec group; sims kept in order",
          "groups=%r" % [g["key"] for g in ms_groups])
    ms_tbl = tables.render_datatable(Document(), ms_data, comp_cfg)["table"]
    ms_texts = [c.text.strip() for row in ms_tbl.rows for c in row.cells]
    check("160.6" in ms_texts and "212.3" in ms_texts,
          "per-sim values render distinctly (post 160.6 AND pre 212.3)",
          "texts=%r" % ms_texts)
    check(ms_texts.count("186.7") == 1,
          "post NTWC (186.7) shown once; pre NTWC blank (not duplicated across sims)")

    return _finish()


def _finish():
    print("\n=== GOLDEN RENDER SUMMARY ===")
    print("  PASS: %d   FAIL: %d" % (PASS, FAIL))
    print("=============================")
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(main())
