# -*- coding: utf-8 -*-
"""
tables.py -- Generic, config-driven table renderers for the structured document builder.

Two public renderers, both writing native Word tables into an existing python-docx
document (portrait body flow -- never a new page, never landscape, never a page-size change):

  * render_datatable(doc, data, cfg)  -- a data-driven "compliance" table: several value
        groups laid out side by side, each group made of N axis columns; auto red-flagging of
        out-of-limit cells; merged category column; fixed (exact) row heights so a wide table
        is shrunk inline rather than spilling onto a new page or rotating the page.

  * render_free_table(doc, rows, cfg, ...) -- an arbitrary rows/cols table with optional
        header shading, cell merges, and per-column widths.

Nothing domain specific is hardcoded here: column widths, fonts, axis labels (including the
unbreakable narrow-axis token), fill colors, limit directions and the flag color all arrive
through the ``cfg`` dict (the template config's table section) or the data itself.

RETURN-SHAPE CONTRACT: both renderers return a result dict
    {"table": <python-docx Table or None>,
     "total_rows": int,          # data rows rendered
     "flagged_rows": int,        # rows with >=1 out-of-spec sim cell (datatable only)
     "warnings": [ {type, detail, location?}, ... ]}  # e.g. row_clip_risk
The engine unpacks this and merges ``warnings``/counts into its render manifest.
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Low-level cell / table helpers
# ---------------------------------------------------------------------------
def _shade(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)


def _vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)


def _vbottom(cell):
    """Bottom-align a cell's content. Used by the image grid so that images of
    unequal height still leave their (a)(b) sub-captions on one baseline."""
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "bottom")
    tcPr.append(va)


def _table_fixed_layout(table):
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _table_borders(table, val="single", sz=4, color="000000"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def _cell_margins(table, top=0, bottom=0, left=28, right=28):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        m = OxmlElement("w:" + edge)
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tblPr.append(mar)


def _set_cell_text(cell, val, font_pt, ascii_font="Arial", eastasia="SimSun",
                   bold=False, color=None, align="center"):
    """Write a single compact run into a cell; sizes the paragraph mark too so empty
    cells do not push the (exact) row height up. Never edits the global Normal style."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = ALIGN.CENTER if align == "center" else ALIGN.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(font_pt + 1)
    _vcenter(cell)
    # paragraph-mark run properties: small size so an empty cell stays short
    ppr = p._p.get_or_add_pPr()
    mrpr = ppr.find(qn("w:rPr"))
    if mrpr is None:
        mrpr = OxmlElement("w:rPr")
        ppr.append(mrpr)
    msz = OxmlElement("w:sz")
    msz.set(qn("w:val"), str(int(font_pt * 2)))
    mrpr.append(msz)
    if val is None or val == "":
        return
    run = p.add_run(str(val))
    f = run.font
    f.name = ascii_font
    f.size = Pt(font_pt)
    f.bold = bold
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), ascii_font)
    rf.set(qn("w:hAnsi"), ascii_font)
    rf.set(qn("w:eastAsia"), eastasia)
    if color:
        f.color.rgb = RGBColor.from_string(color)


def _set_cell_runs(cell, runs, font_pt, ascii_font="Arial", eastasia="SimSun",
                   header_bold=False):
    """Render styled runs [{t, b, i, color}] into a cell (per-run bold/italic/color).
    Mirrors _set_cell_text but supports several runs, so a free-table cell can carry
    the same inline styling as a body paragraph."""
    cell.text = ""
    p = cell.paragraphs[0]
    if font_pt:
        p.alignment = ALIGN.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(font_pt + 1)
        _vcenter(cell)
    for rn in (runs or []):
        run = p.add_run(rn.get("t", ""))
        f = run.font
        if font_pt:
            f.name = ascii_font
            f.size = Pt(font_pt)
            rpr = run._r.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.insert(0, rf)
            rf.set(qn("w:ascii"), ascii_font)
            rf.set(qn("w:hAnsi"), ascii_font)
            rf.set(qn("w:eastAsia"), eastasia)
        if bool(rn.get("b")) or header_bold:   # only set when true -> no stray <w:b w:val="0"/>
            f.bold = True
        if rn.get("i"):
            f.italic = True
        if rn.get("color"):
            f.color.rgb = RGBColor.from_string(rn["color"])


# ---------------------------------------------------------------------------
# Compliance data model helpers (group / axis / limit logic) -- config driven
# ---------------------------------------------------------------------------
def _numv(v):
    """Numeric part of an axis value, supporting a ``[value, "CORNER"]`` 2-tuple."""
    if isinstance(v, (tuple, list)) and len(v) == 2:
        v = v[0]
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


def _fmt_val(v):
    """Render an axis value: ``[value, "CORNER"]`` -> ``"value(CORNER)"``; else as-is."""
    if isinstance(v, (tuple, list)) and len(v) == 2:
        return f"{v[0]}({v[1]})"
    return v


def _violates(limit, nv, smin, smax, en, styp):
    """True when a single simulated value ``nv`` is out of the row's limit.

    Directions: ``le`` (<= upper bound) / ``ge`` (>= target) / ``range`` ([MIN,MAX]).
    Thresholds come from the spec triple / scalar spec -- no hardcoded numbers."""
    if nv is None:
        return False
    if limit == "le":
        thr = smax if smax is not None else (en if en is not None else styp)
        return thr is not None and nv > thr
    if limit == "ge":
        thr = smin if smin is not None else (smax if smax is not None else en)
        return thr is not None and nv < thr
    if limit == "range":
        return (smin is not None and nv < smin) or (smax is not None and nv > smax)
    return False


def flag_positions(row):
    """Axis indices in the simulation group that violate the row's limit.

    Indices 0/1/2 are the simulation [MIN,TYP,MAX] triple; index 3 is the NTWC
    corner (``_axis_value`` lays the sim group out as [MIN,TYP,MAX,NTWC]). NTWC
    is evaluated against its own spec threshold (``spec_ntwc``) when present,
    otherwise it falls back to the same MTM-derived thresholds as the triple --
    so a None NTWC (no corner value) is never flagged. Backward-tolerant: callers
    that only ever looked at {0,1,2} keep working; index 3 is added only when an
    NTWC value is out of spec.

    Directions: ``le`` (<= upper bound) / ``ge`` (>= target) / ``range`` ([MIN,MAX]).
    Thresholds are taken from the row's spec triple / scalar spec. No hardcoded numbers.
    """
    limit = row.get("limit")
    if not limit:
        return set()
    sm = row.get("spec_mtm") or [None, None, None]
    sim = row.get("sim_mtm") or [None, None, None]
    smin, styp, smax = _numv(sm[0]), _numv(sm[1]), _numv(sm[2])
    en = _numv(row.get("spec"))
    flags = set()
    for i, v in enumerate(sim):
        if _violates(limit, _numv(v), smin, smax, en, styp):
            flags.add(i)
    # NTWC corner = axis index 3. Prefer the NTWC-specific spec bound; if the
    # template carries no spec_ntwc, reuse the MTM-derived thresholds.
    nt = _numv(row.get("sim_ntwc"))
    if nt is not None:
        nspec = _numv(row.get("spec_ntwc"))
        n_smin = nspec if nspec is not None else smin
        n_smax = nspec if nspec is not None else smax
        n_en = nspec if nspec is not None else en
        if _violates(limit, nt, n_smin, n_smax, n_en, styp):
            flags.add(3)
    return flags


def _default_axes(cfg):
    return list(cfg.get("axis_labels", ["MIN", "TYP", "MAX", "NTWC"]))


# Dependency-free clip heuristic. The compliance table keeps EXACTLY row heights
# (iron rule 2): a cell that holds more text than one line can hold is silently
# clipped by Word rather than spilling. We can't measure glyph widths without a
# font engine, so we estimate the chars that fit on one line from the cell's
# usable width and the font size, and warn (never resize) when the value is
# longer. Conservative on purpose -- under-warn rather than cry wolf.
#   chars/cm at the reference 7pt is ~8 (calibrated so a real, tight compliance
#   layout's legitimate numeric content -- e.g. "-148.5" in a 0.88cm axis column
#   at 7pt -- does NOT warn, while genuinely oversized free text still does). It
#   scales inversely with font size, so a bigger font fits fewer chars/cm.
_CLIP_CHARS_PER_CM_AT_7PT = 8.0
_CLIP_CELL_PAD_CM = 0.1   # left+right cell margins (~28+28 dxa) eat usable width


def _clip_capacity(width_cm, font_pt):
    """Estimated max single-line characters for a cell of ``width_cm`` at ``font_pt``."""
    usable = max(0.0, float(width_cm) - _CLIP_CELL_PAD_CM)
    if usable <= 0 or not font_pt:
        return 0
    cpc = _CLIP_CHARS_PER_CM_AT_7PT * (7.0 / float(font_pt))
    return int(usable * cpc)


def make_groups(data, cfg):
    """Build the ordered list of value groups: one ``spec`` group + one per simulation."""
    default_axes = _default_axes(cfg)
    groups = [dict(key="spec", title=data.get("spec_name", "Spec"), stage=None,
                   role="spec", axes=list(default_axes))]
    sims = data.get("sims") or [{"key": "sim", "title": "Sim", "stage": None}]
    for sim in sims:
        groups.append(dict(key=sim["key"], title=sim.get("title", sim["key"]),
                           stage=sim.get("stage"), role="sim",
                           axes=list(sim.get("axes", default_axes))))
    return groups


def _axis_value(row, gkey, ai):
    if gkey == "spec":
        arr = list(row.get("spec_mtm") or [None, None, None]) + [row.get("spec_ntwc")]
    else:
        arr = list(row.get("sim_mtm") or [None, None, None]) + [row.get("sim_ntwc")]
    return arr[ai] if ai < len(arr) else None


def _plan_columns(groups, show_spec_col, w):
    """Column plan; each entry carries a render kind and width (cm).  ``w`` = per-kind cm map."""
    plan = [dict(kind="cat", label="Category", w=w["cat"]),
            dict(kind="item", label="Item", w=w["item"])]
    if show_spec_col:
        plan.append(dict(kind="spec", label="Spec", w=w["spec"]))
    for g in groups:
        plan.append(dict(kind="spacer", w=w["spacer"]))
        for ai, ax in enumerate(g["axes"]):
            plan.append(dict(kind="axis", group=g["key"], role=g["role"],
                             axis=ai, label=ax, w=w["axis"]))
    plan.append(dict(kind="spacer", w=w["spacer"]))
    plan.append(dict(kind="unit", label="Unit", w=w["unit"]))
    return plan


# ---------------------------------------------------------------------------
# Public: compliance datatable renderer (inline, portrait, exact row heights)
# ---------------------------------------------------------------------------
def render_datatable(doc, data, cfg):
    """Render a data-driven compliance table inline in the document's portrait body.

    ``cfg`` is the template config's ``compliance`` section:
        col_w_cm, font_pt, row_h_pt{header,data}, axis_labels, fills{header,setting,result},
        setting_kinds, flag_color, borders{val,sz,color}.
    """
    w = cfg["col_w_cm"]
    font_pt = cfg.get("font_pt", 7)
    fills = cfg["fills"]
    setting_kinds = set(cfg.get("setting_kinds", ["common_setting", "module_setting", "tb"]))
    flag_color = cfg.get("flag_color", "FF0000")
    row_h = cfg.get("row_h_pt", {"header": 12, "data": 10})
    bd = cfg.get("borders", {"val": "single", "sz": 4, "color": "000000"})

    groups = make_groups(data, cfg)
    show_spec_col = not any(g["role"] == "spec" for g in groups)
    plan = _plan_columns(groups, show_spec_col, w)
    ncols = len(plan)
    rows = data["rows"]
    nrows = 3 + len(rows)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = 1  # center
    _table_fixed_layout(table)
    _table_borders(table, val=bd.get("val", "single"), sz=bd.get("sz", 4),
                   color=bd.get("color", "000000"))
    _cell_margins(table, top=0, bottom=0, left=28, right=28)

    # column widths (per-kind), applied to every row so the fixed layout sticks
    for idx, p in enumerate(plan):
        cw = Cm(p["w"])
        for r in range(nrows):
            table.cell(r, idx).width = cw

    col_of = {}
    group_axis_cols = {g["key"]: [] for g in groups}
    for idx, p in enumerate(plan):
        if p["kind"] in ("cat", "item", "spec", "unit"):
            col_of[p["kind"]] = idx
        elif p["kind"] == "axis":
            col_of[("axis", p["group"], p["axis"])] = idx
            group_axis_cols[p["group"]].append(idx)

    # ---- header band: 3 rows ----
    for r in range(3):
        for c in range(ncols):
            _shade(table.cell(r, c), fills["header"])
    for key, label in (("cat", "Category"), ("item", "Item"),
                       ("spec", "Spec"), ("unit", "Unit")):
        if key in col_of:
            c = col_of[key]
            table.cell(0, c).merge(table.cell(2, c))
            _set_cell_text(table.cell(0, c), label, font_pt, bold=True)
    for g in groups:
        cc = group_axis_cols[g["key"]]
        table.cell(0, cc[0]).merge(table.cell(0, cc[-1]))
        _set_cell_text(table.cell(0, cc[0]), g["title"], font_pt, bold=True)
        if g.get("stage"):
            table.cell(1, cc[0]).merge(table.cell(1, cc[-1]))
            _set_cell_text(table.cell(1, cc[0]), g["stage"], font_pt, bold=True)
        for ai, ax in enumerate(g["axes"]):
            _set_cell_text(table.cell(2, col_of[("axis", g["key"], ai)]), ax,
                           font_pt, bold=True)

    # repeat the 3 header rows on page breaks
    for hr in range(3):
        trPr = table.rows[hr]._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader")
        th.set(qn("w:val"), "true")
        trPr.append(th)

    # ---- data rows ----
    start = 3
    catg = []
    i = 0
    while i < len(rows):
        j = i
        while j + 1 < len(rows) and rows[j + 1]["cat"] == rows[i]["cat"]:
            j += 1
        catg.append((i, j))
        i = j + 1

    warnings = []
    flagged_rows = 0

    def _clip_check(val, width_cm, kind, row_label):
        """Record a row_clip_risk warning when a cell value is too long for one
        line (EXACTLY height => Word clips it). Metadata only -- never resizes."""
        if val is None:
            return
        s = str(val)
        cap = _clip_capacity(width_cm, font_pt)
        if cap and len(s) > cap:
            shown = s if len(s) <= 30 else s[:30] + "..."
            warnings.append({
                "type": "row_clip_risk",
                "detail": '%s "%s" may clip (%d chars, ~%d fit at %.2gcm)'
                          % (kind, shown, len(s), cap, width_cm),
                "location": row_label,
            })

    for (g0, g1) in catg:
        for gi in range(g0, g1 + 1):
            row = rows[gi]
            r = start + gi
            band = fills["setting"] if row["kind"] in setting_kinds else fills["result"]
            flags = flag_positions(row)
            if flags:
                flagged_rows += 1
            row_label = "row %d (%s)" % (gi, row.get("item", ""))
            # When a row merges its sim axis cells into one wide cell, the
            # per-axis width no longer bounds the text, so skip the clip check
            # for sim cells (would otherwise over-warn on a non-clipping merge).
            span = bool(row.get("sim_span"))
            for idx, p in enumerate(plan):
                if p["kind"] == "cat":
                    continue
                _shade(table.cell(r, idx), band)
                if p["kind"] == "item":
                    _set_cell_text(table.cell(r, idx), row["item"], font_pt)
                    _clip_check(row.get("item"), p["w"], "item", row_label)
                elif p["kind"] == "unit":
                    _set_cell_text(table.cell(r, idx), row["unit"], font_pt)
                    _clip_check(row.get("unit"), p["w"], "unit", row_label)
                elif p["kind"] == "spec":
                    _set_cell_text(table.cell(r, idx), row.get("spec"), font_pt)
                    _clip_check(row.get("spec"), p["w"], "spec", row_label)
                elif p["kind"] == "axis":
                    v = _axis_value(row, p["group"], p["axis"])
                    red = (p["role"] == "sim" and p["axis"] in flags)
                    # B&W-safe marker: flagged sim values are red AND bold so the
                    # flag survives a grayscale print. Bold adds no measurable
                    # width here (fixed column layout) so it cannot push a cell to
                    # wrap/clip -- the EXACTLY row-height no-spill guarantee holds.
                    _set_cell_text(table.cell(r, idx), _fmt_val(v), font_pt,
                                   bold=red, color=(flag_color if red else None))
                    # skip the clip check for merged sim cells (no per-axis bound)
                    if not (span and p["role"] == "sim"):
                        _clip_check(_fmt_val(v), p["w"],
                                    "%s.%s" % (p["group"], p["label"]), row_label)
            if row.get("sim_span"):
                for g in groups:
                    if g["role"] == "sim":
                        cc = group_axis_cols[g["key"]]
                        table.cell(r, cc[0]).merge(table.cell(r, cc[2]))
                        break
        # vertical merge of the category column
        cc = col_of["cat"]
        r0, r1 = start + g0, start + g1
        band = fills["setting"] if rows[g0]["kind"] in setting_kinds else fills["result"]
        table.cell(r0, cc).merge(table.cell(r1, cc))
        _shade(table.cell(r0, cc), band)
        _set_cell_text(table.cell(r0, cc), rows[g0]["cat"], font_pt, bold=True)

    # ---- exact fixed row heights: shrink inline, never spill / rotate ----
    for r in range(nrows):
        table.rows[r].height = Pt(row_h["header"] if r < 3 else row_h["data"])
        table.rows[r].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # RETURN SHAPE (CONTRACT): a result dict. ``table`` is the python-docx Table
    # object (same one previously returned bare); the rest is render metadata the
    # engine feeds into its warnings manifest.
    return {
        "table": table,
        "total_rows": len(rows),
        "flagged_rows": flagged_rows,
        "warnings": warnings,
    }


# ---------------------------------------------------------------------------
# Public: free-table renderer (arbitrary rows/cols)
# ---------------------------------------------------------------------------
def render_free_table(doc, rows, cfg, header_rows=1, merges=None, col_w=None):
    """Render an arbitrary table. ``cfg`` = template config's ``free_table`` section:
        header_fill, border{val,sz,color}, font_pt(optional).

    RETURN SHAPE (CONTRACT): the same result-dict shape as render_datatable --
    {"table": <Table or None>, "total_rows": int, "flagged_rows": 0,
     "warnings": []}. A free table has no compliance limits, so flagged_rows is
     always 0 and warnings is always empty; the keys exist for a uniform caller.
     ``table`` is None when there are no rows to render."""
    if not rows:
        return {"table": None, "total_rows": 0, "flagged_rows": 0, "warnings": []}
    ncols = max(len(r) for r in rows)
    nrows = len(rows)
    header_fill = cfg.get("header_fill", "D9D9D9")
    bd = cfg.get("border", {"val": "single", "sz": 4, "color": "000000"})
    font_pt = cfg.get("font_pt")

    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = 1
    if col_w:
        _table_fixed_layout(table)
    _table_borders(table, val=bd.get("val", "single"), sz=bd.get("sz", 4),
                   color=bd.get("color", "000000"))

    if col_w:
        for idx, cw in enumerate(col_w):
            for r in range(nrows):
                if idx < ncols:
                    table.cell(r, idx).width = Cm(cw)

    for r, rowvals in enumerate(rows):
        for c in range(ncols):
            val = rowvals[c] if c < len(rowvals) else ""
            cell = table.cell(r, c)
            if r < header_rows:
                _shade(cell, header_fill)
            runs = val.get("runs") if isinstance(val, dict) else None
            if isinstance(runs, list):
                _set_cell_runs(cell, runs, font_pt, header_bold=(r < header_rows))
            elif font_pt:
                _set_cell_text(cell, val, font_pt, bold=(r < header_rows), align="center")
            else:
                cell.text = "" if val is None else str(val)
                if r < header_rows:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True

    for m in (merges or []):
        r, c, rs, cs = m["r"], m["c"], m.get("rs", 1), m.get("cs", 1)
        table.cell(r, c).merge(table.cell(r + rs - 1, c + cs - 1))

    return {"table": table, "total_rows": nrows, "flagged_rows": 0, "warnings": []}
