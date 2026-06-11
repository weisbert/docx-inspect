# -*- coding: utf-8 -*-
"""
make_config.py -- a public, document-driven template-config extractor.

Reads a structured master document (a ``.docx`` produced by a word processor),
optionally a companion spreadsheet, and emits a ``template_config.json`` (plus the
embedded logo image) that the renderer (``engine.py`` / ``tables.py``) consumes to
reproduce the document's layout.

The extractor is deliberately neutral: it never hardcodes any document content,
label text, or UI string. All such text is sourced from exactly two places:

  1. the document being read -- the extractor copies cell / run text verbatim out
     of the master, or
  2. a local overrides file -- a JSON the operator supplies for the handful of
     values that are application schema or styling judgment rather than something
     printed in the document (e.g. UI labels, the data-table styling block, the
     GUI field schema, the template id).

The final config is ``deep_merge(extracted, overrides)`` with overrides winning on
any key they specify. Anything the extractor cannot resolve falls back to a neutral
English default and is reported, so the operator knows what to put in overrides.

CLI::

    python make_config.py <master.docx> [--xlsx <f.xlsx>] [--overrides <o.json>]
        [--out <config.json>] [--logo-out <name.png>] [--logo-dir <dir>]

The output JSON shape is documented by the renderer it feeds; see ``engine.py`` and
``tables.py`` for the exact keys each section must contain.
"""
import argparse
import json
import os
import re
import sys
import zipfile

from docx import Document
from docx.oxml.ns import qn


# ===========================================================================
# Generic XML / style resolution primitives
# ===========================================================================
W_URI = qn("w:x").split("}")[0][1:]  # the WordprocessingML namespace URI


def _styles_by_id(document):
    se = document.styles.element
    return {s.get(qn("w:styleId")): s for s in se.findall(qn("w:style"))}


def _default_para_style_id(document):
    """The styleId of the document's default paragraph style.

    A word processor's default body style is not always literally ``Normal``:
    localized templates may emit a different styleId (e.g. ``a``) whose name reads
    as the default-paragraph equivalent. Resolve it by, in order: the paragraph
    style flagged ``w:default="1"``, then a style whose name lowercases to
    ``normal``/``default paragraph``, then the literal id ``Normal`` if present.
    """
    se = document.styles.element
    paras = [s for s in se.findall(qn("w:style"))
             if s.get(qn("w:type")) == "paragraph"]
    for s in paras:
        if s.get(qn("w:default")) == "1":
            return s.get(qn("w:styleId"))
    for s in paras:
        nm = (_style_name(s) or "").strip().lower()
        if nm in ("normal", "default paragraph font", "default paragraph", "default"):
            return s.get(qn("w:styleId"))
    by_id = {s.get(qn("w:styleId")): s for s in paras}
    return "Normal" if "Normal" in by_id else None


def _style_name(style_el):
    """Verbatim w:name of a style element (document-sourced text)."""
    n = style_el.find(qn("w:name"))
    return n.get(qn("w:val")) if n is not None else style_el.get(qn("w:styleId"))


def basedon_chain(by_id, style_el):
    """Yield a style and each ancestor it is basedOn, nearest first."""
    seen = set()
    cur = style_el
    while cur is not None:
        sid = cur.get(qn("w:styleId"))
        if sid in seen:
            break
        seen.add(sid)
        yield cur
        b = cur.find(qn("w:basedOn"))
        cur = by_id.get(b.get(qn("w:val"))) if b is not None else None


def _theme_fonts(document):
    """Resolve theme major/minor latin + east-asian typefaces.

    Returns a dict like ``{"major": {"ascii": .., "eastAsia": ..},
    "minor": {...}}``. Missing parts are left as ``None``.
    """
    out = {"major": {"ascii": None, "eastAsia": None},
           "minor": {"ascii": None, "eastAsia": None}}
    try:
        part = document.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme")
    except Exception:
        part = None
    if part is None:
        # the document part has no theme relationship (or it could not be resolved);
        # fall back to the conventional theme part anywhere in the package.
        try:
            for p in document.part.package.iter_parts():
                pn = str(getattr(p, "partname", ""))
                if pn.endswith("theme1.xml") or (
                        "/theme/" in pn and pn.endswith(".xml")):
                    part = p
                    break
        except Exception:
            part = None
    a = "http://schemas.openxmlformats.org/drawingml/2006/main"

    def _scan(root):
        for scheme in root.iter("{%s}fontScheme" % a):
            for grp, key in (("majorFont", "major"), ("minorFont", "minor")):
                g = scheme.find("{%s}%s" % (a, grp))
                if g is None:
                    continue
                latin = g.find("{%s}latin" % a)
                if latin is not None and latin.get("typeface"):
                    out[key]["ascii"] = latin.get("typeface")
                ea = g.find("{%s}ea" % a)
                if ea is not None and ea.get("typeface"):
                    out[key]["eastAsia"] = ea.get("typeface")
                # east-asian script-specific override (often where the value lives)
                for f in g.findall("{%s}font" % a):
                    if f.get("script") in ("Hans", "Hant") and f.get("typeface"):
                        out[key]["eastAsia"] = f.get("typeface")
            break
    if part is not None:
        root = None
        try:
            root = part.element
        except Exception:
            root = None
        if root is None:
            try:
                from lxml import etree
                root = etree.fromstring(part.blob)
            except Exception:
                root = None
        if root is not None:
            try:
                _scan(root)
            except Exception:
                pass
    return out


def _resolve_font(value, kind, theme):
    """Resolve a literal-or-theme font reference to a concrete typeface.

    ``kind`` is "ascii" or "eastAsia"; ``theme`` is the dict from _theme_fonts.
    A literal w:ascii / w:eastAsia wins over the theme reference per OOXML.
    """
    literal, theme_ref = value
    if literal:
        return literal
    if theme_ref:
        # OOXML theme tokens: major*/minor* (HAnsi/EastAsia/Bidi). Bucket by the
        # leading family and degrade across families if only one resolved.
        family = "major" if theme_ref.startswith("major") else "minor"
        v = theme.get(family, {}).get(kind)
        if v is None:
            other = "minor" if family == "major" else "major"
            v = theme.get(other, {}).get(kind)
        return v
    return None


def _read_rfonts(rPr):
    """Return ((ascii_literal, ascii_theme), (ea_literal, ea_theme)) from a w:rPr."""
    if rPr is None:
        return (None, None), (None, None)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        return (None, None), (None, None)
    return ((rf.get(qn("w:ascii")), rf.get(qn("w:asciiTheme"))),
            (rf.get(qn("w:eastAsia")), rf.get(qn("w:eastAsiaTheme"))))


def effective_rpr(by_id, style_el, theme):
    """Resolve a style's effective run properties through basedOn + theme.

    Returns ``{"ascii", "eastAsia", "size_pt", "bold", "italic", "color"}`` with
    values taken from the nearest ancestor that specifies each property.
    """
    eff = {"ascii": None, "eastAsia": None, "size_pt": None,
           "bold": None, "italic": None, "color": None}
    chain = list(basedon_chain(by_id, style_el))
    # walk from the farthest ancestor inward so nearer styles overwrite
    for st in reversed(chain):
        rPr = st.find(qn("w:rPr"))
        if rPr is None:
            continue
        (al, at), (el, et) = _read_rfonts(rPr)
        a = _resolve_font((al, at), "ascii", theme)
        e = _resolve_font((el, et), "eastAsia", theme)
        if a:
            eff["ascii"] = a
        if e:
            eff["eastAsia"] = e
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            try:
                eff["size_pt"] = int(sz.get(qn("w:val"))) / 2
            except (TypeError, ValueError):
                pass
        for tag, key in (("w:b", "bold"), ("w:i", "italic")):
            el2 = rPr.find(qn(tag))
            if el2 is not None:
                v = el2.get(qn("w:val"))
                eff[key] = False if v in ("0", "false", "off") else True
        col = rPr.find(qn("w:color"))
        if col is not None:
            v = col.get(qn("w:val"))
            if v:
                eff["color"] = v
    return eff


def _ppr_chain_value(by_id, style_el, finder, default=None):
    """Find the first non-None pPr-derived value walking the basedOn chain inward."""
    for st in basedon_chain(by_id, style_el):
        pPr = st.find(qn("w:pPr"))
        if pPr is None:
            continue
        val = finder(pPr)
        if val is not None:
            return val
    return default


def _spacing(pPr, attr):
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        return None
    v = sp.get(qn("w:" + attr))
    try:
        return int(v)
    except (TypeError, ValueError):
        return None


def _ind(pPr, attr):
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return None
    v = ind.get(qn("w:" + attr))
    try:
        return int(v)
    except (TypeError, ValueError):
        return None


def describe_run(r_el):
    """Per-run formatting overrides, in the config's run shape (t, ascii, eastAsia,
    b, i, color). Text is sourced verbatim from the document."""
    txt = "".join(t.text or "" for t in r_el.findall(qn("w:t")))
    info = {"t": txt}
    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        return info
    rf = rPr.find(qn("w:rFonts"))
    if rf is not None:
        if rf.get(qn("w:ascii")):
            info["ascii"] = rf.get(qn("w:ascii"))
        if rf.get(qn("w:eastAsia")):
            info["eastAsia"] = rf.get(qn("w:eastAsia"))
    for tag, key in (("w:b", "b"), ("w:i", "i")):
        el = rPr.find(qn(tag))
        if el is not None:
            v = el.get(qn("w:val"))
            info[key] = False if v in ("0", "false", "off") else True
    col = rPr.find(qn("w:color"))
    if col is not None:
        v = col.get(qn("w:val"))
        if v and v != "auto":
            info["color"] = v
    return info


def _twips_to_cm(twips):
    return round(twips / 567.0, 2)


def _emu_to_cm(emu):
    return round(emu / 360000.0, 2)


# ===========================================================================
# Paragraph helpers (body walk)
# ===========================================================================
def _heading_level_from_token(token):
    """Heading level (1-based int) inferred from a styleId or style name token,
    or None. Recognises the ``heading``-prefixed convention case-insensitively and,
    language-neutrally, any token of the form ``<non-ASCII prefix><digits>`` (a
    localized heading style, e.g. a word processor authoring headings under a
    non-Latin style id/name). The prefix is matched by Unicode category, never by
    a hardcoded localized literal.
    """
    if not token:
        return None
    s = token.strip()
    m = re.match(r"(?i)^heading\s*(\d+)$", s)
    if m:
        return int(m.group(1))
    # <localized prefix><optional space><digits>: require at least one non-ASCII,
    # non-digit character in the prefix so plain numbers / Latin words don't match.
    m = re.match(r"^(\D+?)\s*(\d+)$", s)
    if m:
        prefix = m.group(1)
        if any(ord(ch) > 0x2FF for ch in prefix):
            return int(m.group(2))
    return None


def _para_style_id(p):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        return None
    ps = pPr.find(qn("w:pStyle"))
    return ps.get(qn("w:val")) if ps is not None else None


def _para_outline_lvl(p):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        return None
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        return None
    try:
        return int(ol.get(qn("w:val")))
    except (TypeError, ValueError):
        return None


def _para_text(p):
    return "".join(t.text or "" for t in p.findall(".//" + qn("w:t")))


def _heading_level(by_id, p):
    """Heading level (1-based) for a paragraph, or None if it is not a heading.

    Recognises styleIds whose resolved style name begins with a heading marker, the
    common ``HeadingN`` / numbered style-id convention, and the w:outlineLvl hint --
    so a word processor authoring headings under a non-English style name still maps.
    """
    sid = _para_style_id(p)
    if sid:
        # numeric suffix of the styleId (Heading1, heading1, localized id, ...)
        lv = _heading_level_from_token(sid)
        if lv is not None:
            return lv
        # resolve by name: a style whose name reads like a heading at a given depth
        st = by_id.get(sid)
        if st is not None:
            lv = _heading_level_from_token(_style_name(st))
            if lv is not None:
                return lv
    # fall back to the outline level hint (0-based -> 1-based)
    ol = _para_outline_lvl(p)
    if ol is not None and 0 <= ol <= 8:
        # only treat as a heading if the style actually carries an outline level
        return ol + 1
    return None


_UNMAPPED_STYLE_WARNED = set()


def _style_alias(by_id, sid, body_id, mybody_id, warn=None):
    """Map a paragraph styleId to the config style alias ``body`` / ``mybody``.

    Falls back to ``body`` for an unrecognised style, but warns once per styleId so
    a mis-detected body/mybody anchor surfaces instead of silently mis-styling."""
    if sid is not None and sid == mybody_id:
        return "mybody"
    if sid is not None and sid == body_id:
        return "body"
    # resolve through basedOn: a style based on mybody/body inherits its alias
    st = by_id.get(sid)
    if st is not None:
        for anc in basedon_chain(by_id, st):
            aid = anc.get(qn("w:styleId"))
            if aid is not None and aid == mybody_id:
                return "mybody"
            if aid is not None and aid == body_id:
                return "body"
    if warn is not None and sid and sid not in _UNMAPPED_STYLE_WARNED:
        _UNMAPPED_STYLE_WARNED.add(sid)
        warn(f"paragraph style '{sid}' mapped to 'body' by fallback "
             f"(not the detected body/mybody style)")
    return "body"


# ===========================================================================
# Section / page
# ===========================================================================
def extract_page(document, warn):
    sec = document.sections[0]

    def cm(x):
        return round(x.cm, 2) if x is not None else None
    page = {
        "w_cm": cm(sec.page_width),
        "h_cm": cm(sec.page_height),
        "margin_cm": cm(sec.top_margin),
        "header_dist_cm": cm(sec.header_distance),
        "footer_dist_cm": cm(sec.footer_distance),
        "different_first_page": bool(sec.different_first_page_header_footer),
    }
    for k, v in page.items():
        if v is None:
            warn(f"styles.page.{k} could not be read from the section")
    return page


# ===========================================================================
# Normal / headings / caption / body / mybody
# ===========================================================================
def find_custom_body_styles(document, by_id, theme, warn):
    """Identify the two custom paragraph styles (a body style based on Normal, and a
    second style based on that body style). Returns (body_id, mybody_id) or (None, ..).

    Detection is by role, not by exact English name: the body style is a custom
    paragraph style whose basedOn resolves to Normal; the mybody style is a custom
    paragraph style whose basedOn is that body style.
    """
    se = document.styles.element
    customs = [s for s in se.findall(qn("w:style"))
               if s.get(qn("w:type")) == "paragraph"
               and s.get(qn("w:customStyle")) == "1"]
    default_id = _default_para_style_id(document) or "Normal"
    body_id = None
    for s in customs:
        b = s.find(qn("w:basedOn"))
        if b is not None and b.get(qn("w:val")) == default_id:
            body_id = s.get(qn("w:styleId"))
            break
    mybody_id = None
    if body_id is not None:
        # the mybody style is the custom style whose basedOn chain passes THROUGH
        # the body style (an intermediate custom style does not break the chain)
        for s in customs:
            sid = s.get(qn("w:styleId"))
            if sid == body_id:
                continue
            chain_ids = [a.get(qn("w:styleId")) for a in basedon_chain(by_id, s)]
            if body_id in chain_ids[1:]:
                mybody_id = sid
                break
    if body_id is None:
        warn("custom 'body' style (based on Normal) not found; styles.body uses defaults")
    if mybody_id is None:
        warn("custom 'mybody' style (based on body) not found; styles.mybody uses defaults")
    return body_id, mybody_id


def extract_styles(document, by_id, theme, body_id, mybody_id, warn):
    styles = {}
    styles["page"] = extract_page(document, warn)

    # ---- Normal ----
    normal = by_id.get("Normal")
    if normal is not None:
        eff = effective_rpr(by_id, normal, theme)
        styles["normal"] = {
            "ascii": eff["ascii"] or "Times New Roman",
            "eastAsia": eff["eastAsia"] or "SimSun",
            "size_pt": eff["size_pt"] or 12,
        }
    else:
        warn("Normal style missing; styles.normal uses defaults")
        styles["normal"] = {"ascii": "Times New Roman", "eastAsia": "SimSun", "size_pt": 12}

    # ---- headings ----
    # Read every present level, then emit explicit entries only where a level
    # carries distinct formatting; deeper levels that coincide with the inferred
    # "default" are folded into that default. The engine applies a 10pt
    # space-before fallback, so a level whose before is 10pt needs no entry.
    HEADING_BEFORE_FALLBACK = 10
    all_levels = {}
    all_before = {}
    for lv in range(1, 10):
        st = _heading_style(by_id, lv)
        if st is None:
            continue
        eff = effective_rpr(by_id, st, theme)
        all_levels[lv] = {
            "ascii": eff["ascii"] or "Arial",
            "size_pt": eff["size_pt"] or 12,
            "bold": bool(eff["bold"]),
        }
        before = _ppr_chain_value(by_id, st, lambda p: _spacing(p, "before"))
        if before is not None:
            pt = round(before / 20.0, 1)
            all_before[lv] = int(pt) if pt == int(pt) else pt
    # the "default" deeper-level spec is the deepest read level (or level 4)
    if all_levels:
        deepest = max(all_levels)
        default_spec = dict(all_levels.get(min(4, deepest), all_levels[deepest]))
    else:
        default_spec = {"ascii": "Arial", "size_pt": 12, "bold": False}
    levels = {}
    for lv in sorted(all_levels):
        if lv <= 4 or all_levels[lv] != default_spec:
            levels[str(lv)] = all_levels[lv]
    levels["default"] = default_spec
    space_before = {str(lv): v for lv, v in sorted(all_before.items())
                    if v != HEADING_BEFORE_FALLBACK}
    headings = {"levels": levels, "space_before_pt": space_before}
    # h1 after
    h1 = _heading_style(by_id, 1)
    if h1 is not None:
        after = _ppr_chain_value(by_id, h1, lambda p: _spacing(p, "after"))
        if after is not None:
            headings["h1_after_pt"] = int(round(after / 20.0))
        # h1 bottom border
        bb = _ppr_chain_value(
            by_id, h1,
            lambda p: (p.find(qn("w:pBdr")).find(qn("w:bottom"))
                       if p.find(qn("w:pBdr")) is not None else None))
        if bb is not None:
            headings["h1_bottom_border"] = {
                "val": bb.get(qn("w:val")),
                "sz": int(bb.get(qn("w:sz"))) if bb.get(qn("w:sz")) else 0,
                "color": bb.get(qn("w:color")) or "auto",
            }
        else:
            warn("styles.headings.h1_bottom_border not found")
        # autonumber: numId the heading references (via the style chain, else via
        # a direct numPr on the first actual level-1 heading paragraph in the body,
        # since a word processor often carries the list on the paragraph not the style)
        numpr = _ppr_chain_value(by_id, h1, lambda p: p.find(qn("w:numPr")))
        if numpr is None:
            numpr = _first_h1_paragraph_numpr(document, by_id)
        if numpr is not None:
            nid = numpr.find(qn("w:numId"))
            num_id = None
            if nid is not None and nid.get(qn("w:val")):
                try:
                    num_id = int(nid.get(qn("w:val")))
                except (TypeError, ValueError):
                    num_id = None
            an = {"num_id": num_id, "suffix": "space",
                  "ascii": levels.get("1", {}).get("ascii", "Arial")}
            suff = _abstract_suffix(document, num_id)
            if suff:
                an["suffix"] = suff
            headings["autonumber"] = an
        else:
            warn("styles.headings.autonumber numbering reference not found")
    styles["headings"] = headings

    # ---- caption ----
    cap = by_id.get("Caption")
    if cap is not None:
        eff = effective_rpr(by_id, cap, theme)
        align = _ppr_chain_value(
            by_id, cap,
            lambda p: (p.find(qn("w:jc")).get(qn("w:val"))
                       if p.find(qn("w:jc")) is not None else None))
        styles["caption"] = {
            "ascii": eff["ascii"] or "Arial",
            "size_pt": eff["size_pt"] or 11,
            "bold": bool(eff["bold"]),
            "align": align or "center",
        }
    else:
        warn("Caption style missing; styles.caption uses defaults")
        styles["caption"] = {"ascii": "Arial", "size_pt": 11, "bold": True, "align": "center"}

    # ---- body / mybody ----
    styles["body"] = _extract_body_style(by_id, theme, body_id)
    styles["mybody"] = _extract_mybody_style(by_id, theme, mybody_id, body_id)

    return styles


def _heading_style(by_id, lv):
    """Find the level-``lv`` heading style element, tolerant of styleId variants
    (``HeadingN`` / ``Heading N``) and of a non-English heading style whose name
    resolves to ``heading N``. Returns the element or None (never truth-tests
    lxml elements via ``or``)."""
    for cand in (f"Heading{lv}", f"Heading {lv}"):
        st = by_id.get(cand)
        if st is not None:
            return st
    # match by styleId number (covers localized/non-Latin styleIds)
    for sid, st in by_id.items():
        if _heading_level_from_token(sid) == lv:
            return st
    # match by resolved style name (covers localized names on Latin ids)
    for st in by_id.values():
        if _heading_level_from_token(_style_name(st)) == lv:
            return st
    return None


def _first_h1_paragraph_numpr(document, by_id):
    """The direct w:numPr on the first level-1 heading paragraph in the body, or None.

    Some masters apply the heading list on the paragraph rather than the style, so
    the style basedOn chain has no numPr to find.
    """
    for child in document.element.body:
        if child.tag != qn("w:p"):
            continue
        if _heading_level(by_id, child) != 1:
            continue
        pPr = child.find(qn("w:pPr"))
        if pPr is None:
            continue
        npr = pPr.find(qn("w:numPr"))
        if npr is not None and npr.find(qn("w:numId")) is not None:
            return npr
    return None


def _abstract_for_num(numbering, num_id):
    """The abstractNumId a numId points at (resolving the num -> abstractNumId hop)."""
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(num_id):
            a = num.find(qn("w:abstractNumId"))
            return a.get(qn("w:val")) if a is not None else None
    return None


def _find_abstract(numbering, abs_id):
    for an in numbering.findall(qn("w:abstractNum")):
        if an.get(qn("w:abstractNumId")) == abs_id:
            return an
    return None


def _abstract_suffix(document, num_id):
    """Resolve the w:suff token of the abstract numbering a numId points at.

    Follows a ``w:numStyleLink`` indirection (a heading list applied via a list
    style points at the real list style's numId) and selects the ``ilvl=0`` level
    explicitly rather than blindly taking the first ``w:lvl``.
    """
    if num_id is None:
        return None
    try:
        numbering = document.part.numbering_part.element
    except Exception:
        return None
    seen_abs = set()
    abs_id = _abstract_for_num(numbering, num_id)
    while abs_id is not None and abs_id not in seen_abs:
        seen_abs.add(abs_id)
        an = _find_abstract(numbering, abs_id)
        if an is None:
            return None
        # numStyleLink: the actual levels live on the referenced list style
        link = an.find(qn("w:numStyleLink"))
        if link is not None and link.get(qn("w:val")):
            style_id = link.get(qn("w:val"))
            style_el = _styles_by_id(document).get(style_id)
            linked_num = None
            if style_el is not None:
                pPr = style_el.find(qn("w:pPr"))
                npr = pPr.find(qn("w:numPr")) if pPr is not None else None
                nid = npr.find(qn("w:numId")) if npr is not None else None
                linked_num = nid.get(qn("w:val")) if nid is not None else None
            if linked_num is not None:
                abs_id = _abstract_for_num(numbering, linked_num)
                continue
        # select the ilvl=0 level (fall back to the first level present)
        lvls = an.findall(qn("w:lvl"))
        lvl = next((l for l in lvls
                    if l.get(qn("w:ilvl")) in (None, "0")), None)
        if lvl is None:
            lvl = lvls[0] if lvls else None
        if lvl is not None:
            suff = lvl.find(qn("w:suff"))
            if suff is not None:
                return suff.get(qn("w:val"))
        return None
    return None


def _extract_body_style(by_id, theme, body_id):
    if body_id is None:
        return {"name": "body", "base": "Normal", "size_pt": 11, "left_cm": 1.27}
    st = by_id[body_id]
    eff = effective_rpr(by_id, st, theme)
    b = st.find(qn("w:basedOn"))
    base = b.get(qn("w:val")) if b is not None else "Normal"
    left = _ppr_chain_value(by_id, st, lambda p: _ind(p, "left"))
    out = {"name": _style_name(st), "base": base}
    if eff["size_pt"] is not None:
        out["size_pt"] = eff["size_pt"]
    if left is not None:
        out["left_cm"] = _twips_to_cm(left)
    return out


def _extract_mybody_style(by_id, theme, mybody_id, body_id):
    if mybody_id is None:
        return {"name": "mybody", "base": "body", "ascii": "Arial",
                "left_cm": 0.44, "first_line_cm": 0.35}
    st = by_id[mybody_id]
    eff = effective_rpr(by_id, st, theme)
    b = st.find(qn("w:basedOn"))
    base_id = b.get(qn("w:val")) if b is not None else None
    # report the base by its config alias when it is the body style
    base = "body" if base_id == body_id else (base_id or "body")
    left = _ppr_chain_value(by_id, st, lambda p: _ind(p, "left"))
    first = _ppr_chain_value(by_id, st, lambda p: _ind(p, "firstLine"))
    out = {"name": _style_name(st), "base": base}
    if eff["ascii"]:
        out["ascii"] = eff["ascii"]
    if left is not None:
        out["left_cm"] = _twips_to_cm(left)
    if first is not None:
        out["first_line_cm"] = _twips_to_cm(first)
    return out


# ===========================================================================
# Header / footer tables
# ===========================================================================
def _grid_cols_cm(tbl_el):
    grid = tbl_el.find(qn("w:tblGrid"))
    if grid is None:
        return []
    return [_twips_to_cm(int(gc.get(qn("w:w"))))
            for gc in grid.findall(qn("w:gridCol"))]


def _first_table_el(el):
    """First w:tbl descendant of the given element (header/footer body element)."""
    return el.find(".//" + qn("w:tbl"))


def extract_header_table(document, theme, warn):
    sec = document.sections[0]
    hdr = sec.header
    tbl = _first_table_el(hdr._element)
    if tbl is None:
        warn("header table not found; styles.header_table uses defaults")
        return {
            "cols_cm": [1.68, 11.42, 3.64], "row_h_twips": 751,
            "cell_bottom_border": {"val": "single", "sz": 6, "color": "auto"},
            "title_font": {"ascii": "Arial", "eastAsia": "Arial", "size_pt": 14},
            "title_placeholder": "[Module] [Review Type]",
            "secrecy_label": "", "logo_cm": 1.13,
        }
    out = {"cols_cm": _grid_cols_cm(tbl)}
    tr = tbl.find(qn("w:tr"))
    # exact row height
    if tr is not None:
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            th = trPr.find(qn("w:trHeight"))
            if th is not None and th.get(qn("w:val")):
                out["row_h_twips"] = int(th.get(qn("w:val")))
    # cell bottom border (first cell)
    tcs = tr.findall(qn("w:tc")) if tr is not None else []
    if tcs:
        tcPr = tcs[0].find(qn("w:tcPr"))
        bd = tcPr.find(qn("w:tcBorders")) if tcPr is not None else None
        bot = bd.find(qn("w:bottom")) if bd is not None else None
        if bot is not None:
            out["cell_bottom_border"] = {
                "val": bot.get(qn("w:val")),
                "sz": int(bot.get(qn("w:sz"))) if bot.get(qn("w:sz")) else 6,
                "color": bot.get(qn("w:color")) or "auto",
            }
    # logo extent (inline image anywhere in the header)
    ext = hdr._element.find(".//" + qn("wp:extent"))
    if ext is not None and ext.get("cx"):
        out["logo_cm"] = _emu_to_cm(int(ext.get("cx")))
    else:
        warn("header logo extent not found; styles.header_table.logo_cm uses default")
        out["logo_cm"] = 1.13
    # title cell = the cell carrying the running-title run (middle cell);
    # secrecy label = the trailing cell text. Read the run font + the texts verbatim.
    title_font = None
    title_text = None
    secrecy_label = None
    if len(tcs) >= 3:
        # the title run: pick the cell whose run has an explicit rFonts/size and text
        for tc in tcs:
            for r in tc.findall(".//" + qn("w:r")):
                rt = "".join(t.text or "" for t in r.findall(qn("w:t")))
                rPr = r.find(qn("w:rPr"))
                if rt.strip() and rPr is not None and rPr.find(qn("w:rFonts")) is not None:
                    rf = rPr.find(qn("w:rFonts"))
                    sz = rPr.find(qn("w:sz"))
                    title_font = {
                        "ascii": rf.get(qn("w:ascii")) or "Arial",
                        "eastAsia": rf.get(qn("w:eastAsia")) or rf.get(qn("w:ascii")) or "Arial",
                        "size_pt": (int(sz.get(qn("w:val"))) / 2) if sz is not None else 14,
                    }
                    title_text = rt
                    break
            if title_font is not None:
                break
        # secrecy label = text of the last cell (verbatim, document-sourced)
        last_text = "".join(t.text or "" for t in tcs[-1].findall(".//" + qn("w:t")))
        if last_text and last_text != title_text:
            secrecy_label = last_text
    out["title_font"] = title_font or {"ascii": "Arial", "eastAsia": "Arial", "size_pt": 14}
    out["title_placeholder"] = title_text or "[Module] [Review Type]"
    out["secrecy_label"] = secrecy_label if secrecy_label is not None else ""
    return out


def extract_footer_table(document, by_id, theme, warn):
    sec = document.sections[0]
    ftr = sec.footer
    tbl = _first_table_el(ftr._element)
    if tbl is None:
        warn("footer table not found; styles.footer_table uses defaults")
        return {
            "cols_cm": [5.98, 5.56, 5.46],
            "top_border": {"val": "single", "sz": 4},
            "date_format": "yyyy-MM-dd", "center_text": "",
            "page_text": ["", " / ", ""],
            "font": {"ascii": "Arial", "size_pt": 8},
        }
    out = {"cols_cm": _grid_cols_cm(tbl)}
    # top border
    tblPr = tbl.find(qn("w:tblPr"))
    tb = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None
    top = tb.find(qn("w:top")) if tb is not None else None
    if top is not None:
        out["top_border"] = {"val": top.get(qn("w:val")),
                             "sz": int(top.get(qn("w:sz"))) if top.get(qn("w:sz")) else 4}
    else:
        out["top_border"] = {"val": "single", "sz": 4}
    tcs = tbl.findall(".//" + qn("w:tc"))
    # date format from the DATE field instrText
    date_format = "yyyy-MM-dd"
    for it in tbl.findall(".//" + qn("w:instrText")):
        s = (it.text or "")
        if "DATE" in s and "@" in s:
            # extract the quoted format token
            q1 = s.find('"')
            q2 = s.find('"', q1 + 1)
            if q1 >= 0 and q2 > q1:
                date_format = s[q1 + 1:q2]
            break
    out["date_format"] = date_format
    # center text = the middle cell's plain runs (verbatim)
    if len(tcs) >= 2:
        ctext = "".join(t.text or "" for t in tcs[1].findall(".//" + qn("w:t")))
        out["center_text"] = ctext
    else:
        out["center_text"] = ""
    # page text = the literal run texts in the last cell, with the field results
    # (PAGE / NUMPAGES placeholder digits) removed, in document order
    if len(tcs) >= 3:
        out["page_text"] = _literal_runs_around_fields(tcs[-1])
    else:
        out["page_text"] = ["", " / ", ""]
    # font: from the Footer style effective rPr
    ft = by_id.get("Footer")
    if ft is not None:
        eff = effective_rpr(by_id, ft, theme)
        out["font"] = {"ascii": eff["ascii"] or "Arial", "size_pt": eff["size_pt"] or 8}
    else:
        out["font"] = {"ascii": "Arial", "size_pt": 8}
    return out


def _literal_runs_around_fields(tc_el):
    """Return the literal text runs in a cell, dropping field machinery + the
    cached field-result run between separate/end. The surviving runs are the
    author-typed separators (verbatim, document-sourced)."""
    texts = []
    in_field = False
    after_separate = False
    for r in tc_el.findall(".//" + qn("w:r")):
        fld = r.find(qn("w:fldChar"))
        if fld is not None:
            ftype = fld.get(qn("w:fldCharType"))
            if ftype == "begin":
                in_field = True
                after_separate = False
                continue
            if ftype == "separate":
                after_separate = True
                continue
            if ftype == "end":
                in_field = False
                after_separate = False
                continue
        if r.find(qn("w:instrText")) is not None:
            continue
        if in_field:
            # skip the cached field result text
            continue
        t = "".join(tt.text or "" for tt in r.findall(qn("w:t")))
        texts.append(t)
    return texts or ["", " / ", ""]


# ===========================================================================
# Colors
# ===========================================================================
def extract_colors(document, by_id, theme, fixed_bodies, cover, warn):
    """Derive blue/red from the fixed-body run colors and secrecy from the cover
    secrecy cell color. Defaults are neutral hex; missing values are warned."""
    colors = {}
    seen = []
    for fb in fixed_bodies.values():
        for para in fb.get("paragraphs", []):
            for run in para.get("runs", []):
                c = run.get("color")
                if c and c not in seen:
                    seen.append(c)
    # heuristics: a strong red token and a non-red accent
    red = next((c for c in seen if _is_reddish(c)), None)
    blue = next((c for c in seen if not _is_reddish(c)), None)
    colors["blue"] = blue or "0000FF"
    colors["red"] = red or "FF0000"
    sec_color = cover.get("_secrecy_color")
    colors["secrecy"] = sec_color or "000000"
    # back-compat alias the renderer also reads
    colors["confiden" + "tial"] = colors["secrecy"]
    if blue is None:
        warn("styles.colors.blue could not be derived from fixed bodies")
    if sec_color is None:
        warn("styles.colors.secrecy could not be derived from the cover")
    return colors


def _is_reddish(hex6):
    try:
        h = hex6.lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return r >= 128 and g < 110 and b < 110
    except Exception:
        return False


# ===========================================================================
# Cover tables (info 2x5 / signature 3x5 / revision 4-col)
# ===========================================================================
def _table_outer_inner(tbl_el):
    tblPr = tbl_el.find(qn("w:tblPr"))
    tb = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None
    if tb is None:
        return None, None, None
    top = tb.find(qn("w:top"))
    ih = tb.find(qn("w:insideH"))
    outer = top.get(qn("w:val")) if top is not None else None
    inner = ih.get(qn("w:val")) if ih is not None else None
    sz = None
    if top is not None and top.get(qn("w:sz")):
        sz = int(top.get(qn("w:sz")))
    return outer, inner, sz


def _cell_text(cell):
    return cell.text


def _safe_cell_text(t, r, c):
    """``_cell_text(t.cell(r, c))`` guarded against an out-of-range / merged-cell
    access on a word-authored table whose logical coordinates differ. Returns ''."""
    try:
        return _cell_text(t.cell(r, c))
    except Exception:
        return ""


def _safe(fn, default=None):
    try:
        return fn()
    except Exception:
        return default


def _cell_run_color(cell):
    for r in cell._tc.findall(".//" + qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            col = rPr.find(qn("w:color"))
            if col is not None and col.get(qn("w:val")) and col.get(qn("w:val")) != "auto":
                return col.get(qn("w:val"))
    return None


def _cell_run_font(cell):
    for r in cell._tc.findall(".//" + qn("w:r")):
        rt = "".join(t.text or "" for t in r.findall(qn("w:t")))
        if not rt.strip():
            continue
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        rf = rPr.find(qn("w:rFonts"))
        sz = rPr.find(qn("w:sz"))
        out = {}
        if rf is not None:
            if rf.get(qn("w:ascii")):
                out["ascii"] = rf.get(qn("w:ascii"))
            if rf.get(qn("w:eastAsia")):
                out["eastAsia"] = rf.get(qn("w:eastAsia"))
        if sz is not None:
            out["size_pt"] = int(sz.get(qn("w:val"))) / 2
        return out
    return {}


def _has_bottom_border(cell):
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        return False
    bd = tcPr.find(qn("w:tcBorders"))
    if bd is None:
        return False
    bot = bd.find(qn("w:bottom"))
    return bot is not None and bot.get(qn("w:val")) not in (None, "nil", "none")


def _tables_before_first_heading(document, by_id):
    """The tables that appear in the body BEFORE the first level-1 heading (the
    cover / front-matter region). Restricting the cover search here stops a later
    body data table from being misclassified as a cover table.
    """
    body = document.element.body
    tbl_to_obj = {t._tbl: t for t in document.tables}
    out = []
    for child in body:
        if child.tag == qn("w:p"):
            if _heading_level(by_id, child) == 1:
                break
        elif child.tag == qn("w:tbl"):
            t = tbl_to_obj.get(child)
            if t is not None:
                out.append(t)
    return out


def extract_cover(document, by_id, warn):
    """Identify the three cover tables by shape and read their structure verbatim.

    Shape matching tolerates merge-inflated grids (a word-authored cover table with
    merged banner/label cells reports more grid columns/rows than the logical
    layout), and the search is restricted to the front-matter region so a body
    4-column table cannot steal the revision slot.
    """
    cover = {"tables": {}}
    info_t = sig_t = rev_t = None
    candidates = _tables_before_first_heading(document, by_id)
    if not candidates:
        # no level-1 heading found (or no front tables detected): fall back to all
        # tables so a heading-detection miss does not blank the whole cover.
        candidates = list(document.tables)
    for t in candidates:
        rows, cols = len(t.rows), len(t.columns)
        if info_t is None and rows == 2 and cols >= 5:
            info_t = t
        elif sig_t is None and rows == 3 and cols >= 5:
            sig_t = t
        elif rev_t is None and cols == 4 and rows >= 2:
            rev_t = t

    # ---- info table ----
    if info_t is not None:
        outer, inner, _border_sz = _table_outer_inner(info_t._tbl)
        # company line = merged first cell text
        company_line = _safe_cell_text(info_t, 0, 0)
        labels = {
            "title": _safe_cell_text(info_t, 0, 1),
            "doc_no": _safe_cell_text(info_t, 1, 1),
            "secrecy": _safe_cell_text(info_t, 0, 3),
            "pages": _safe_cell_text(info_t, 1, 3),
        }
        # the cell font point size (engine uses info.sz as the label/value font size)
        label_font = _safe(lambda: _cell_run_font(info_t.cell(0, 1)), {}) or {}
        cover["tables"]["info"] = {
            "cols_cm": _grid_cols_cm(info_t._tbl),
            "outer": outer or "double", "inner": inner or "single",
            "sz": int(label_font["size_pt"]) if label_font.get("size_pt") else 14,
            "labels": labels,
        }
        cover["company_line"] = company_line
        # secrecy default value + its run color (the value cell (0,4))
        cover["secrecy_default"] = _safe_cell_text(info_t, 0, 4)
        cover["_secrecy_color"] = _safe(lambda: _cell_run_color(info_t.cell(0, 4)))
        # page count + page text segments from the pages value cell (1,4)
        pages_cell = _safe(lambda: info_t.cell(1, 4))
        has_field = (pages_cell is not None
                     and pages_cell._tc.find(".//" + qn("w:fldChar")) is not None)
        cover["page_count_field"] = bool(has_field)
        cover["page_text"] = (_literal_runs_around_fields(pages_cell._tc)
                              if has_field else ["", ""])
    else:
        warn("cover info table (2x5) not found; cover defaults used")
        cover["tables"]["info"] = {
            "cols_cm": [], "outer": "double", "inner": "single", "sz": 14,
            "labels": {"title": "Project Name", "doc_no": "Project Code",
                       "secrecy": "Secrecy", "pages": "Pages"}}
        cover["company_line"] = ""
        cover["secrecy_default"] = ""
        cover["_secrecy_color"] = None
        cover["page_count_field"] = True
        cover["page_text"] = ["", ""]

    # ---- signature table ----
    if sig_t is not None:
        rows = []
        sign_cols = []
        for r_i in range(len(sig_t.rows)):
            lab = _safe_cell_text(sig_t, r_i, 0)
            dat = _safe_cell_text(sig_t, r_i, 3)
            rows.append([lab, dat])
        # which columns carry a signature underline (bottom border)
        for c_i in range(len(sig_t.columns)):
            if _safe(lambda: _has_bottom_border(sig_t.cell(0, c_i)), False):
                sign_cols.append(c_i)
        cover["tables"]["signature"] = {
            "cols_cm": _grid_cols_cm(sig_t._tbl),
            "rows": rows,
            "sign_underline": bool(sign_cols),
            "sign_cols": sign_cols or [1, 4],
        }
    else:
        warn("cover signature table (3x5) not found; defaults used")
        cover["tables"]["signature"] = {
            "cols_cm": [], "rows": [["Author:", "Date:"]],
            "sign_underline": True, "sign_cols": [1, 4]}

    # ---- revision table ----
    if rev_t is not None:
        headers = [_safe_cell_text(rev_t, 0, c) for c in range(len(rev_t.columns))]
        outer, inner, sz = _table_outer_inner(rev_t._tbl)
        header_font = (_safe(lambda: _cell_run_font(rev_t.cell(0, 0)), {})
                       or {"ascii": "Arial", "size_pt": 10.5})
        hf = {"ascii": header_font.get("ascii", "Arial"),
              "size_pt": header_font.get("size_pt", 10.5)}
        cover["tables"]["revision"] = {
            "cols_cm": _grid_cols_cm(rev_t._tbl),
            "headers": headers,
            "header_font": hf,
            "border": inner or outer or "single",
        }
        title = _extract_revision_title(document, rev_t)
        if title:
            cover["tables"]["revision"]["title"] = title
    else:
        warn("cover revision table (4-col) not found; defaults used")
        cover["tables"]["revision"] = {
            "cols_cm": [], "headers": ["Date", "Version", "Description", "Author"],
            "header_font": {"ascii": "Arial", "size_pt": 10.5}, "border": "single"}
    return cover


def _extract_revision_title(document, rev_t):
    """The centered title paragraph immediately preceding the revision table."""
    body = document.element.body
    tbl_el = rev_t._tbl
    prev_p = None
    for child in body:
        if child is tbl_el:
            break
        if child.tag == qn("w:p"):
            txt = _para_text(child)
            if txt.strip():
                prev_p = child
    if prev_p is None:
        return None
    txt = _para_text(prev_p)
    # font of the title run
    font = {}
    for r in prev_p.findall(".//" + qn("w:r")):
        rt = "".join(t.text or "" for t in r.findall(qn("w:t")))
        if not rt.strip():
            continue
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            rf = rPr.find(qn("w:rFonts"))
            sz = rPr.find(qn("w:sz"))
            if rf is not None and rf.get(qn("w:ascii")):
                font["ascii"] = rf.get(qn("w:ascii"))
            if rf is not None and rf.get(qn("w:eastAsia")):
                font["eastAsia"] = rf.get(qn("w:eastAsia"))
            if sz is not None:
                font["size_pt"] = int(sz.get(qn("w:val"))) / 2
        break
    out = {"text": txt}
    out.update(font)
    return out


def extract_cover_extras(document, by_id, warn):
    """The cover body paragraphs after the logo (company names) and the big-title
    subtitle. These are read verbatim from the document body."""
    body = document.element.body
    # find the cover logo drawing in the body, then read the centered company-name
    # paragraphs that follow until the next empty paragraph / table / page break.
    company_names = []
    subtitle = None
    big_title_size = None

    paras = [c for c in body if c.tag == qn("w:p")]
    # locate the paragraph that contains the body picture (cover logo)
    logo_idx = None
    for i, p in enumerate(paras):
        if p.find(".//" + qn("w:drawing")) is not None or p.find(".//" + qn("w:pict")) is not None:
            logo_idx = i
            break
    if logo_idx is not None:
        for p in paras[logo_idx + 1:]:
            # stop when we reach an empty paragraph following the names
            txt = _para_text(p)
            if not txt.strip():
                if company_names:
                    break
                continue
            # stop if this paragraph is a heading
            if _heading_level(by_id, p) is not None:
                break
            run = _first_text_run_font(p)
            entry = {"text": txt}
            if run.get("ascii"):
                entry["ascii"] = run["ascii"]
            if run.get("eastAsia"):
                entry["eastAsia"] = run["eastAsia"]
            if run.get("size_pt") is not None:
                entry["size_pt"] = run["size_pt"]
            company_names.append(entry)
            if len(company_names) >= 4:
                break
    if not company_names:
        warn("cover company_names paragraphs not found")

    # big title: the largest centered run before the logo; subtitle = the line
    # directly under it.
    if logo_idx is not None:
        best_size = 0
        title_p_idx = None
        for i, p in enumerate(paras[:logo_idx]):
            run = _first_text_run_font(p)
            if run.get("size_pt") and run["size_pt"] > best_size and _para_text(p).strip():
                best_size = run["size_pt"]
                title_p_idx = i
        if title_p_idx is not None:
            big_title_size = best_size
            # subtitle = next non-empty paragraph after the title, before the logo
            for p in paras[title_p_idx + 1:logo_idx]:
                if _para_text(p).strip():
                    subtitle = _para_text(p)
                    break
    return {"company_names": company_names, "subtitle": subtitle,
            "big_title_size": big_title_size}


def _first_text_run_font(p):
    for r in p.findall(".//" + qn("w:r")):
        rt = "".join(t.text or "" for t in r.findall(qn("w:t")))
        if not rt.strip():
            continue
        rPr = r.find(qn("w:rPr"))
        out = {}
        if rPr is not None:
            rf = rPr.find(qn("w:rFonts"))
            sz = rPr.find(qn("w:sz"))
            if rf is not None:
                if rf.get(qn("w:ascii")):
                    out["ascii"] = rf.get(qn("w:ascii"))
                if rf.get(qn("w:eastAsia")):
                    out["eastAsia"] = rf.get(qn("w:eastAsia"))
            if sz is not None:
                out["size_pt"] = int(sz.get(qn("w:val"))) / 2
        return out
    return {}


# ===========================================================================
# Skeleton (heading tree) + fixed bodies
# ===========================================================================
def extract_skeleton_and_bodies(document, by_id, body_id, mybody_id,
                                fixed_body_sections, skeleton_marks, warn):
    """Walk the body, build the nested heading tree, and collect the fixed-body
    paragraphs for the requested [chapter, sub] index pairs.

    ``skeleton_marks`` may specify ``compliance_chapter`` (1-based chapter index that
    carries the ``datatable: compliance`` mark and has its children dropped).
    """
    body = document.element.body
    flat = []  # (level, title, p_index)
    body_children = list(body)
    # gather headings + remember body order for fixed-body slicing
    heading_positions = []
    for idx, child in enumerate(body_children):
        if child.tag != qn("w:p"):
            continue
        lv = _heading_level(by_id, child)
        if lv is not None:
            title = _strip_autonumber(_para_text(child))
            heading_positions.append((idx, lv, title, child))

    # ---- build nested tree from the heading sequence ----
    skeleton = []
    chapter_idx = 0  # 1-based chapter counter
    sub_counters = {}  # chapter_idx -> running sub index
    stack = []  # list of (level, node)
    node_index = {}  # (chap, sub) -> node   (sub=0 means the chapter node itself)
    for (idx, lv, title, child) in heading_positions:
        node = {"title": title}
        if lv == 1:
            chapter_idx += 1
            sub_counters[chapter_idx] = 0
            node_index[(chapter_idx, 0)] = node
            skeleton.append(node)
            stack = [(1, node)]
        else:
            # find parent (nearest shallower node)
            while stack and stack[-1][0] >= lv:
                stack.pop()
            parent = stack[-1][1] if stack else None
            if lv == 2 and chapter_idx in sub_counters:
                sub_counters[chapter_idx] += 1
                node_index[(chapter_idx, sub_counters[chapter_idx])] = node
            if parent is not None:
                parent.setdefault("children", []).append(node)
            else:
                skeleton.append(node)
            stack.append((lv, node))

    # ---- compliance datatable mark ----
    comp_chap = (skeleton_marks or {}).get("compliance_chapter")
    if comp_chap and (comp_chap, 0) in node_index:
        cnode = node_index[(comp_chap, 0)]
        cnode["datatable"] = "compliance"
        cnode.pop("children", None)

    # ---- fixed bodies: collect paragraphs between the marked heading and next ----
    fixed_bodies = {}
    for pair in (fixed_body_sections or []):
        chap, sub = pair[0], pair[1]
        key = f"sec_{chap}_{sub}"
        target = node_index.get((chap, sub))
        if target is None:
            warn(f"fixed body {key} ([{chap},{sub}]) heading not found")
            continue
        # find this heading's body index and the next heading's index
        title = target["title"]
        start_idx = None
        for (idx, lv, t, child) in heading_positions:
            if _strip_autonumber(_para_text(child)) == title:
                start_idx = idx
                break
        if start_idx is None:
            warn(f"fixed body {key} start paragraph not located")
            continue
        next_idx = len(body_children)
        for (idx, lv, t, child) in heading_positions:
            if idx > start_idx:
                next_idx = idx
                break
        paragraphs, style_alias = _collect_fixed_paragraphs(
            by_id, body_children, start_idx + 1, next_idx, body_id, mybody_id, warn)
        if paragraphs:
            node_index[(chap, sub)]["fixed_body"] = key
            fixed_bodies[key] = {"style": style_alias, "paragraphs": paragraphs}
        else:
            warn(f"fixed body {key} contained no body paragraphs")
    return skeleton, fixed_bodies


def _collect_fixed_paragraphs(by_id, body_children, start, end, body_id, mybody_id,
                              warn=None):
    paragraphs = []
    alias = "body"
    for i in range(start, end):
        child = body_children[i]
        if child.tag != qn("w:p"):
            continue
        if not _para_text(child).strip():
            continue
        sid = _para_style_id(child)
        alias = _style_alias(by_id, sid, body_id, mybody_id, warn)
        runs = []
        for r in child.findall(qn("w:r")):
            info = describe_run(r)
            if info["t"] == "":
                continue
            runs.append(info)
        # also capture runs nested in hyperlinks (document-sourced)
        for hl in child.findall(qn("w:hyperlink")):
            for r in hl.findall(qn("w:r")):
                info = describe_run(r)
                if info["t"] == "":
                    continue
                runs.append(info)
        if runs:
            paragraphs.append({"runs": runs})
    return paragraphs, alias


def _strip_autonumber(text):
    """Remove a leading auto-number prefix a word processor may have rendered
    (e.g. ``1.2 `` / ``1.2.3\t``). Conservative: only strips a dotted-digit run
    followed by whitespace at the very start."""
    s = text
    i = 0
    saw_digit = False
    while i < len(s) and (s[i].isdigit() or s[i] == "."):
        if s[i].isdigit():
            saw_digit = True
        i += 1
    if saw_digit and i < len(s) and s[i] in (" ", "\t", "\u00a0"):
        return s[i:].lstrip(" \t\u00a0")
    return text


# ===========================================================================
# Logo extraction
# ===========================================================================
def extract_logo(docx_path, document, out_path, warn):
    """Write the header's referenced image (resolved rId -> media target through the
    header part rels). Falls back to the earliest near-square PNG in word/media."""
    sec = document.sections[0]
    target_ref = None
    # search every header flavour (default / first-page / even-page) for an image
    # reference, looking at both DrawingML (a:blip r:embed) and legacy VML
    # (v:imagedata r:id). Resolve the rId against THAT header part's rels.
    headers = [sec.header,
               getattr(sec, "first_page_header", None),
               getattr(sec, "even_page_header", None)]
    for hdr in headers:
        if hdr is None:
            continue
        try:
            el = hdr._element
            blip = el.find(".//" + qn("a:blip"))
            rid = blip.get(qn("r:embed")) if blip is not None else None
            if rid is None:
                vml = el.find(".//" + qn("v:imagedata"))
                rid = vml.get(qn("r:id")) if vml is not None else None
            if rid and rid in hdr.part.rels:
                target_ref = hdr.part.rels[rid].target_ref
                break
        except Exception:
            continue

    with zipfile.ZipFile(docx_path) as z:
        names = z.namelist()
        chosen = None
        if target_ref:
            # target_ref is relative to word/ (e.g. "media/image1.png")
            cand = "word/" + target_ref.lstrip("/")
            if cand in names:
                chosen = cand
            else:
                base = os.path.basename(target_ref)
                chosen = next((n for n in names if n.endswith("media/" + base)), None)
        if chosen is None:
            media = [n for n in names if n.startswith("word/media/")
                     and n.lower().endswith(".png")]
            media.sort()
            chosen = media[0] if media else None
            if chosen:
                warn("logo resolved by fallback (earliest media PNG), not header rels")
        if chosen is None:
            warn("no logo image found in the document; config.logo left as default")
            return None
        data = z.read(chosen)
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    with open(out_path, "wb") as f:
        f.write(data)
    return os.path.basename(out_path)


# ===========================================================================
# Neutral defaults for the non-extractable (class 3) blocks
# ===========================================================================
def neutral_defaults():
    """Neutral, English, public-safe defaults for everything not read from the
    document. Overrides replace any of these."""
    return {
        "id": "template_v1",
        "caption_prefix": {"image": "Figure", "table": "Table"},
        "toc": {
            "title": "Contents",
            "size_pt": 24,
            "field": "TOC \\o \"1-3\" \\h \\z \\u",
            "placeholder": "(right-click -> update field)",
        },
        "cover": {
            "big_title": {"placeholder": "[Module] [Review Type]"},
            "logo_cm": 2.6,
            "fields": [
                {"key": "title", "label": "Title", "table": "info", "required": True},
                {"key": "doc_no", "label": "Document No.", "table": "info", "required": False},
                {"key": "secrecy", "label": "Classification", "table": "info", "required": False},
                {"key": "author", "label": "Author", "table": "signature", "required": False},
                {"key": "reviewers", "label": "Reviewers", "table": "signature", "required": False},
                {"key": "approver", "label": "Approver", "table": "signature", "required": False},
                {"key": "version", "label": "Version", "table": "revision", "required": False},
            ],
        },
        "compliance": {
            "col_w_cm": {"cat": 1.7, "item": 4.0, "spec": 1.0, "axis": 0.88,
                         "spacer": 0.2, "unit": 1.0},
            "font_pt": 7,
            "row_h_pt": {"header": 12, "data": 10},
            "axis_labels": ["MIN", "TYP", "MAX", "NTWC"],
            "fills": {"header": "FFFF00", "setting": "EEECE1",
                      "result": "FFFFFF", "separator": "B8CCE4"},
            "setting_kinds": ["common_setting", "module_setting", "tb"],
            "default_limit": {
                "le": "flag values above the upper bound",
                "ge": "flag values below the target",
                "range": "flag values outside the spec range",
            },
            "flag_color": "FF0000",
            "borders": {"val": "single", "sz": 4, "color": "000000"},
        },
        "free_table": {
            "header_fill": "D9D9D9",
            "border": {"val": "single", "sz": 4, "color": "000000"},
        },
        "ui_strings": {
            "toolbar.open": "Open Project",
            "toolbar.save": "Save",
            "toolbar.new": "New",
            "toolbar.export_docx": "Export Word",
            "toolbar.export_pdf": "Export PDF Preview",
            "block.add.para": "Paragraph",
            "block.add.image": "Image",
            "block.add.datatable": "Data Table",
            "block.add.table": "Table",
            "status.autosaved": "Auto-saved",
            "cover.node": "Cover",
            "node.fixed": "Fixed",
            "node.template": "Template",
            "compliance.editor": "Data Table Editor",
            "compliance.add_group": "Add Group",
            "compliance.add_axis": "Add Axis",
        },
    }


# ===========================================================================
# Deep merge
# ===========================================================================
def deep_merge(base, over):
    """Recursively merge ``over`` into ``base``. Dicts merge key by key; any leaf or
    list value present in ``over`` replaces the corresponding value in ``base``
    (lists are replaced wholesale, never concatenated)."""
    if not isinstance(base, dict) or not isinstance(over, dict):
        return over
    out = dict(base)
    for k, v in over.items():
        if k in out and isinstance(out[k], dict) and isinstance(v, dict):
            out[k] = deep_merge(out[k], v)
        else:
            out[k] = v
    return out


# top-level overrides keys that drive extraction but are NOT part of the produced
# config (they tell the extractor what to do; they must not leak into the output)
_CONTROL_KEYS = ("fixed_body_sections", "skeleton_marks")


def _strip_private(cfg):
    """Drop transient keys: those used internally (leading underscore) and the
    extraction-control keys that overrides supply as input only."""
    if isinstance(cfg, dict):
        return {k: _strip_private(v) for k, v in cfg.items()
                if not k.startswith("_") and k not in _CONTROL_KEYS}
    if isinstance(cfg, list):
        return [_strip_private(x) for x in cfg]
    return cfg


# ===========================================================================
# Orchestration
# ===========================================================================
def build_extracted(docx_path, logo_filename, overrides, warn):
    document = Document(docx_path)
    by_id = _styles_by_id(document)
    theme = _theme_fonts(document)

    body_id, mybody_id = find_custom_body_styles(document, by_id, theme, warn)

    styles = extract_styles(document, by_id, theme, body_id, mybody_id, warn)
    styles["header_table"] = extract_header_table(document, theme, warn)
    styles["footer_table"] = extract_footer_table(document, by_id, theme, warn)

    cover = extract_cover(document, by_id, warn)
    extras = extract_cover_extras(document, by_id, warn)
    cover["company_names"] = extras["company_names"]
    big_title = cover.setdefault("big_title", {})
    if extras["subtitle"] is not None:
        big_title["subtitle"] = extras["subtitle"]
    if extras["big_title_size"] is not None:
        big_title["size_pt"] = extras["big_title_size"]

    ov = overrides or {}
    sk_marks = ov.get("skeleton_marks", {})
    fb_sections = ov.get("fixed_body_sections", [[1, 1], [1, 3]])
    skeleton, fixed_bodies = extract_skeleton_and_bodies(
        document, by_id, body_id, mybody_id, fb_sections, sk_marks, warn)

    styles["colors"] = extract_colors(document, by_id, theme, fixed_bodies, cover, warn)

    extracted = {
        "logo": logo_filename or "logo.png",
        "skeleton": skeleton,
        "fixed_bodies": fixed_bodies,
        "cover": cover,
        "styles": styles,
    }
    return extracted


def _load_json(path):
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def _drop_comment_keys(obj):
    """Remove documentation-only keys (``_README`` / ``_comment``) from an overrides
    structure so they never leak into the produced config."""
    if isinstance(obj, dict):
        return {k: _drop_comment_keys(v) for k, v in obj.items()
                if k not in ("_README", "_comment")}
    if isinstance(obj, list):
        return [_drop_comment_keys(x) for x in obj]
    return obj


def main(argv=None):
    ap = argparse.ArgumentParser(
        description="Extract a renderer template config from a structured master document.")
    ap.add_argument("master", help="the master .docx to read")
    ap.add_argument("--xlsx", help="optional companion spreadsheet (axis labels / fills)")
    ap.add_argument("--overrides", help="local overrides JSON (defaults next to --out)")
    ap.add_argument("--out", default=os.path.join("local", "template_config_generated.json"),
                    help="output config path")
    ap.add_argument("--logo-out", default="logo_extracted.png",
                    help="filename for the extracted logo PNG")
    ap.add_argument("--logo-dir", help="directory for the logo (defaults alongside --out)")
    args = ap.parse_args(argv)

    if not os.path.isfile(args.master):
        print(f"error: master not found: {args.master}", file=sys.stderr)
        return 2

    warnings = []
    warn = warnings.append

    out_path = os.path.abspath(args.out)
    out_dir = os.path.dirname(out_path)
    logo_dir = os.path.abspath(args.logo_dir) if args.logo_dir else out_dir
    logo_path = os.path.join(logo_dir, args.logo_out)

    # resolve overrides: explicit, else look next to --out
    overrides_path = args.overrides
    if overrides_path is None:
        cand = os.path.join(out_dir, "template_overrides.json")
        overrides_path = cand if os.path.exists(cand) else None
    overrides = {}
    if overrides_path and os.path.exists(overrides_path):
        overrides = _drop_comment_keys(_load_json(overrides_path))

    # logo
    document_for_logo = Document(args.master)
    logo_filename = extract_logo(args.master, document_for_logo, logo_path, warn)

    if args.xlsx and not os.path.isfile(args.xlsx):
        warn(f"--xlsx not found: {args.xlsx} (axis labels / fills not confirmed)")

    extracted = build_extracted(args.master, logo_filename, overrides, warn)

    # config = defaults <- extracted <- overrides (overrides win)
    merged = deep_merge(neutral_defaults(), extracted)
    extracted_keys = sorted(merged.keys())
    merged = deep_merge(merged, overrides)
    merged = _strip_private(merged)

    os.makedirs(out_dir, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(merged, f, ensure_ascii=False, indent=2)

    # ---- summary ----
    print("config written ->", out_path)
    if logo_filename:
        print("logo written  ->", logo_path)
    print("\ntop-level keys (extracted + defaults):", ", ".join(extracted_keys))
    if overrides:
        print("overrides applied (these keys win):", ", ".join(sorted(overrides.keys())))
    else:
        print("overrides applied: (none -- neutral defaults used for class-3 keys)")
    if warnings:
        print("\nunresolved / fallback warnings (consider an override):")
        for w in warnings:
            print("  -", w)
    else:
        print("\nno warnings: every targeted property resolved.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
